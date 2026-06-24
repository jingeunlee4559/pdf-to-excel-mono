from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Tuple

from fastapi import UploadFile

from app.services.llm_client import call_llm_json, get_llm_config
from app.services.storage_service import repair_mojibake_filename, save_upload_file, validate_storage_path
from app.services.unit_normalizer import clean_cell_text, enrich_row_units

DEFAULT_COLUMNS = [
    {"key": "vendor_name", "label": "업체명"},
    {"key": "item_name", "label": "품목명"},
    {"key": "spec", "label": "규격"},
    {"key": "quantity", "label": "수량"},
    {"key": "unit", "label": "단위"},
    {"key": "standard_unit_price", "label": "기준/표준단가"},
    {"key": "vendor_unit_price", "label": "업체 견적단가"},
    {"key": "unit_price", "label": "일반 단가"},
    {"key": "amount", "label": "금액"},
    {"key": "price_diff", "label": "차이"},
    {"key": "diff_rate", "label": "대비율"},
    {"key": "remark", "label": "비고"},
]

STANDARD_MARKET_PRICE_COLUMNS = [
    {"key": "construction_code", "label": "공종코드"},
    {"key": "item_name", "label": "공종명칭"},
    {"key": "spec", "label": "규격"},
    {"key": "unit", "label": "단위"},
    {"key": "unit_price", "label": "단가"},
    {"key": "labor_ratio", "label": "노무비율"},
    {"key": "remark", "label": "비고"},
]

REFERENCE_GUIDELINE_COLUMNS = [
    {"key": "section", "label": "구분/장절"},
    {"key": "basis_item", "label": "기준 항목"},
    {"key": "application_basis", "label": "적용 기준"},
    {"key": "calculation_method", "label": "계산/적용 방식"},
    {"key": "unit_price_basis", "label": "단가 기준"},
    {"key": "source_page", "label": "근거 페이지"},
    {"key": "remark", "label": "비고"},
]

TEXT_VENDOR_COMPARISON_COLUMNS = [
    {"key": "vendor_name", "label": "업체명"},
    {"key": "total_amount", "label": "총 견적금액"},
    {"key": "price_diff", "label": "표준단가 대비"},
    {"key": "diff_rate", "label": "대비율"},
    {"key": "remark", "label": "검토의견"},
]

TEXT_VENDOR_COMPARISON_TABLE_TYPE = "TEXT_VENDOR_COMPARISON_REPORT"

REFERENCE_TABLE_TYPES = {"REFERENCE_GUIDELINE_TABLE", "GUIDELINE_SUMMARY_TABLE"}
STANDARD_MARKET_TABLE_TYPES = {"STANDARD_MARKET_PRICE_TABLE"}

MULTI_VENDOR_COMPARE_TABLE_TYPE = "MULTI_VENDOR_PRICE_COMPARISON"


def _safe_compare_key(label: str, index: int, suffix: str = "unit_price") -> str:
    raw = compact_text(label) or f"vendor{index}"
    m = re.search(r"([A-Za-z0-9]+)회사", str(label or ""), re.I)
    if m:
        base = f"company_{m.group(1).lower()}"
    else:
        base = re.sub(r"[^A-Za-z0-9]+", "_", raw).strip("_").lower() or f"vendor_{index}"
    return f"{base}_{suffix}"


def _extract_company_name(filename: str, text: str = "", rows: List[Dict[str, Any]] | None = None, index: int = 1) -> str:
    """원문/행/파일명에서 실제 업체명을 추출한다.

    A회사/B회사 같은 비교용 별칭을 업체명으로 확정하지 않는다. 우선순위는
    업체명·회사명 라벨 → 행 필드 → 상단의 alias · 실제명 패턴 → 파일명 앞 토큰이다.
    """
    source = f"{filename}\n{text[:1800]}"
    label_patterns = [
        r"업체명\s*[:：]?\s*([^\n\r\t]{1,80})",
        r"회사명\s*[:：]?\s*([^\n\r\t]{1,80})",
        r"상호\s*[:：]?\s*([^\n\r\t]{1,80})",
        r"공급자\s*[:：]?\s*([^\n\r\t]{1,80})",
        r"견적업체\s*[:：]?\s*([^\n\r\t]{1,80})",
    ]
    for pattern in label_patterns:
        m = re.search(pattern, source)
        if m:
            value = clean_cell_text(m.group(1))
            value = re.split(r"\s{2,}|\t|사업자|대표|견적일자|주소|전화|연락처|TEL|Fax|FAX", value)[0].strip()
            if value:
                return value[:40]

    for row in rows or []:
        value = clean_cell_text(row.get("vendor_name") or row.get("company_name") or "")
        if value:
            return value[:40]

    # 예: "A회사 · ㈜에이건설", "1안 : ○○건설". alias 자체가 아니라 separator 뒤 실제명을 사용한다.
    for line in [line.strip() for line in source.splitlines()[:30] if line.strip()]:
        m = re.search(r"[^\s·:\-]{1,30}\s*[·:\-]\s*([㈜\(\)가-힣A-Za-z0-9\s&._-]{2,70})", line)
        if m:
            value = clean_cell_text(m.group(1))
            if value and not any(skip in compact_text(value) for skip in ["건설공사견적서", "표준시장단가비교용"]):
                return value[:40]

    stem = Path(filename or f"업체{index}").stem.strip()
    parts = [clean_cell_text(part) for part in re.split(r"[_\-]+", stem) if clean_cell_text(part)]
    if parts:
        return parts[0][:40]
    return stem[:40] or f"업체{index}"


def _request_wants_company_comparison(user_request: str, llm_intent: Dict[str, Any] | None = None) -> bool:
    request = compact_text(user_request)
    intent_name = ""
    if isinstance(llm_intent, dict):
        intent_name = str(llm_intent.get("intent") or "").strip().upper()
    if intent_name == "COMPANY_COMPARISON":
        return True
    comparison_terms = ["회사별", "업체별", "거래처별", "견적비교", "단가비교", "비교표"]
    if any(term in request for term in comparison_terms):
        return True
    if "비교" in request and any(term in request for term in ["회사", "업체", "견적", "단가", "가격", "금액"]):
        return True
    # A회사/B회사/C회사 또는 실제 건설사명을 지정하고 단가/수량/표 생성을 요청한 경우도 업체 비교로 본다.
    company_mentioned = bool(re.search(r"[A-Za-z]\s*회사", str(user_request or ""), re.I))
    company_mentioned = company_mentioned or any(term in request for term in ["에이건설", "비테크건설", "씨엔씨종합건설", "대동종합건설", "이엔지건설", "건설㈜", "종합건설"])
    return company_mentioned and any(term in request for term in ["단가", "견적", "표", "수량", "기준", "금액", "가격"])


def _request_wants_standard_price(user_request: str, llm_intent: Dict[str, Any] | None = None) -> bool:
    """표준시장단가 컬럼 표시 여부를 판단한다.

    원본 PDF에 표준시장단가 컬럼이 있어도, 사용자가 명시적으로 요청하지 않으면
    업체 단가 비교 화면/엑셀에는 표시하지 않는다.
    "단가 기준", "비교 기준" 같은 일반 표현은 표준시장단가 요청으로 보지 않는다.
    """
    request = compact_text(user_request)
    explicit_terms = [
        "표준시장단가",
        "표준시장",
        "표준단가",
        "기준단가",
        "시장단가",
        "표준가",
    ]
    if any(term in request for term in explicit_terms):
        return True
    # LLM 판단은 보조 정보다. 원문/파일명에 표준시장단가가 들어 있다는 이유만으로
    # requiresStandardPrice=true가 되는 경우가 있어 사용자 요청 문구를 우선한다.
    if isinstance(llm_intent, dict) and llm_intent.get("requiresStandardPrice") is True:
        return any(term in request for term in explicit_terms)
    return False


def _prune_standard_price_from_compare_table(table: Dict[str, Any], include_standard_price: bool) -> Dict[str, Any]:
    """업체 비교표에서 표준시장단가 관련 표시 컬럼을 조건부로 제거한다."""
    if include_standard_price or not isinstance(table, dict):
        return table
    hide_keys = {"standard_unit_price", "lowest_vs_standard"}
    next_table = dict(table)
    next_table["columns"] = [col for col in (table.get("columns") or []) if col.get("key") not in hide_keys]
    next_rows = []
    for row in table.get("rows") or []:
        if not isinstance(row, dict):
            next_rows.append(row)
            continue
        next_row = {k: v for k, v in row.items() if k not in hide_keys}
        next_rows.append(next_row)
    next_table["rows"] = next_rows
    meta = dict(next_table.get("meta") or {})
    meta["standardPriceHidden"] = True
    meta["standardPriceDisplayPolicy"] = "사용자가 표준시장단가/기준단가를 명시 요청할 때만 표시"
    next_table["meta"] = meta
    return next_table


def _row_has_price_value(row: Dict[str, Any]) -> bool:
    return any(to_number(row.get(key)) > 0 for key in ["vendor_unit_price", "unit_price", "amount", "standard_unit_price"])


def _get_vendor_price_text(row: Dict[str, Any]) -> str:
    """업체 비교에 사용할 가격을 가져온다. standard_unit_price는 업체 가격으로 사용하지 않는다."""
    for key in PRICE_VALUE_PRIORITY:
        value = clean_number(row.get(key) or "")
        if value:
            return value
    return ""


def _is_standard_market_file(filename: str, text: str, rows: List[Dict[str, Any]] | None = None, user_request: str = "", llm_intent: Dict[str, Any] | None = None) -> bool:
    rows = rows or []
    compact = compact_text(f"{filename}\n{text[:5000]}")
    standard_markers = sum(1 for marker in ["건설공사표준시장단가", "표준시장단가", "공종코드", "공종명칭", "노무비율"] if compact_text(marker) in compact)
    estimate_markers = sum(1 for marker in ["견적서", "견적", "업체명", "회사명", "공급가액", "합계금액"] if compact_text(marker) in compact)
    price_rows = sum(1 for row in rows if isinstance(row, dict) and clean_cell_text(row.get("item_name") or row.get("construction_code") or "") and _row_has_price_value(row))

    if _request_wants_company_comparison(user_request, llm_intent) and price_rows >= 2:
        # 비교 요청에서는 표준시장단가 문구가 포함된 견적서를 기준자료로 오분류하지 않는다.
        return standard_markers >= 4 and estimate_markers == 0
    return standard_markers >= 3 and estimate_markers == 0


def _is_estimate_file(filename: str, text: str, rows: List[Dict[str, Any]], user_request: str = "", llm_intent: Dict[str, Any] | None = None) -> bool:
    compact = compact_text(f"{filename}\n{text[:4000]}")
    has_estimate_word = any(word in compact for word in ["견적서", "견적", "업체명", "회사명", "합계금액", "공사견적"])
    has_price_rows = sum(1 for row in rows if isinstance(row, dict) and _row_has_price_value(row) and str(row.get("item_name") or row.get("construction_code") or "").strip()) >= 2
    if _request_wants_company_comparison(user_request, llm_intent) and has_price_rows:
        return True
    if _is_standard_market_file(filename, text, rows, user_request=user_request, llm_intent=llm_intent):
        return False
    return bool(has_estimate_word or has_price_rows)


def _row_match_key(row: Dict[str, Any]) -> str:
    # 같은 공종코드가 여러 품목/규격에 반복되는 견적서가 있다.
    # code만 키로 쓰면 P.P마대(쌓기), P.P마대(헐기), 톤마대(쌓기) 같은 행이 한 행으로 합쳐져
    # 최초 요청/후속 추가 요청에서 품목이 사라진다.
    code = compact_text(row.get("construction_code"))
    item = compact_text(row.get("item_name"))
    spec = compact_text(row.get("spec"))
    unit = compact_text(row.get("unit_normalized") or row.get("unit"))
    if item or spec or unit:
        return f"code:{code}|item:{item}|spec:{spec}|unit:{unit}" if code else f"text:{item}|{spec}|{unit}"
    return f"code:{code}" if code else "text:||"


def _row_relaxed_key(row: Dict[str, Any]) -> str:
    return f"{compact_text(row.get('item_name'))}|{compact_text(row.get('spec'))}"


def _row_item_key(row: Dict[str, Any]) -> str:
    return compact_text(row.get("item_name"))


def _find_reference_row(company_row: Dict[str, Any], reference_rows: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    if not reference_rows:
        return None
    code = compact_text(company_row.get("construction_code"))
    if code:
        for ref in reference_rows:
            if compact_text(ref.get("construction_code")) == code:
                return ref
    relaxed = _row_relaxed_key(company_row)
    if relaxed != "|":
        for ref in reference_rows:
            if _row_relaxed_key(ref) == relaxed:
                return ref
    item = _row_item_key(company_row)
    unit = clean_cell_text(company_row.get("unit_normalized") or company_row.get("unit"))
    if item:
        candidates = [ref for ref in reference_rows if _row_item_key(ref) == item]
        if unit:
            same_unit = [ref for ref in candidates if clean_cell_text(ref.get("unit_normalized") or ref.get("unit")) == unit]
            if same_unit:
                return same_unit[0]
        if candidates:
            return candidates[0]
    return None


def _extract_focus_terms(
    user_request: str,
    company_rows: List[Dict[str, Any]],
    reference_rows: List[Dict[str, Any]],
    llm_terms: List[str] | None = None,
) -> List[str]:
    """사용자 요청에서 비교 대상 품목을 추출한다.

    원칙:
    - 문서에서 실제 추출된 품목명이 사용자 문장에 그대로 있으면 그 정확한 품목명을 최우선으로 사용한다.
    - 괄호가 열린 채 끝나는 부분어(예: P.P마대()는 다른 품목까지 오염시키므로 버린다.
    - LLM targetKeywords는 원문/요청 근거가 있을 때만 보조로 사용한다.
    """
    raw_request = str(user_request or "")
    req = compact_text(raw_request)
    if not req:
        return []

    row_names = []
    seen_names = set()
    for row in [*company_rows, *reference_rows]:
        name = clean_cell_text(row.get("item_name") or "")
        key = compact_text(name)
        if key and key not in seen_names:
            row_names.append((name, key))
            seen_names.add(key)

    # 1) 원문 표에 존재하는 품목명이 사용자 요청에 그대로 있으면 정확 매칭만 사용한다.
    exact_terms = [key for _name, key in row_names if len(key) >= 2 and key in req and _valid_focus_term(key)]
    if exact_terms:
        return sorted(set(exact_terms), key=len, reverse=True)

    candidates: List[str] = []

    def add_candidate(value: Any) -> None:
        term = compact_text(value)
        if not _valid_focus_term(term):
            return
        if term in GENERIC_FOCUS_TERMS:
            return
        if term not in req:
            if not any(term in name_key or name_key in term for _name, name_key in row_names if name_key):
                return
        candidates.append(term)

    # 2) LLM이 뽑은 키워드는 보조로만 사용한다.
    for term in llm_terms or []:
        add_candidate(term)

    # 3) 쉼표/줄바꿈으로 직접 나열된 표현 보강.
    for token in re.split(r"[,，/\\n]+", raw_request):
        token = token.strip()
        for m in re.finditer(r"([가-힣A-Za-z0-9·./_\-]+(?:\([^\)]+\))?)", token):
            add_candidate(m.group(1))

    # 4) 조사/명령어 앞 표현 보강.
    stop_words = {
        "표", "표로", "비교", "비교표", "파일", "각각", "다시", "만들어", "만들어서",
        "기준", "회사", "업체", "자료", "문서", "보고", "봐서", "해줘", "해줄래", "정리",
        "추가", "이것도", "저것도", "나오게", "보여줘", "첨부", "분석", "업체별", "회사별",
        "단가", "견적", "견적서", "비교견적서", "자사양식", "엑셀"
    }
    for m in re.finditer(r"([가-힣A-Za-z0-9·./()_\-]{2,40})(?:만|를|을|로|기준|비교|표|나오게|보여줘|추가|포함|같이|함께)", raw_request):
        term = compact_text(m.group(1))
        if term not in stop_words and not any(sw in term for sw in ["파일보고", "다시비교", "표로만들"]):
            add_candidate(term)

    # 5) 후보 중 실제 품목명과 정확히 일치하는 후보가 있으면 부분 후보는 제거한다.
    row_key_set = {key for _name, key in row_names}
    exact_again = [term for term in candidates if term in row_key_set]
    if exact_again:
        return sorted(set(exact_again), key=len, reverse=True)

    return sorted(set(candidates), key=len, reverse=True)

def _filter_rows_by_focus(rows: List[Dict[str, Any]], focus_terms: List[str]) -> List[Dict[str, Any]]:
    terms = [term for term in (focus_terms or []) if _valid_focus_term(term)]
    if not terms:
        return rows

    row_item_keys = {compact_text(row.get("item_name")) for row in rows if compact_text(row.get("item_name"))}
    exact_terms = {term for term in terms if term in row_item_keys}

    filtered: List[Dict[str, Any]] = []
    for row in rows:
        item = compact_text(row.get("item_name"))
        code = compact_text(row.get("construction_code"))
        spec = compact_text(row.get("spec"))
        haystack = compact_text(" ".join([str(row.get("item_name") or ""), str(row.get("construction_code") or ""), str(row.get("spec") or "")]))
        if exact_terms:
            if item in exact_terms or code in exact_terms:
                filtered.append(row)
        elif any(term in haystack or haystack in term or term == spec for term in terms):
            filtered.append(row)
    return filtered

def _format_price_diff(company_price: float, standard_price: float) -> str:
    if not company_price or not standard_price:
        return ""
    diff = company_price - standard_price
    pct = (diff / standard_price) * 100 if standard_price else 0
    sign = "+" if diff > 0 else ""
    return f"{sign}{diff:,.0f} ({sign}{pct:.1f}%)"


def _extract_requested_quantity(user_request: str = "", llm_intent: Dict[str, Any] | None = None) -> Tuple[str, str]:
    """사용자가 지정한 수량을 추출한다.

    "4개 업체"의 4는 회사/파일 개수이지 산출 수량이 아니므로 제외하고,
    "각 50개씩", "수량 50개", "50개 기준" 같은 표현을 우선한다.
    """
    text = str(user_request or "")

    def valid_match(match: re.Match) -> Tuple[str, str] | None:
        qty = clean_number(match.group("qty"))
        unit = str(match.groupdict().get("unit") or "").strip()
        if not qty:
            return None
        after = text[match.end(): match.end() + 12]
        before = text[max(0, match.start() - 8): match.start()]
        if re.match(r"\s*(업체|회사|파일|문서|자료|개사)", after):
            return None
        if re.search(r"(업체|회사|파일|문서)\s*$", before):
            return None
        return qty, unit

    # 명시적 수량 표현을 LLM보다 우선한다. LLM이 '4개 업체'의 4를 qty로 잘못 줄 수 있기 때문이다.
    patterns = [
        r"각\s*(?P<qty>[0-9][0-9,]*(?:\.[0-9]+)?)\s*(?P<unit>개|EA|ea|㎡|m2|m²|㎥|m3|m³|톤|ton|kg|KG|대|명|식|시간|일|세트|SET)?\s*씩?",
        r"(?P<qty>[0-9][0-9,]*(?:\.[0-9]+)?)\s*(?P<unit>개|EA|ea|㎡|m2|m²|㎥|m3|m³|톤|ton|kg|KG|대|명|식|시간|일|세트|SET)\s*씩",
        r"수량\s*[:=]?\s*(?P<qty>[0-9][0-9,]*(?:\.[0-9]+)?)\s*(?P<unit>개|EA|ea|㎡|m2|m²|㎥|m3|m³|톤|ton|kg|KG|대|명|식|시간|일|세트|SET)?",
        r"(?P<qty>[0-9][0-9,]*(?:\.[0-9]+)?)\s*(?P<unit>개|EA|ea|㎡|m2|m²|㎥|m3|m³|톤|ton|kg|KG|대|명|식|시간|일|세트|SET)?\s*(?:기준|으로|만큼|수량)",
        r"(?P<qty>[0-9][0-9,]*(?:\.[0-9]+)?)\s*(?P<unit>개|EA|ea|㎡|m2|m²|㎥|m3|m³|톤|ton|kg|KG|대|명|식|시간|일|세트|SET)"
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            found = valid_match(m)
            if found:
                return found

    if isinstance(llm_intent, dict):
        for key in ("quantity", "requestedQuantity", "requested_quantity", "qty"):
            raw = llm_intent.get(key)
            if isinstance(raw, dict):
                value = raw.get("value") or raw.get("quantity") or raw.get("qty")
                unit = str(raw.get("unit") or "").strip()
                if value not in (None, ""):
                    return clean_number(value), unit
            elif raw not in (None, ""):
                return clean_number(raw), ""
    return "", ""


def _dedupe_vendor_sources(vendor_sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """같은 회사가 이전 작업 파일과 새 파일에 중복 포함되면 한 회사로 합친다.

    같은 공종 키는 뒤에 들어온 파일의 값을 우선한다. 백엔드가 기존 파일을 같이 넘겨도
    회사 컬럼이 중복 생성되지 않도록 하는 안전장치다.
    """
    ordered: List[str] = []
    by_key: Dict[str, Dict[str, Any]] = {}
    for source in vendor_sources or []:
        key = _company_compare_key(source.get("name")) or compact_text(source.get("name"))
        if not key:
            continue
        if key not in by_key:
            ordered.append(key)
            by_key[key] = {**source, "rows": []}
        current = by_key[key]
        current["name"] = source.get("name") or current.get("name")
        current["filename"] = source.get("filename") or current.get("filename")
        row_map = {_row_match_key(row): row for row in current.get("rows") or []}
        for row in source.get("rows") or []:
            row_map[_row_match_key(row)] = row
        current["rows"] = list(row_map.values())
    return [by_key[key] for key in ordered]



def _clean_item_display(value: Any) -> str:
    """PDF 셀 줄바꿈으로 분리된 공종명을 화면/엑셀에 보기 좋게 정리한다."""
    text = clean_cell_text(value)
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s*([()])\s*", r"\1", text)
    # pdfplumber가 `강 관 동 바 리`, `톤 마 대`처럼 한글을 글자 단위로 분리한 경우 복구한다.
    text = re.sub(r"(?<=[가-힣])\s+(?=[가-힣])", "", text)
    text = re.sub(r"(?<=P\.P)\s+(?=마대)", "", text, flags=re.I)
    return text.strip()


def _strip_company_alias_prefix(value: Any) -> str:
    """A회사/B회사 같은 비교용 별칭을 화면 표시·검색용 이름에서 제거한다."""
    text = clean_cell_text(value)
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"^\s*[A-Za-z]\s*회사\s*[·ㆍ:：\-–—]*\s*", "", text, flags=re.I)
    # pdfplumber가 `A회사㈜한국전기`처럼 붙여 읽는 경우 보정
    text = re.sub(r"^\s*[A-Za-z]\s*회사(?=㈜|\(주\)|주식회사|[가-힣A-Za-z0-9])\s*", "", text, flags=re.I)
    return text.strip()


def _clean_vendor_label_display(value: Any) -> str:
    text = clean_cell_text(value)
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"([A-Za-z])\s*회사", r"\1회사", text)
    text = re.sub(r"㈜\s+", "㈜", text)
    text = re.sub(r"(?<=[가-힣])\s+(?=[가-힣])", "", text)
    text = re.sub(r"([A-Za-z]회사)(?=[가-힣㈜])", r"\1 ", text)
    text = _strip_company_alias_prefix(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip()

def _has_vendor_like_price_columns(row: Dict[str, Any]) -> bool:
    return bool(_collect_dynamic_vendor_columns([row]))


def _vendor_column_canonical_key(label: Any, key: Any) -> str:
    """업체 컬럼을 원본 헤더 기준으로 묶는다.

    특정 업체명을 코드에 넣지 않는다. A회사/B회사 같은 문서 내 별칭이 있으면 그 별칭으로
    같은 업체 컬럼(단가/금액 alias)을 묶고, 별칭이 없으면 실제 업체명으로 묶는다.
    """
    joined = f"{label or ''} {key or ''}"
    alias = re.search(r"([A-Za-z])\s*회사", str(joined or ""), re.I)
    if alias:
        return f"company_{alias.group(1).lower()}"
    display = _clean_vendor_label_display(joined)
    canonical = compact_text(display)
    canonical = re.sub(r"(?:단가|금액|견적가|견적단가|업체견적단가)$", "", canonical)
    return canonical or compact_text(key or "")

def _vendor_column_sort_key(canonical: str, first_index: int) -> Tuple[int, int]:
    m = re.match(r"company_([a-z])$", canonical or "")
    if m:
        return (ord(m.group(1)) - ord("a"), first_index)
    return (1000 + first_index, first_index)


def _collect_dynamic_vendor_columns(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """한 파일 안의 업체 단가 컬럼을 감지하고, 깨진 중복 컬럼은 같은 업체로 묶는다."""
    reserved = {
        "no", "no.", "number", "construction_code", "item_name", "spec", "quantity", "unit",
        "standard_unit_price", "vendor_unit_price", "unit_price", "amount", "price_diff", "diff_rate",
        "labor_ratio", "remark", "unit_original", "unit_normalized", "unit_group", "unit_confidence",
        "request_quantity", "field", "이", "e회이엔설",
    }
    groups: Dict[str, Dict[str, Any]] = {}
    column_position = 0
    for row_index, row in enumerate(rows or []):
        if not isinstance(row, dict):
            continue
        source_headers = row.get("_source_headers") if isinstance(row.get("_source_headers"), dict) else {}
        for key, value in row.items():
            if key.startswith("_") or key in reserved or re.sub(r"_\d+$", "", key) in reserved:
                continue
            if to_number(value) <= 0:
                continue
            header = _clean_vendor_label_display(source_headers.get(key) or key)
            compact_header = compact_text(header)
            compact_key = compact_text(key)
            # 특정 업체명/업종명을 하드코딩하지 않고, 헤더가 비어 있지 않은 숫자 동적 컬럼을 업체 후보로 본다.
            # 표준단가/최저/합계 등은 reserved/semantic stopword에서 제외한다.
            semantic_stop = ["표준", "기준", "최저", "최고", "합계", "소계", "수량", "단위", "규격", "공종", "코드", "노무", "차이", "대비", "금액합계"]
            looks_like_vendor = bool(compact_header or compact_key)
            if any(stop in compact_header for stop in semantic_stop) or any(stop in compact_key for stop in semantic_stop):
                looks_like_vendor = False
            if not looks_like_vendor:
                continue
            canonical = _vendor_column_canonical_key(header, key)
            if not canonical:
                continue
            if canonical not in groups:
                groups[canonical] = {
                    "canonical": canonical,
                    "key": key,
                    "label": header or key,
                    "aliases": [],
                    "valueCount": 0,
                    "firstIndex": column_position,
                    "labelScore": 0,
                }
                column_position += 1
            group = groups[canonical]
            if key not in group["aliases"]:
                group["aliases"].append(key)
            group["valueCount"] += 1
            # 더 온전한 헤더를 대표 라벨로 선택한다. A회사/B회사 표기는 최우선이다.
            score = 0
            if re.search(r"[A-Za-z]\s*회사", header, re.I):
                score += 100
            if "㈜" in header or "건설" in header:
                score += 20
            score += min(len(compact_header), 40)
            if score > group.get("labelScore", 0):
                group["label"] = header or group["label"]
                group["key"] = key
                group["labelScore"] = score

    result = list(groups.values())
    # A회사/B회사 같은 비교표 별칭이 존재하는 문서는 그 별칭 그룹만 업체 컬럼으로 사용한다.
    # pdfplumber가 헤더 줄 일부를 별도 숫자 컬럼처럼 읽어 `사지건` 같은 가짜 업체 후보가 생기는 경우를 방지한다.
    has_alias_groups = any(re.match(r"company_[a-z]$", str(item.get("canonical") or ""), re.I) for item in result)
    if has_alias_groups:
        result = [item for item in result if re.match(r"company_[a-z]$", str(item.get("canonical") or ""), re.I)]
    # 값이 지나치게 적은 깨진 컬럼은 같은 canonical에 흡수되었으면 alias로만 남고 대표 컬럼으로는 나오지 않는다.
    result.sort(key=lambda item: _vendor_column_sort_key(str(item.get("canonical") or ""), int(item.get("firstIndex") or 0)))
    return result

def _vendor_label_search_terms(label: Any, key: Any = "", aliases: List[Any] | None = None) -> List[str]:
    """원본 업체 컬럼 라벨에서 검색 가능한 토큰을 만든다.

    특정 업체명은 코드에 넣지 않는다. 실제 매칭은 사용자가 입력한 문자열과
    원본 컬럼 라벨/키/alias에서 생성한 토큰의 포함 관계로만 판단한다.
    """
    raw_values = [str(label or ""), str(key or "")]
    raw_values.extend([str(item or "") for item in (aliases or [])])
    pieces: List[str] = []
    stopwords = {
        "회사", "업체", "단가", "금액", "견적", "견적가", "견적단가", "업체견적단가",
        "표준", "기준", "최저", "최고", "요청", "수량", "단위", "규격", "공종", "코드", "일반",
        "전기", "설비", "기술", "건설", "종합", "시스템", "이엔지", "엔지니어링", "주식회사",
    }
    for raw in raw_values:
        if not raw:
            continue
        cleaned = clean_cell_text(raw)
        cleaned = re.sub(r"\s*(단가|금액|견적가|견적단가|업체견적단가)$", "", cleaned)
        cleaned = _strip_company_alias_prefix(cleaned)
        no_corp = re.sub(r"주식회사|\(주\)|㈜|（주）|유한회사|합자회사|합명회사", " ", cleaned)
        # 전체명과 법인표기 제거명을 모두 보존한다.
        for candidate in [cleaned, no_corp]:
            token = compact_text(candidate)
            if token and len(token) >= 2 and token not in stopwords:
                pieces.append(token)
        for piece in re.split(r"[\s·ㆍ/()\[\]{},._\-]+", no_corp):
            token = compact_text(piece)
            if token and len(token) >= 2 and token not in stopwords:
                pieces.append(token)
        alias = re.search(r"([A-Za-z])\s*회사", raw, re.I)
        if alias:
            pieces.append(compact_text(alias.group(0)))
    return list(dict.fromkeys(pieces))


def _requested_vendor_count(user_request: str) -> int | None:
    """사용자 요청에서 '회사 3개', '3개 업체' 같은 업체 수만 추출한다."""
    text = str(user_request or "")
    patterns = [
        r"(?:회사|업체|거래처)\s*(?P<count>[0-9]{1,2})\s*개",
        r"(?P<count>[0-9]{1,2})\s*개\s*(?:회사|업체|거래처)",
        r"(?P<count>[0-9]{1,2})\s*(?:개사|개 업체|개 회사)",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            try:
                count = int(m.group("count"))
                if 1 <= count <= 50:
                    return count
            except Exception:
                pass
    korean_counts = {"한": 1, "두": 2, "세": 3, "네": 4, "다섯": 5}
    for word, count in korean_counts.items():
        if re.search(fr"{word}\s*(?:개\s*)?(?:회사|업체|거래처)", text):
            return count
    return None



def _request_vendor_terms_from_message(user_request: str) -> List[str]:
    """사용자 요청에서 업체명으로 보이는 후보를 추출한다.

    고정 업체명 사전은 사용하지 않는다. 쉼표/조사/명령어를 제거한 뒤
    원본 업체 컬럼명과 대조할 후보만 만든다.
    """
    text = str(user_request or "")
    if not text.strip():
        return []
    cleaned = re.sub(r"(비교해서|비교해줘|비교|단가|금액|수량|각각|기준|표|정리|만들어줘|해줄래|해줘|보여줘|보여줄래|추가|제외|빼줘|삭제|제거|회사|업체|거래처|개사|개씩|개)", " ", text)
    cleaned = re.sub(r"[0-9]+\s*(?:개|개씩|개사|회사|업체)?", " ", cleaned)
    parts = re.split(r"[,，/\n\t]+|\s{2,}", cleaned)
    tokens: List[str] = []
    stop = {"이", "그", "저", "현재", "문서", "파일", "첨부", "전체", "항목", "공종", "기준", "으로", "로", "만", "그리고", "랑", "와", "과", "및", "기중차단기", "몰드변압기", "설치"}
    for part in parts:
        part = part.strip()
        if not part:
            continue
        # 공백으로 이어 쓴 업체명도 살리되, 일반 명령어 조각은 제거한다.
        for piece in re.split(r"\s+(?:랑|와|과|및)\s+|\s+", part):
            token = compact_text(piece)
            if token and len(token) >= 2 and token not in stop and not _looks_like_work_item_term(token):
                tokens.append(token)
        joined = compact_text(part)
        if joined and len(joined) >= 3 and joined not in stop and not _looks_like_work_item_term(joined):
            tokens.append(joined)
    # A회사/B회사 같은 별칭도 보존
    for m in re.finditer(r"[A-Za-z]\s*회사", text):
        tokens.append(compact_text(m.group(0)))
    return list(dict.fromkeys(tokens))



_TEXT_WORK_ITEM_HINTS = ("설치", "포설", "배선", "조립", "제작", "철거", "교체", "시공", "공사", "공종", "항목")


def _looks_like_work_item_term(term: str) -> bool:
    token = compact_text(term)
    if not token:
        return False
    return any(hint in token for hint in _TEXT_WORK_ITEM_HINTS)


def _strip_corp_prefix(name: str) -> str:
    value = clean_cell_text(name)
    value = re.sub(r"^(?:㈜|\(주\)|주식회사)\s*", "", value)
    value = re.sub(r"\s*(?:㈜)$", "", value)
    return value.strip()


def _text_report_vendor_match_terms(vendor: Dict[str, Any]) -> List[str]:
    name = clean_cell_text(vendor.get('name'))
    alias = clean_cell_text(vendor.get('alias'))
    stripped = _strip_corp_prefix(name)
    terms = [compact_text(name), compact_text(stripped), compact_text(alias)]
    # 문장 안에서는 'B회사 ㈜대한전기설비'처럼 나오고, 사용자 요청은 보통 '대한전기설비'처럼 입력된다.
    if stripped and name != stripped:
        terms.append(compact_text(f"㈜{stripped}"))
        terms.append(compact_text(f"(주){stripped}"))
    return [term for term in dict.fromkeys(terms) if term and len(term) >= 2]


def _request_has_explicit_vendor_mention(vendor_columns: List[Dict[str, Any]], user_request: str) -> bool:
    compact_request = compact_text(user_request)
    if not compact_request:
        return False
    if _request_vendor_terms_from_message(user_request):
        return True
    for col in vendor_columns or []:
        terms = _vendor_label_search_terms(col.get("label"), col.get("key"), col.get("aliases") or [])
        for term in terms:
            if term and len(term) >= 2 and term in compact_request:
                return True
    return False


def _vendor_column_matches_request(col: Dict[str, Any], request_terms: List[str], compact_request: str) -> bool:
    terms = _vendor_label_search_terms(col.get("label"), col.get("key"), col.get("aliases") or [])
    for term in terms:
        if not term or len(term) < 2:
            continue
        if compact_request and term in compact_request:
            return True
        for req in request_terms:
            if not req or len(req) < 2:
                continue
            if term in req or req in term:
                return True
    return False


def _filter_vendor_columns_by_request(vendor_columns: List[Dict[str, str]], user_request: str) -> List[Dict[str, str]]:
    """사용자가 입력한 업체명과 원본 컬럼 라벨을 대조해 컬럼을 선택한다.

    - 사용자가 업체명을 직접 입력하면 그 업체만 선택한다.
    - '2개 회사' 같은 수량은 업체명이 없을 때만 좌측부터 N개 선택한다.
    - 업체명 매칭이 일부만 성공해도 임의 업체로 채우지 않는다.
    - 업체명을 입력했는데 매칭이 0개이면 전체 컬럼을 반환하지 않는다.
    """
    request = str(user_request or "")
    compact_request = compact_text(request)
    exclude_fragments = re.split(r"(?:말고|빼고|제외하고|제외|빼줘|삭제|제거)", request, maxsplit=1)
    compact_exclude = compact_text(exclude_fragments[0]) if len(exclude_fragments) > 1 else ""
    requested_count = _requested_vendor_count(request)
    request_terms = _request_vendor_terms_from_message(request)
    selected: List[Dict[str, str]] = []

    for col in vendor_columns or []:
        terms = _vendor_label_search_terms(col.get("label"), col.get("key"), col.get("aliases") or [])
        if compact_exclude and any(term and term in compact_exclude for term in terms):
            continue
        if _vendor_column_matches_request(col, request_terms, compact_request):
            selected.append(col)

    if selected:
        deduped: List[Dict[str, str]] = []
        seen = set()
        for col in selected:
            ident = col.get("canonical") or col.get("key") or col.get("label")
            if ident in seen:
                continue
            seen.add(ident)
            deduped.append(col)
        return deduped[:requested_count] if requested_count else deduped

    has_explicit_vendor_text = bool(request_terms)
    if requested_count and not has_explicit_vendor_text:
        return vendor_columns[:requested_count]
    if has_explicit_vendor_text:
        return []
    return vendor_columns


def _requested_company_terms(user_request: str) -> List[str]:
    """디버그/메타 표시용: 사용자가 직접 입력한 업체명 후보를 원문 기준으로 보존한다."""
    text = str(user_request or "")
    tokens = []
    cleaned = re.sub(r"(비교|단가|금액|수량|회사|업체|개|각각|으로|로|만|해줘|보여줘|추가|제외|빼줘|표|만들어줘)", " ", text)
    for piece in re.split(r"[\s,，/]+", cleaned):
        token = compact_text(piece)
        if token and len(token) >= 2:
            tokens.append(token)
    return list(dict.fromkeys(tokens))

def _get_dynamic_vendor_price(row: Dict[str, Any], col: Dict[str, Any]) -> str:
    """동적 업체 컬럼의 대표 key/alias key에서 단가를 가져온다."""
    keys = []
    key = col.get("key")
    if key:
        keys.append(str(key))
    for alias in col.get("aliases") or []:
        alias_key = str(alias or "")
        if alias_key and alias_key not in keys:
            keys.append(alias_key)
    for candidate_key in keys:
        value = clean_number(row.get(candidate_key) or "")
        if value:
            return value
    return ""


def _repair_page_split_item_names(rows: List[Dict[str, Any]], vendor_columns: List[Dict[str, str]]) -> List[Dict[str, Any]]:
    """페이지 경계 때문에 `톤마` / `대(만들기)`가 두 행으로 분리된 경우 병합한다."""
    repaired: List[Dict[str, Any]] = []
    idx = 0
    while idx < len(rows):
        row = dict(rows[idx])
        next_row = rows[idx + 1] if idx + 1 < len(rows) and isinstance(rows[idx + 1], dict) else None
        if next_row:
            current_item = _clean_item_display(row.get("item_name"))
            next_item = _clean_item_display(next_row.get("item_name"))
            current_has_prices = any(to_number(_get_dynamic_vendor_price(row, col)) > 0 for col in vendor_columns) or to_number(row.get("standard_unit_price")) > 0
            next_has_code_or_prices = bool(clean_cell_text(next_row.get("construction_code"))) or any(to_number(_get_dynamic_vendor_price(next_row, col)) > 0 for col in vendor_columns) or to_number(next_row.get("standard_unit_price")) > 0
            if current_has_prices and current_item and next_item and not next_has_code_or_prices:
                # 주로 표가 페이지를 넘어갈 때 앞 행의 공종명이 잘린 케이스다.
                if next_item.startswith("대(") or next_item.startswith("(") or compact_text(current_item + next_item) in compact_text(f"{current_item}{next_item}"):
                    row["item_name"] = _clean_item_display(f"{current_item} {next_item}")
                    idx += 2
                    repaired.append(row)
                    continue
        if row.get("item_name"):
            row["item_name"] = _clean_item_display(row.get("item_name"))
        repaired.append(row)
        idx += 1
    return repaired


def build_single_file_multi_vendor_price_comparison(
    parsed_files: List[Dict[str, Any]],
    user_request: str = "",
    llm_intent: Dict[str, Any] | None = None,
) -> Dict[str, Any] | None:
    """단일 PDF/엑셀 안에 여러 업체 단가 컬럼이 가로로 들어간 비교표를 생성한다."""
    if not parsed_files:
        return None
    request = str(user_request or "")
    if not _request_wants_company_comparison(request, llm_intent):
        return None
    include_standard_price = _request_wants_standard_price(request, llm_intent)

    target_file = None
    target_rows: List[Dict[str, Any]] = []
    target_vendor_columns: List[Dict[str, str]] = []
    for file in parsed_files:
        rows = [enrich_row_units(dict(row)) for row in (file.get("parsedRows") or file.get("rows") or []) if isinstance(row, dict)]
        vendor_columns = _collect_dynamic_vendor_columns(rows)
        if len(vendor_columns) >= 2:
            target_file = file
            target_rows = rows
            target_vendor_columns = vendor_columns
            break
    if not target_file or len(target_vendor_columns) < 2:
        return None

    selected_vendor_columns = _filter_vendor_columns_by_request(target_vendor_columns, request)
    explicit_vendor_request = _request_has_explicit_vendor_mention(target_vendor_columns, request)
    # 업체명을 명시하지 않은 일반 비교 요청에서만 전체 업체를 유지한다.
    # 업체명을 명시했는데 매칭이 실패한 경우에는 엉뚱한 업체 전체를 보여주지 않는다.
    if not selected_vendor_columns and not explicit_vendor_request:
        selected_vendor_columns = target_vendor_columns

    target_rows = _repair_page_split_item_names(target_rows, target_vendor_columns)

    if not selected_vendor_columns:
        return {
            "tableName": "업체별 단가 비교표",
            "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
            "columns": [
                {"key": "item_name", "label": "공종명칭"},
                {"key": "remark", "label": "비고"},
            ],
            "rows": [{
                "item_name": "요청 업체",
                "remark": "사용자가 입력한 업체명과 일치하는 원본 업체 컬럼을 찾지 못했습니다.",
            }],
            "meta": {
                "sourceMode": "single_file_multi_vendor",
                "vendorCount": 0,
                "vendors": [],
                "allVendors": [{"name": _clean_vendor_label_display(col["label"]), "columnKey": col["key"]} for col in target_vendor_columns],
                "requestedCompanyTerms": _requested_company_terms(request),
                "requestedVendorCount": _requested_vendor_count(request),
                "standardPriceShown": include_standard_price,
                "standardPriceHidden": not include_standard_price,
            },
        }

    llm_terms: List[str] = []
    if isinstance(llm_intent, dict):
        raw_terms = llm_intent.get("targetKeywords") or llm_intent.get("target_keywords") or []
        if isinstance(raw_terms, list):
            llm_terms = [str(term) for term in raw_terms if str(term or "").strip()]
    focus_terms = _extract_focus_terms(request, target_rows, [], llm_terms=llm_terms)
    request_quantity, request_quantity_unit = _extract_requested_quantity(request, llm_intent)

    all_candidate_rows: List[Dict[str, Any]] = []
    for row in target_rows:
        if not clean_cell_text(row.get("item_name") or row.get("construction_code")):
            continue
        if str(row.get("no.") or row.get("no") or "").strip().startswith("소"):
            continue
        if not any(to_number(_get_dynamic_vendor_price(row, col)) > 0 for col in target_vendor_columns):
            continue
        all_candidate_rows.append(row)

    candidate_rows = list(all_candidate_rows)
    if focus_terms:
        candidate_rows = _filter_rows_by_focus(candidate_rows, focus_terms)

    def build_output(rows_for_output: List[Dict[str, Any]], vendor_cols: List[Dict[str, Any]]) -> Tuple[List[Dict[str, str]], List[Dict[str, Any]], List[Dict[str, Any]]]:
        columns: List[Dict[str, str]] = [
            {"key": "construction_code", "label": "공종코드"},
            {"key": "item_name", "label": "공종명칭"},
            {"key": "spec", "label": "규격"},
            {"key": "quantity", "label": "요청 수량"},
            {"key": "unit", "label": "단위"},
        ]
        if include_standard_price:
            columns.append({"key": "standard_unit_price", "label": "표준시장단가"})

        vendor_key_map: Dict[str, Dict[str, str]] = {}
        vendor_meta: List[Dict[str, Any]] = []
        for idx, col in enumerate(vendor_cols, start=1):
            clean_label = _clean_vendor_label_display(col.get("label") or col.get("key") or f"업체{idx}")
            unit_price_key = _safe_compare_key(clean_label, idx, "unit_price")
            amount_key = _safe_compare_key(clean_label, idx, "amount")
            vendor_key_map[col["key"]] = {"unitPriceKey": unit_price_key, "amountKey": amount_key, "label": clean_label}
            vendor_meta.append({
                "name": clean_label,
                "index": idx - 1,
                "sourceColumnKey": col.get("key"),
                "unitPriceKey": unit_price_key,
                "amountKey": amount_key,
                "canonical": col.get("canonical"),
            })
            columns.append({"key": unit_price_key, "label": f"{clean_label} 단가"})
            columns.append({"key": amount_key, "label": f"{clean_label} 금액"})

        columns.extend([
            {"key": "lowest_vendor", "label": "최저 업체"},
            {"key": "lowest_unit_price", "label": "최저 단가"},
            {"key": "lowest_amount", "label": "최저 금액"},
        ])
        if include_standard_price:
            columns.append({"key": "lowest_vs_standard", "label": "최저-표준 차이"})
        columns.append({"key": "remark", "label": "비고"})

        output_rows: List[Dict[str, Any]] = []
        for row in rows_for_output:
            standard_price = to_number(row.get("standard_unit_price"))
            out: Dict[str, Any] = {
                "construction_code": clean_cell_text(row.get("construction_code") or ""),
                "item_name": _clean_item_display(row.get("item_name") or ""),
                "spec": clean_cell_text(row.get("spec") or ""),
                "quantity": request_quantity or clean_number(row.get("quantity") or ""),
                "request_quantity": request_quantity or clean_number(row.get("quantity") or ""),
                "unit": clean_cell_text(row.get("unit_normalized") or row.get("unit") or request_quantity_unit or ""),
                "remark": "",
            }
            if include_standard_price:
                out["standard_unit_price"] = clean_number(row.get("standard_unit_price") or "")
            vendor_prices: List[Tuple[float, str]] = []
            qty_value = to_number(out.get("quantity"))
            for col in vendor_cols:
                price_text = _get_dynamic_vendor_price(row, col)
                key_info = vendor_key_map[col["key"]]
                out[key_info["unitPriceKey"]] = price_text
                price = to_number(price_text)
                if price:
                    amount = price * qty_value if qty_value else 0
                    out[key_info["amountKey"]] = f"{amount:,.0f}" if amount else ""
                    vendor_prices.append((price, key_info["label"]))
                else:
                    out[key_info["amountKey"]] = ""
            if vendor_prices:
                lowest_price, lowest_vendor = min(vendor_prices, key=lambda item: item[0])
                out["lowest_vendor"] = lowest_vendor
                out["lowest_unit_price"] = f"{lowest_price:,.0f}"
                out["lowest_amount"] = f"{lowest_price * qty_value:,.0f}" if qty_value else ""
                if include_standard_price:
                    out["lowest_vs_standard"] = _format_price_diff(lowest_price, standard_price)
            output_rows.append(out)
        return columns, output_rows, vendor_meta

    if not candidate_rows:
        return {
            "tableName": "업체별 단가 비교표",
            "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
            "columns": [
                {"key": "item_name", "label": "공종명칭"},
                {"key": "remark", "label": "비고"},
            ],
            "rows": [{
                "item_name": ", ".join(focus_terms) if focus_terms else "요청 공종",
                "remark": "요청한 공종/업체 조건과 일치하는 단가 행을 찾지 못했습니다.",
            }],
            "meta": {
                "sourceMode": "single_file_multi_vendor",
                "vendorCount": len(selected_vendor_columns),
                "vendors": [{"name": _clean_vendor_label_display(col["label"]), "columnKey": col["key"]} for col in selected_vendor_columns],
                "allVendors": [{"name": _clean_vendor_label_display(col["label"]), "columnKey": col["key"]} for col in target_vendor_columns],
                "focusTerms": focus_terms,
            },
        }

    columns, result_rows, selected_vendor_meta = build_output(candidate_rows, selected_vendor_columns)
    all_columns, all_rows, all_vendor_meta = build_output(all_candidate_rows, target_vendor_columns)

    return {
        "tableName": "업체별 단가 비교표",
        "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
        "columns": columns,
        "rows": result_rows,
        "meta": {
            "sourceMode": "single_file_multi_vendor",
            "sourceFileName": str(target_file.get("originalName") or target_file.get("original_name") or ""),
            "vendorCount": len(selected_vendor_columns),
            "vendors": selected_vendor_meta,
            "allVendors": all_vendor_meta,
            "allColumns": all_columns,
            "allRows": all_rows,
            "focusTerms": focus_terms,
            "requestedCompanyTerms": _requested_company_terms(request),
            "requestedVendorCount": _requested_vendor_count(request),
            "standardPriceShown": include_standard_price,
            "standardPriceHidden": not include_standard_price,
        },
    }


def build_multi_vendor_price_comparison(
    parsed_files: List[Dict[str, Any]],
    user_request: str = "",
    llm_intent: Dict[str, Any] | None = None,
    include_all_rows: bool = True,
) -> Dict[str, Any] | None:
    if len(parsed_files or []) < 2:
        return None
    request = str(user_request or "")
    include_standard_price = _request_wants_standard_price(request, llm_intent)
    wants_compare = any(word in request for word in ["비교", "단가", "최저", "견적", "업체", "회사"])
    if not wants_compare:
        return None

    reference_rows: List[Dict[str, Any]] = []
    vendor_sources: List[Dict[str, Any]] = []
    for idx, file in enumerate(parsed_files, start=1):
        filename = str(file.get("originalName") or file.get("original_name") or "")
        text = str(file.get("extractedText") or file.get("extracted_text") or "")
        rows = [enrich_row_units(dict(row)) for row in (file.get("parsedRows") or file.get("rows") or []) if isinstance(row, dict)]
        if _is_estimate_file(filename, text, rows, user_request=request, llm_intent=llm_intent):
            company = _extract_company_name(filename, text, rows, idx)
            vendor_rows: List[Dict[str, Any]] = []
            for row in rows:
                if not str(row.get("item_name") or "").strip():
                    continue
                if not _get_vendor_price_text(row):
                    continue
                next_row = dict(row)
                next_row["vendor_name"] = company
                next_row["source_file_name"] = filename
                vendor_rows.append(enrich_row_units(_normalize_price_role_fallback(next_row)))
            if vendor_rows:
                vendor_sources.append({"name": company, "filename": filename, "rows": vendor_rows})
            continue
        if _is_standard_market_file(filename, text, rows, user_request=request, llm_intent=llm_intent):
            text_rows = extract_standard_market_rows_from_text(text)
            if rows:
                rows = [row for row in rows if str(row.get("item_name") or "").strip() and str(row.get("unit_price") or row.get("standard_unit_price") or "").strip()]
                reference_rows.extend(merge_standard_market_rows(rows, text_rows))
            else:
                reference_rows.extend(text_rows)
            continue

    vendor_sources = _dedupe_vendor_sources(vendor_sources)

    if len(vendor_sources) < 2:
        return None

    all_vendor_rows = [row for source in vendor_sources for row in source["rows"]]
    llm_terms = []
    if isinstance(llm_intent, dict):
        raw_terms = llm_intent.get("targetKeywords") or llm_intent.get("target_keywords") or []
        if isinstance(raw_terms, list):
            llm_terms = [str(term) for term in raw_terms if str(term or "").strip()]
    focus_terms = _extract_focus_terms(request, all_vendor_rows, reference_rows, llm_terms=llm_terms)
    request_quantity, request_quantity_unit = _extract_requested_quantity(request, llm_intent)
    # 후속 질문에서 다른 공종으로 다시 바꿀 수 있도록, 최초 요청으로 필터링되기 전의 전체 비교표를 보존한다.
    # 화면에는 focus_terms에 맞춘 행만 보여주되, table_json.meta.allRows에는 전체 행을 저장한다.
    all_columns_for_meta: List[Dict[str, Any]] = []
    all_rows_for_meta: List[Dict[str, Any]] = []
    if include_all_rows and focus_terms:
        full_table = build_multi_vendor_price_comparison(
            parsed_files,
            user_request="업체별 단가 비교표",
            llm_intent=None,
            include_all_rows=False,
        )
        if full_table:
            all_columns_for_meta = list(full_table.get("columns") or [])
            all_rows_for_meta = list(full_table.get("rows") or [])
    if focus_terms:
        for source in vendor_sources:
            source["rows"] = _filter_rows_by_focus(source["rows"], focus_terms)
        vendor_sources = [source for source in vendor_sources if source.get("rows")]
        reference_rows = _filter_rows_by_focus(reference_rows, focus_terms)
        if len(vendor_sources) < 2:
            return {
                "tableName": "업체별 단가 비교표",
                "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
                "columns": [
                    {"key": "item_name", "label": "공종명칭"},
                    {"key": "remark", "label": "비고"},
                ],
                "rows": [{
                    "item_name": ", ".join(focus_terms),
                    "remark": "요청한 공종이 2개 이상 업체 견적서에서 동시에 확인되지 않아 비교표를 만들 수 없습니다.",
                }],
                "meta": {
                    "vendorCount": len(vendor_sources),
                    "vendors": [{"name": item["name"], "filename": item["filename"]} for item in vendor_sources],
                    "referenceRows": len(reference_rows),
                    "focusTerms": focus_terms,
                    "standardPriceShown": include_standard_price,
                    "standardPriceHidden": not include_standard_price,
                    "allColumns": all_columns_for_meta,
                    "allRows": all_rows_for_meta,
                },
            }

    groups: Dict[str, Dict[str, Any]] = {}
    for source in vendor_sources:
        for row in source["rows"]:
            key = _row_match_key(row)
            if key in {"text:||", "text:|"}:
                continue
            group = groups.setdefault(key, {"sample": row, "vendors": {}, "sourceFiles": set()})
            group["vendors"][source["name"]] = row
            group["sourceFiles"].add(source["filename"])
    if not groups:
        return None

    columns: List[Dict[str, str]] = [
        {"key": "construction_code", "label": "공종코드"},
        {"key": "item_name", "label": "공종명칭"},
        {"key": "spec", "label": "규격"},
        {"key": "quantity", "label": "요청 수량"},
        {"key": "unit", "label": "단위"},
    ]
    if include_standard_price:
        columns.append({"key": "standard_unit_price", "label": "표준시장단가"})
    vendor_key_map: Dict[str, str] = {}
    for idx, source in enumerate(vendor_sources, start=1):
        key = _safe_compare_key(source["name"], idx, "unit_price")
        vendor_key_map[source["name"]] = key
        columns.append({"key": key, "label": f"{source['name']} 단가"})
    columns.extend([
        {"key": "lowest_vendor", "label": "최저 업체"},
        {"key": "lowest_unit_price", "label": "최저 단가"},
    ])
    if include_standard_price:
        columns.append({"key": "lowest_vs_standard", "label": "최저-표준 차이"})
    columns.extend([
        {"key": "labor_ratio", "label": "노무비율"},
        {"key": "remark", "label": "비고"},
    ])

    result_rows: List[Dict[str, Any]] = []
    for _, group in sorted(groups.items(), key=lambda kv: (str(kv[1]["sample"].get("item_name") or ""), str(kv[1]["sample"].get("spec") or ""))):
        sample = group["sample"]
        ref = _find_reference_row(sample, reference_rows)
        standard_price = to_number(sample.get("standard_unit_price") or "") or (to_number(ref.get("unit_price")) if ref else 0.0)
        unit = clean_cell_text(sample.get("unit_normalized") or sample.get("unit") or (ref or {}).get("unit") or "")
        out: Dict[str, Any] = {
            "construction_code": clean_cell_text(sample.get("construction_code") or (ref or {}).get("construction_code") or ""),
            "item_name": clean_cell_text(sample.get("item_name") or (ref or {}).get("item_name") or ""),
            "spec": clean_cell_text(sample.get("spec") or (ref or {}).get("spec") or ""),
            "quantity": request_quantity or clean_number(sample.get("quantity") or ""),
            "request_quantity": request_quantity or clean_number(sample.get("quantity") or ""),
            "unit": unit,
            "labor_ratio": clean_cell_text((ref or {}).get("labor_ratio") or sample.get("labor_ratio") or ""),
            "remark": "",
        }
        if include_standard_price:
            out["standard_unit_price"] = clean_number(sample.get("standard_unit_price") or (ref or {}).get("unit_price") or "")
        vendor_prices: List[Tuple[float, str]] = []
        units = {unit} if unit else set()
        for source in vendor_sources:
            vendor = source["name"]
            vrow = group["vendors"].get(vendor)
            col_key = vendor_key_map[vendor]
            if not vrow:
                out[col_key] = ""
                continue
            price_text = _get_vendor_price_text(vrow)
            out[col_key] = price_text
            price = to_number(price_text)
            if price:
                vendor_prices.append((price, vendor))
            vunit = clean_cell_text(vrow.get("unit_normalized") or vrow.get("unit") or "")
            if vunit:
                units.add(vunit)
        if vendor_prices:
            lowest_price, lowest_vendor = min(vendor_prices, key=lambda item: item[0])
            out["lowest_vendor"] = lowest_vendor
            out["lowest_unit_price"] = f"{lowest_price:,.0f}"
            if include_standard_price:
                out["lowest_vs_standard"] = _format_price_diff(lowest_price, standard_price)
        else:
            out["lowest_vendor"] = ""
            out["lowest_unit_price"] = ""
            if include_standard_price:
                out["lowest_vs_standard"] = ""
        remarks = []
        if request_quantity and request_quantity_unit and unit and compact_text(request_quantity_unit) not in compact_text(unit) and compact_text(unit) not in compact_text(request_quantity_unit):
            remarks.append(f"요청 수량 단위 확인 필요: {request_quantity}{request_quantity_unit} / 원문 단위 {unit}")
        if len([u for u in units if u]) >= 2:
            remarks.append(f"단위 확인 필요: {', '.join(sorted(units))}")
        if reference_rows and not ref:
            remarks.append("기준자료 행 미매칭")
        out["remark"] = " / ".join(remarks)
        result_rows.append(out)
        if len(result_rows) >= _env_int("MULTI_COMPARE_MAX_ROWS", 300):
            break

    if not result_rows:
        return None
    columns = prune_empty_columns(columns, result_rows)
    if not all_rows_for_meta:
        all_rows_for_meta = result_rows
    if not all_columns_for_meta:
        all_columns_for_meta = columns
    meta_vendors = []
    for vendor_index, item in enumerate(vendor_sources, start=1):
        name = item["name"]
        meta_vendors.append({
            "name": name,
            "filename": item.get("filename"),
            "index": vendor_index - 1,
            "unitPriceKey": vendor_key_map.get(name),
            "priceKey": vendor_key_map.get(name),
        })

    return {
        "tableName": "업체별 단가 비교표",
        "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
        "columns": columns,
        "rows": result_rows,
        "meta": {
            "vendorCount": len(vendor_sources),
            "vendors": meta_vendors,
            "referenceRows": len(reference_rows),
            "focusTerms": focus_terms,
            "standardPriceShown": include_standard_price,
            "standardPriceHidden": not include_standard_price,
            "allColumns": all_columns_for_meta,
            "allRows": all_rows_for_meta,
        },
    }


REFERENCE_TABLE_KEYWORDS = [
    "단가", "가격", "계약단가", "거래실례가격", "노임단가", "정부노임단가",
    "요율", "산정", "산출", "계산", "적용기준", "산정기준", "기준",
    "할증", "품셈", "경비", "노무비", "재료비", "자재", "전력비", "기본요금",
]

FIELD_ALIASES = {
    "construction_code": ["공종코드", "코드", "code"],
    "labor_ratio": ["노무비율", "노무 비율", "노무율", "labor"],
    "vendor_name": ["업체", "업체명", "거래처", "공급업체", "회사명", "상호", "시공사", "견적업체"],
    "item_name": ["품목", "품목명", "자재", "자재명", "내역", "공종", "작업명", "명칭", "품명", "공종명칭", "공종명"],
    "spec": ["규격", "사양", "모델", "크기", "치수", "규격명"],
    "quantity": ["수량", "물량", "개수", "qty", "수량산출", "소요량"],
    "unit": ["단위", "uom", "unit"],
    # 가격 역할 분리: 특정 업체명/A회사/B회사를 코드에 넣지 않고, 헤더의 의미만 분리한다.
    "standard_unit_price": ["표준단가", "기준단가", "시장단가", "표준시장단가", "기준가격"],
    "vendor_unit_price": ["견적단가", "견적가", "견적가격", "견적금액", "제안단가", "제안가", "제시단가", "제시가", "업체단가", "업체견적", "공급단가", "견적"],
    "unit_price": ["단가", "단위금액", "원가", "가격"],
    "price_diff": ["차이", "차액", "증감", "증감액", "차이금액"],
    "diff_rate": ["대비", "대비율", "비율", "증감률", "율", "%"],
    "amount": ["금액", "합계", "합계금액", "총금액", "공급가액", "금 액"],
    "remark": ["비고", "메모", "특이사항", "참고사항", "비 고"],
}

NUMBER_KEYS = {"quantity", "unit_price", "vendor_unit_price", "standard_unit_price", "amount", "supply_amount", "tax_amount", "price_diff"}
PRICE_COMPARE_WORDS = ["단가", "견적", "비교", "업체", "최저", "최고", "가격", "견적서"]

PRICE_VALUE_PRIORITY = ("vendor_unit_price", "unit_price", "amount")


logger = logging.getLogger("app.document_analyzer")


def _env_int(name: str, default: int) -> int:
    try:
        return int(str(os.getenv(name, default)).strip())
    except (TypeError, ValueError):
        return default


def _env_float(name: str, default: float) -> float:
    try:
        return float(str(os.getenv(name, default)).strip())
    except (TypeError, ValueError):
        return default


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "y", "on"}


def _new_parse_logger(scope: str) -> Tuple[List[Dict[str, Any]], Callable[..., None]]:
    # 기본은 터미널 로그만 남긴다.
    # 프론트/API 응답으로 로그를 보내고 싶을 때만 RETURN_PARSE_LOGS=true로 켠다.
    logs: List[Dict[str, Any]] = []
    capture_logs = _env_bool("RETURN_PARSE_LOGS", False)

    def write(level: str, message: str, **data: Any) -> None:
        clean_data = {k: v for k, v in data.items() if v is not None}
        suffix = " ".join(f"{key}={value}" for key, value in clean_data.items())
        text = f"[{scope}] {message}" + (f" {suffix}" if suffix else "")
        if capture_logs:
            logs.append({
                "time": datetime.now().isoformat(timespec="seconds"),
                "level": level.upper(),
                "message": text,
                "data": clean_data,
            })
        log_fn = getattr(logger, level.lower(), logger.info)
        log_fn(text)

    return logs, write


def _limit_pages(total_pages: int, env_name: str = "PDF_MAX_PAGES") -> int:
    # 0 또는 음수는 전체 페이지 처리. 기존 50페이지 제한을 제거하기 위한 기본값이다.
    limit = _env_int(env_name, 0)
    if limit <= 0:
        return total_pages
    return min(total_pages, limit)



DOCUMENT_TYPE_RULES = [
    ("표준시장단가표", ["건설공사표준시장단가", "표준시장단가", "공종코드", "노무비율"]),
    ("기준서/지침서", ["지침서", "적용기준", "표준품셈", "품셈", "적산", "공사원가", "건축견적지침서"]),
    ("Use Case 명세서", ["use case", "유스케이스", "use case id", "actor 정의", "main flow", "alternative flow"]),
    ("업무 프로세스 명세서", ["프로세스", "as-is", "to-be", "업무 흐름"]),
    ("요구사항 정의서", ["요구사항", "기능 요구", "비기능 요구", "요구사항 id"]),
    ("보고서", ["보고서", "kpi", "pain point", "기대 효과"]),
]

BUSINESS_TABLE_REQUIRED_KEYS = {"item_name", "quantity", "unit", "unit_price", "vendor_unit_price", "standard_unit_price", "amount", "vendor_name"}


def compact_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _company_compare_key(value: Any) -> str:
    return re.sub(r"[\s._\-()（）\[\]{}·,㈜]+|주식회사|\(주\)|（주）", "", str(value or "")).lower()


def _has_unbalanced_parenthesis(value: Any) -> bool:
    text = str(value or "")
    return text.count("(") != text.count(")") or text.count("[") != text.count("]") or text.endswith(("(", "[", "{"))


GENERIC_FOCUS_TERMS = {
    "단가", "견적", "비교", "표준", "시장", "표준시장", "파일", "회사", "업체",
    "공종", "표", "분석", "엑셀", "문서", "자료", "자사양식", "비교표",
    "업체별", "회사별", "공종별", "견적서", "비교견적서", "단가비교",
    "업체비교", "회사비교", "업체별단가", "업체별비교", "업체별단가비교",
    "공종명칭", "공종코드", "규격", "수량", "단위", "금액", "최저", "작성자",
}

def _valid_focus_term(value: Any) -> bool:
    term = compact_text(value)
    if len(term) < 2:
        return False
    if term in GENERIC_FOCUS_TERMS:
        return False
    if any(term == generic or term.endswith(generic) and generic in {"비교표", "견적서"} for generic in GENERIC_FOCUS_TERMS):
        return False
    if _has_unbalanced_parenthesis(term):
        return False
    if re.fullmatch(r"[0-9, .]+(?:개|건|장|식|개사|업체|회사)?", str(value or "").strip()):
        return False
    return True


def source_has_value(source_text: str, value: Any) -> bool:
    """LLM이 만든 값이 실제 원문에 존재하는지 보수적으로 확인한다."""
    text = compact_text(source_text)
    raw = str(value or "").strip()
    if not raw:
        return True
    compact = compact_text(raw)
    if not compact:
        return True
    # 너무 일반적인 단어/단위는 근거 검증 대상으로 보지 않는다.
    if compact in {"개", "개소", "본", "m", "m2", "m3", "㎡", "㎥", "공m3", "공㎥", "ea", "pcs", "box", "set", "lot", "식", "원", "-", "none", "null"}:
        return True
    # 숫자는 구분자를 제거하고 확인한다.
    if re.fullmatch(r"[0-9,\.\-]+", compact):
        digits = re.sub(r"[^0-9]", "", compact)
        return bool(digits) and digits in re.sub(r"[^0-9]", "", text)
    return compact in text


def infer_document_profile(text: str, user_request: str = "") -> Dict[str, Any]:
    """문서 유형을 키워드 하나로 과잉 분류하지 않고, 제목/구조 기반으로 보수적으로 판별한다."""
    lower = text.lower()
    compact = compact_text(text)

    if is_text_only_vendor_comparison_report(text):
        return {
            "documentType": "업체별 단가 비교 검토보고서",
            "purpose": "서술형 업체별 단가 비교 결과와 확인 필요 사항 검토",
            "confidence": 0.9,
        }

    if is_standard_market_price_document(text):
        return {
            "documentType": "표준시장단가표",
            "purpose": "공종별 표준시장단가 표 추출 및 단가 확인",
            "confidence": 0.9,
        }

    for doc_type, keywords in DOCUMENT_TYPE_RULES:
        score = sum(1 for kw in keywords if compact_text(kw) in compact or kw in lower)
        if score >= 2:
            return {
                "documentType": doc_type,
                "purpose": "문서 내용 분석 및 엑셀화 가능 여부 검토",
                "confidence": 0.86,
            }

    # 견적서/단가표는 단순히 본문에 '견적서'라는 예시 단어가 있다는 이유만으로 분류하지 않는다.
    price_structural_score = 0
    for kw in ["품명", "품목", "규격", "수량", "단위", "단가", "금액", "공급가액", "합계"]:
        if kw in text:
            price_structural_score += 1
    request_price = any(word in (user_request or "") for word in ["단가", "견적", "비교", "가격"])
    if price_structural_score >= 5 and ("견적" in text or "단가" in text or request_price) and not is_reference_or_guideline_document(text):
        return {
            "documentType": "견적서/단가표",
            "purpose": "단가 및 금액 비교용 표 데이터 생성",
            "confidence": 0.82,
        }

    return {
        "documentType": "업무 문서",
        "purpose": "문서 내용 요약 및 표 데이터 추출 가능 여부 확인",
        "confidence": 0.68,
    }


def build_file_profiles(
    parsed_files: List[Dict[str, Any]],
    user_request: str = "",
    llm_intent: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    """첨부파일별 유형/역할 요약을 만든다.

    - 특정 파일명, 특정 업체명, A회사/B회사 같은 값을 코드에 박지 않는다.
    - 사용자 요청, 파일 수, 표 행 구조, 문서 내 업체명/가격행 근거로 역할을 정한다.
    - LLM이 timeout이어도 PDF 파싱 결과만으로 최소 파일별 분석을 제공한다.
    """
    profiles: List[Dict[str, Any]] = []
    wants_compare = _request_wants_company_comparison(user_request, llm_intent)

    for index, file in enumerate(parsed_files or [], start=1):
        filename = str(file.get("originalName") or file.get("original_name") or file.get("filename") or f"파일{index}")
        text = str(file.get("extractedText") or file.get("extracted_text") or "")
        rows = [row for row in (file.get("parsedRows") or file.get("rows") or []) if isinstance(row, dict)]
        profile = infer_document_profile(text, user_request)
        company_name = _extract_company_name(filename, text, rows, index)
        page_count = int(file.get("pageCount") or file.get("page_count") or 0)
        char_count = len(text or "")
        row_count = len(rows or [])
        price_row_count = sum(
            1
            for row in rows
            if isinstance(row, dict)
            and clean_cell_text(row.get("item_name") or row.get("construction_code") or "")
            and _row_has_price_value(row)
        )

        document_type = str(profile.get("documentType") or "업무 문서")
        role = "SOURCE_DOCUMENT"
        role_label = "분석 대상 문서"
        confidence = float(profile.get("confidence") or 0.68)

        if wants_compare and _is_estimate_file(filename, text, rows, user_request=user_request, llm_intent=llm_intent):
            role = "COMPARE_TARGET"
            role_label = "비교 대상 견적서"
            document_type = "업체 견적서/단가표"
            confidence = max(confidence, 0.86)
        elif _is_standard_market_file(filename, text, rows, user_request=user_request, llm_intent=llm_intent):
            role = "REFERENCE_PRICE"
            role_label = "기준 단가 자료"
            document_type = "표준시장단가/기준단가 자료"
            confidence = max(confidence, 0.84)
        elif is_reference_or_guideline_document(text):
            role = "REFERENCE_GUIDELINE"
            role_label = "기준/지침 자료"
            document_type = "기준서/지침서"
            confidence = max(confidence, 0.8)

        summary_parts = [
            f"{page_count}페이지" if page_count else "페이지 미확인",
            f"표 후보 {row_count}행" if row_count else "표 후보 없음",
        ]
        if price_row_count:
            summary_parts.append(f"가격 행 {price_row_count}행")
        if company_name:
            summary_parts.append(f"업체명 {company_name}")

        profiles.append({
            "index": index,
            "fileName": filename,
            "companyName": company_name,
            "documentType": document_type,
            "role": role,
            "roleLabel": role_label,
            "pageCount": page_count,
            "charCount": char_count,
            "rowCount": row_count,
            "priceRowCount": price_row_count,
            "confidence": round(confidence, 4),
            "summary": " / ".join(summary_parts),
        })

    return profiles


def extract_key_values_from_text(text: str) -> List[Dict[str, Any]]:
    patterns = [
        ("Use Case ID", r"Use\s*Case\s*ID\s*([^\n]+)"),
        ("Use Case 명", r"Use\s*Case\s*명\s*([^\n]+)"),
        ("업무 영역", r"업무\s*영역\s*([^\n]+)"),
        ("작성일", r"작성일\s*([0-9]{4}[-./][0-9]{2}[-./][0-9]{2}|[0-9]{4}[-./][0-9]{2}[-./][0-9]{1,2}|[^\n]{4,30})"),
        ("작성자", r"작성자\s*([^\n]+)"),
        ("우선순위", r"우선순위\s*([^\n]+)"),
        ("자동화 수준", r"자동화\s*수준\s*([^\n]+)"),
    ]
    result: List[Dict[str, Any]] = []
    seen = set()
    for label, pattern in patterns:
        m = re.search(pattern, text, re.I)
        if not m:
            continue
        value = re.sub(r"\s{2,}", " ", m.group(1).strip())[:160]
        # 다음 항목까지 한 줄로 붙은 PDF 추출물을 완화한다.
        value = re.split(r"\s+(?:Use\s*Case\s*명|업무\s*영역|작성일|작성자|관련\s*시스템|우선순위|자동화\s*수준|관련\s*KPI)", value)[0].strip()
        if value and (label, value) not in seen:
            result.append({"label": label, "value": value})
            seen.add((label, value))
    return result[:12]


def is_business_row_supported(row: Dict[str, Any], source_text: str) -> bool:
    """LLM/규칙 파서가 만든 행이 원문 근거를 갖는 실제 업무 표 행인지 확인한다."""
    if not row:
        return False
    meaningful = {k: str(row.get(k) or "").strip() for k in BUSINESS_TABLE_REQUIRED_KEYS if str(row.get(k) or "").strip()}
    if not meaningful:
        return False

    # 품목명/업체명/규격/금액 중 하나 이상은 원문에 있어야 한다.
    anchor_keys = ["item_name", "vendor_name", "spec", "amount", "unit_price", "vendor_unit_price", "standard_unit_price", "price_diff"]
    anchor_values = [row.get(k) for k in anchor_keys if str(row.get(k) or "").strip()]
    if not anchor_values:
        return False
    return any(source_has_value(source_text, v) for v in anchor_values)


def filter_grounded_rows(rows: List[Dict[str, Any]], source_text: str) -> List[Dict[str, Any]]:
    filtered: List[Dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        if is_reference_row_supported(row, source_text) or is_business_row_supported(row, source_text):
            filtered.append(row)
    return filtered


def is_narrative_document(text: str) -> bool:
    compact = compact_text(text)
    narrative_markers = ["유스케이스", "usecase", "프로젝트개요", "actor정의", "mainflow", "alternativeflow", "businessrule", "kpi정의"]
    return sum(1 for marker in narrative_markers if marker in compact) >= 2


def is_reference_or_guideline_document(text: str) -> bool:
    compact = compact_text(text)
    markers = ["건축견적지침서", "지침서", "적용기준", "표준품셈", "품셈의정의", "적산및견적", "공사원가", "관계법령"]
    return sum(1 for marker in markers if compact_text(marker) in compact) >= 2


def is_standard_market_price_document(text: str) -> bool:
    compact = compact_text(text)
    markers = ["건설공사표준시장단가", "표준시장단가", "공종코드", "공종명칭", "노무비율"]
    return ("표준시장단가" in compact and "공종코드" in compact and "노무비율" in compact) or sum(1 for marker in markers if compact_text(marker) in compact) >= 3


def is_text_only_vendor_comparison_report(text: str) -> bool:
    compact = compact_text(text)
    has_compare_title = "업체별단가비교검토보고서" in compact or ("업체별" in compact and "단가비교" in compact and "검토보고서" in compact)
    has_text_only_policy = any(marker in compact for marker in ["텍스트전용", "문장형텍스트", "표형식자료를사용하지", "행열색상강조요약표등은모두제외", "문단형태로만기술"])
    has_vendor_totals = "총견적금액" in compact and "표준시장단가" in compact and "최저가" in compact
    return bool(has_compare_title and (has_text_only_policy or has_vendor_totals))


def extract_text_vendor_total_rows(text: str) -> List[Dict[str, Any]]:
    """서술형 업체별 비교보고서의 총괄 비교 문장에서 업체 총액 행만 보수적으로 추출한다.

    표가 없는 문서에서 임의 품목 행을 만들지 않기 위한 요약용 행이다.
    개별 공종/업체 단가가 원문에 모두 있지 않으면 확정 단가 비교표로 사용하지 않는다.
    """
    rows: List[Dict[str, Any]] = []
    seen = set()
    section_match = re.search(r"3\.\s*총괄\s*비교\s*결과(?P<section>.*?)(?:\n\s*4\.|4\.\s*세부)", text, re.S)
    source_text = section_match.group("section") if section_match else text[:12000]
    # 예: ㈜한국전기가 12,198,571원으로 ... 표준시장단가 대비 -966,869원, -7.3%
    patterns = [
        re.compile(
            r"(?P<vendor>(?:[A-Z]회사\s*)?(?:㈜|\(주\)|주식회사)?[가-힣A-Za-z0-9·ㆍ]+(?:전기|설비|시스템|일렉|기술)?(?:㈜)?)"
            r"(?:의)?\s*(?:총\s*)?견적금액은\s*(?P<amount>[0-9]{1,3}(?:,[0-9]{3})+)원"
            r"(?P<context>.{0,180}?)(?:비교하면|대비)\s*(?P<diff>[+\-]?[0-9]{1,3}(?:,[0-9]{3})+)원\s*,\s*(?P<rate>[+\-]?[0-9]+(?:\.\d+)?)%",
            re.S,
        ),
        re.compile(
            r"(?P<vendor>(?:[A-Z]회사\s*)?(?:㈜|\(주\)|주식회사)?[가-힣A-Za-z0-9·ㆍ]+(?:전기|설비|시스템|일렉|기술)?(?:㈜)?)"
            r"(?:가|이|은|는)?\s*(?P<amount>[0-9]{1,3}(?:,[0-9]{3})+)원으로(?P<context>.{0,180}?)(?:대비|비교하여|비교하면)\s*(?P<diff>[+\-]?[0-9]{1,3}(?:,[0-9]{3})+)원\s*,\s*(?P<rate>[+\-]?[0-9]+(?:\.\d+)?)%",
            re.S,
        ),
    ]
    for pattern in patterns:
        for m in pattern.finditer(source_text):
            vendor = clean_cell_text(m.group("vendor"))
            vendor = re.sub(r"^[A-Z]회사\s*", "", vendor).strip()
            vendor = re.sub(r"(의|가|이|은|는)$", "", vendor).strip()
            amount = clean_number(m.group("amount"))
            diff = clean_number(m.group("diff"))
            rate = f"{m.group('rate')}%"
            key = compact_text(f"{vendor}{amount}{diff}{rate}")
            if not vendor or key in seen:
                continue
            seen.add(key)
            context = _clean_line(m.group("context"), limit=120) if '_clean_line' in globals() else clean_cell_text(m.group("context"))[:120]
            remark = "가격 측면 우선 검토 대상" if "가장 낮" in context or "우선 검토" in context else "공종별 편차 및 조건 재확인 필요"
            rows.append({
                "vendor_name": vendor,
                "total_amount": amount,
                "price_diff": diff,
                "diff_rate": rate,
                "remark": remark,
            })
    return rows[:20]


def _extract_requested_quantity_value(user_request: str) -> Tuple[str, str]:
    """사용자 요청의 작업 수량을 추출한다. 예: '각 수량 50개', '량은 50개' -> ('50', '개').

    '2개 업체' 같은 업체 수는 수량으로 오인하지 않는다.
    """
    text = str(user_request or "")
    patterns = [
        r"(?:각\s*)?(?:수량|량)\s*(?:은|는|:)?\s*(?P<qty>[0-9]+(?:\.[0-9]+)?)\s*(?P<unit>개|대|식|m|M|EA|ea|본|장|조)?",
        r"(?P<qty>[0-9]+(?:\.[0-9]+)?)\s*(?P<unit>개|대|식|m|M|EA|ea|본|장|조)\s*(?:씩|로)?(?!\s*(?:업체|회사|거래처|개사))",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return clean_number(m.group('qty')), (m.groupdict().get('unit') or '').strip()
    return "", ""


def _extract_requested_item_terms(user_request: str) -> List[str]:
    """사용자 요청에서 공종/품목명으로 보이는 토큰을 추출한다."""
    text = str(user_request or "")
    if not text.strip():
        return []
    terms: List[str] = []
    # '유입변압기설치', '기중차단기 설치', '변압기 설치'처럼 설치/포설/배선/조립이 붙은 공종명을 우선 추출한다.
    # 공백이 있어도 compact_text 후 비교하므로 "변압기 설치" → "변압기설치"로 통일
    for m in re.finditer(r"[가-힣A-Za-z0-9·ㆍ()]+(?:\s+[가-힣A-Za-z0-9·ㆍ()]+)*?\s*(?:설치|포설|배선|조립|제작|철거|교체|시공)", text):
        term = compact_text(m.group(0))
        if len(term) >= 3:
            terms.append(term)
    cleaned = re.sub(r"(비교|단가|금액|수량|각|업체|회사|견적|표|정리|만들어|해줘|해줄래|그리고|랑|와|과|별로|기준|개|대|식|로|으로)", " ", text)
    cleaned = re.sub(r"[0-9]+(?:\.[0-9]+)?", " ", cleaned)
    for piece in re.split(r"[\s,，/]+", cleaned):
        token = compact_text(piece)
        if len(token) >= 4 and not any(stop in token for stop in ["대한전기", "파워시스템", "한국전기", "스마트일렉", "전기기술"]):
            terms.append(token)
    return list(dict.fromkeys(terms))


def _extract_text_report_vendor_master(text: str) -> List[Dict[str, Any]]:
    """서술형 비교보고서의 업체 기본사항에서 A/B/C/D/E 업체명을 추출한다."""
    vendors: List[Dict[str, Any]] = []
    seen = set()
    section_match = re.search(r"2\.\s*업체\s*기본사항(?P<section>.*?)(?:\n\s*3\.|3\.\s*총괄)", text, re.S)
    source = section_match.group('section') if section_match else text[:8000]
    pattern = re.compile(r"(?P<alias>[A-Z])회사\s+(?P<name>(?:㈜|\(주\)|주식회사)?[가-힣A-Za-z0-9·ㆍ]+(?:㈜)?)")
    for m in pattern.finditer(source):
        alias = f"{m.group('alias').upper()}회사"
        name = clean_cell_text(m.group('name'))
        name = re.sub(r"(의|가|이|은|는)$", "", name).strip()
        key = compact_text(name)
        if not name or key in seen:
            continue
        seen.add(key)
        vendors.append({"alias": alias, "name": name, "compareKey": key})
    return vendors


def _select_text_report_vendors(text: str, user_request: str) -> List[Dict[str, Any]]:
    all_vendors = _extract_text_report_vendor_master(text)
    if not all_vendors:
        return []
    request_terms = _request_vendor_terms_from_message(user_request)
    compact_request = compact_text(user_request)
    selected: List[Dict[str, Any]] = []

    # 1) 실제 업체명/별칭을 원문 요청에 직접 대조한다. ㈜, (주), 주식회사 유무 차이를 모두 허용한다.
    for vendor in all_vendors:
        vendor_terms = _text_report_vendor_match_terms(vendor)
        if compact_request and any(term and term in compact_request for term in vendor_terms):
            selected.append(vendor)
            continue
        for req in request_terms:
            if any(req and term and (req in term or term in req) for term in vendor_terms):
                selected.append(vendor)
                break

    if selected:
        deduped: List[Dict[str, Any]] = []
        seen = set()
        for vendor in selected:
            ident = compact_text(vendor.get('name') or vendor.get('alias'))
            if ident in seen:
                continue
            seen.add(ident)
            deduped.append(vendor)
        return deduped

    # 2) 업체명이 전혀 없고 '2개 업체'처럼 수량만 있는 경우에만 앞에서부터 N개를 사용한다.
    requested_count = _requested_vendor_count(user_request)
    has_vendor_like_text = bool(request_terms)
    if requested_count and not has_vendor_like_text:
        return all_vendors[:requested_count]
    return all_vendors


def _split_item_name_and_spec(item_text: str) -> Tuple[str, str]:
    item = clean_cell_text(item_text)
    # 마지막 규격 토큰을 분리한다. 예: '유입변압기설치 300kVA' -> ('유입변압기설치', '300kVA')
    m = re.match(r"(?P<name>.*?)(?:\s+(?P<spec>[0-9][0-9A-Za-z.,/×xX㎡㎥㎜㎝㎏kKmMvVaA\- ]+))$", item)
    if m and m.group('name').strip():
        return clean_cell_text(m.group('name')), clean_cell_text(m.group('spec'))
    return item, ""


def extract_text_vendor_item_rows(text: str, user_request: str = "") -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """표가 없는 비교보고서에서 공종별 최저/최고 견적 문장을 행으로 추출한다.

    중요: 원문에 특정 업체의 단가가 직접 적힌 경우에만 해당 업체 단가를 채운다.
    요청 업체가 최저/최고 업체로 등장하지 않는 행은 '확인 필요'로 남긴다.
    """
    selected_vendors = _select_text_report_vendors(text, user_request)
    requested_terms = _extract_requested_item_terms(user_request)
    request_qty, request_unit = _extract_requested_quantity_value(user_request)
    rows: List[Dict[str, Any]] = []
    seen = set()
    # 페이지 헤더 패턴: "[page X / Y]" 포함 형식도 처리
    page_header = r"(?:\s*전기공사\s*업체별\s*단가\s*비교\s*검토보고서\s*-\s*텍스트\s*전용\s*샘플\s*\d+\s*(?:\[page\s*\d+\s*/\s*\d+\])?\s*)?"
    vendor_pat = r"(?:[A-Z]회사\s*)?(?:㈜|\(주\)|주식회사)?[가-힣A-Za-z0-9·ㆍ]+(?:㈜)?"
    pattern = re.compile(
        r"공종코드\s*(?P<code>[A-Z]{2}[0-9]{3}\.[0-9]{5})\s*항목은\s*"
        r"(?P<item>[^에]{1,80}?)에\s*관한\s*단가\s*검토\s*건이다\.\s*"
        r"적용\s*단위는\s*(?P<unit>[^\s\.]+)이고\s*표준시장단가는\s*(?P<std>[0-9]{1,3}(?:,[0-9]{3})+)원이다\.\s*"
        r"이번\s*비교에서\s*최저\s*견적은\s*(?P<low_vendor>" + vendor_pat + r")의" + page_header + r"\s*(?P<low_price>[0-9]{1,3}(?:,[0-9]{3})+)원이며,\s*"
        r"최고\s*견적은\s*(?P<high_vendor>" + vendor_pat + r")의" + page_header + r"\s*(?P<high_price>[0-9]{1,3}(?:,[0-9]{3})+)원이다\."
        r"(?P<context>.{0,500}?)"
        r"(?=공종코드\s*[A-Z]{2}[0-9]{3}\.|[가-힣A-Za-z0-9·ㆍ()]+\s*\([A-Z]{2}[0-9]{2}\*\)에 대한 검토 의견|$)",
        re.S,
    )
    for m in pattern.finditer(text):
        full_item = clean_cell_text(m.group('item'))
        compact_item = compact_text(full_item)
        if requested_terms:
            # 변압기 설치(공백) → 변압기설치(compact)로 통일해서 비교
            compact_terms = [compact_text(t) for t in requested_terms]
            if not any(t in compact_item or compact_item in t for t in compact_terms):
                continue
        item_name, spec = _split_item_name_and_spec(full_item)
        low_vendor = clean_cell_text(m.group('low_vendor'))
        high_vendor = clean_cell_text(m.group('high_vendor'))
        low_vendor = re.sub(r"^[A-Z]회사\s*", "", low_vendor).strip()
        high_vendor = re.sub(r"^[A-Z]회사\s*", "", high_vendor).strip()
        low_vendor = re.sub(r"(의|가|이|은|는)$", "", low_vendor).strip()
        high_vendor = re.sub(r"(의|가|이|은|는)$", "", high_vendor).strip()
        low_price = clean_number(m.group('low_price'))
        high_price = clean_number(m.group('high_price'))
        vendor_prices: Dict[str, Any] = {}
        vendor_amounts: Dict[str, Any] = {}
        direct_vendor_map = {
            compact_text(low_vendor): low_price,
            compact_text(high_vendor): high_price,
        }
        # 선택 업체가 있으면 그 업체만, 없으면 최저/최고 업체를 표시한다.
        vendors_for_row = selected_vendors or [
            {"name": low_vendor, "compareKey": compact_text(low_vendor)},
            {"name": high_vendor, "compareKey": compact_text(high_vendor)},
        ]
        for vendor in vendors_for_row:
            name = clean_cell_text(vendor.get('name'))
            c_name = compact_text(name)
            matched_price = ""
            for direct_name, price in direct_vendor_map.items():
                if c_name and direct_name and (c_name in direct_name or direct_name in c_name):
                    matched_price = price
                    break
            vendor_prices[name] = matched_price if matched_price else "원문 미기재"
            if request_qty and matched_price:
                vendor_amounts[name] = int(float(request_qty) * to_number(matched_price))
            elif request_qty:
                vendor_amounts[name] = "원문 미기재"
        code = clean_cell_text(m.group('code'))
        key = compact_text(f"{code}{full_item}{low_vendor}{high_vendor}")
        if key in seen:
            continue
        seen.add(key)
        remark_parts = []
        if request_unit and clean_cell_text(m.group('unit')) and compact_text(request_unit) != compact_text(m.group('unit')):
            remark_parts.append(f"단위 확인: 원문 {clean_cell_text(m.group('unit'))} / 요청 {request_unit}")
        rows.append(enrich_row_units({
            "construction_code": code,
            "item_name": item_name,
            "spec": spec,
            "quantity": request_qty or "",
            "unit": clean_cell_text(m.group('unit')),
            "standard_unit_price": clean_number(m.group('std')),
            "lowest_vendor": low_vendor,
            "highest_vendor": high_vendor,
            "vendor_prices": vendor_prices,
            "vendor_amounts": vendor_amounts,
            "remark": " / ".join(remark_parts),
        }))
    return rows[:80], selected_vendors


def build_text_vendor_comparison_item_table(text: str, user_request: str = "") -> Dict[str, Any] | None:
    if not is_text_only_vendor_comparison_report(text):
        return None
    rows, selected_vendors = extract_text_vendor_item_rows(text, user_request=user_request)
    if not rows:
        return None

    vendor_names: List[str] = []
    for vendor in selected_vendors or []:
        name = clean_cell_text(vendor.get('name'))
        if name and name not in vendor_names:
            vendor_names.append(name)
    if not vendor_names:
        for row in rows:
            for name in (row.get('vendor_prices') or {}).keys():
                clean_name = clean_cell_text(name)
                if clean_name and clean_name not in vendor_names:
                    vendor_names.append(clean_name)

    vendors_meta = []
    vendor_columns: List[Dict[str, Any]] = []
    for idx, name in enumerate(vendor_names):
        unit_price_key = _safe_compare_key(name, idx + 1, "unit_price")
        amount_key = _safe_compare_key(name, idx + 1, "amount")
        vendors_meta.append({
            "index": idx,
            "name": name,
            "vendorName": name,
            "label": name,
            "unitPriceKey": unit_price_key,
            "amountKey": amount_key,
        })
        vendor_columns.extend([
            {"key": unit_price_key, "label": f"{name} 단가"},
            {"key": amount_key, "label": f"{name} 금액"},
        ])

    flattened_rows: List[Dict[str, Any]] = []
    for row in rows:
        out = dict(row)
        vendor_prices = row.get('vendor_prices') or {}
        vendor_amounts = row.get('vendor_amounts') or {}
        for meta in vendors_meta:
            name = meta["name"]
            out[meta["unitPriceKey"]] = vendor_prices.get(name, "")
            out[meta["amountKey"]] = vendor_amounts.get(name, "")
        flattened_rows.append(out)

    return {
        "tableName": "서술형 업체별 단가 비교표",
        "tableType": MULTI_VENDOR_COMPARE_TABLE_TYPE,
        "columns": [
            {"key": "construction_code", "label": "공종코드"},
            {"key": "item_name", "label": "품목명"},
            {"key": "spec", "label": "규격"},
            {"key": "quantity", "label": "수량"},
            {"key": "unit", "label": "단위"},
            {"key": "standard_unit_price", "label": "표준단가"},
            *vendor_columns,
            {"key": "lowest_vendor", "label": "원문 최저업체"},
            {"key": "highest_vendor", "label": "원문 최고업체"},
            {"key": "remark", "label": "비고"},
        ],
        "rows": flattened_rows,
        "meta": {
            "sourceMode": "text_only_vendor_comparison_item_rows",
            "vendors": vendors_meta,
            "templateLayoutMode": "COMPACT_VENDOR_GROUPS",
            "notice": "표가 없는 보고서에서 공종별 최저/최고 견적 문장을 추출했습니다. 요청 업체 단가가 원문에 직접 없으면 원문 미기재로 표시합니다.",
            "requestedText": user_request or "",
        },
    }


def build_text_vendor_comparison_summary_table(text: str, user_request: str = "") -> Dict[str, Any] | None:
    if not is_text_only_vendor_comparison_report(text):
        return None
    rows = extract_text_vendor_total_rows(text)
    return {
        "tableName": "서술형 업체별 단가 비교 요약",
        "tableType": TEXT_VENDOR_COMPARISON_TABLE_TYPE,
        "columns": TEXT_VENDOR_COMPARISON_COLUMNS,
        "rows": rows,
        "meta": {
            "sourceMode": "text_only_vendor_comparison_report",
            "notice": "표가 없는 보고서에서 총괄 업체별 금액만 추출했습니다. 개별 공종 단가는 원문에 명시된 범위에서만 확인해야 합니다.",
            "requestedText": user_request or "",
        },
    }


def _split_text_pages(text: str) -> List[Dict[str, Any]]:
    """PyMuPDF/pdfplumber가 붙인 [page N / T] 마커 기준으로 페이지를 나눈다."""
    if not text:
        return []
    pattern = re.compile(r"\[page\s+(\d+)\s*/\s*(\d+)(?:\s+OCR)?\]", re.I)
    matches = list(pattern.finditer(text))
    if not matches:
        return [{"page": None, "text": text}]
    pages: List[Dict[str, Any]] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        pages.append({"page": int(match.group(1)), "pageCount": int(match.group(2)), "text": text[start:end]})
    return pages


def _clean_line(value: Any, limit: int = 220) -> str:
    clean = re.sub(r"\s+", " ", str(value or "")).strip(" -·•\t")
    return clean[:limit]




def _coalesce_standard_market_lines(page_text: str) -> List[str]:
    """표준시장단가 PDF의 행을 텍스트 라인 기준으로 재조립한다.

    PyMuPDF/pdfplumber는 m², m³, 공m³ 같은 위첨자와 다단 셀을 줄바꿈으로
    분리하는 경우가 많다. 예: `... m` 다음 줄 `2 24,973 63%`.
    이 함수는 공종코드로 시작하는 라인을 기준으로 뒤따르는 줄을 붙여
    코드/명칭/규격/단위/단가/노무비율 행을 복구한다.
    """
    lines = [_clean_line(line, 260) for line in str(page_text or '').splitlines()]
    lines = [line for line in lines if line]
    merged: List[str] = []
    code_re = re.compile(r"^[A-Z]{1,3}\d{3}\.\d{5}\b")
    tail_re = re.compile(r"\s(?:\d|공?m\s*[23]|[㎡㎥]|개|본|식|시간|hr|ton|kg|개소)\s+[0-9]{1,3}(?:,[0-9]{3})+\s+\d{1,3}(?:\.\d+)?%\s*$", re.I)

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if not code_re.match(line):
            idx += 1
            continue
        buf = [line]
        j = idx + 1
        # 한 행은 길어도 5줄 내외다. 다음 코드/단가정의/단가보정이 나오면 중단한다.
        while j < len(lines) and len(buf) < 7:
            candidate = " ".join(buf)
            if tail_re.search(candidate):
                break
            nxt = lines[j]
            if code_re.match(nxt) or nxt.startswith("【") or nxt.startswith("■") or "공종코드" in nxt:
                break
            buf.append(nxt)
            j += 1
        merged.append(" ".join(buf))
        idx = max(j, idx + 1)
    return merged


def _split_standard_market_body(body: str) -> Tuple[str, str, str, str, str]:
    tail = re.search(
        r"(?P<unit>공\s*m\s*[23]|공㎥|m\s*[23]|㎡|㎥|개소|개|본|식|시간|hr|ton|kg|m|대|조|매|장)\s+"
        r"(?P<price>[0-9]{1,3}(?:,[0-9]{3})+|[0-9]{3,})\s+"
        r"(?P<labor>\d{1,3}(?:\.\d+)?%)\s*$",
        body,
        re.I,
    )
    if not tail:
        return "", "", "", "", ""

    prefix = body[:tail.start()].strip()
    unit = tail.group("unit")
    price = tail.group("price")
    labor = tail.group("labor")

    # 규격은 보통 숫자/H=/ℓ=/Type/φ/Φ/D= 등으로 시작한다.
    spec_marker = re.search(r"(?=\b(?:H|B|L|D)\s*=|[ℓøØφΦ∅]|\bType\b|\bTYPE\b|\d|[-–])", prefix)
    if spec_marker and spec_marker.start() > 0:
        item_name = prefix[:spec_marker.start()].strip()
        spec = prefix[spec_marker.start():].strip()
    else:
        parts = prefix.rsplit(" ", 1)
        item_name = parts[0].strip() if len(parts) > 1 else prefix.strip()
        spec = parts[1].strip() if len(parts) > 1 else "-"

    item_name = re.sub(r"\s+", " ", item_name).strip()
    spec = re.sub(r"\s+", " ", spec).strip() or "-"
    return item_name, spec, unit, price, labor


def extract_standard_market_rows_from_text(text: str, max_rows: int | None = None) -> List[Dict[str, Any]]:
    """표준시장단가 문서 전용 텍스트 파서.

    pdfplumber가 표의 왼쪽 공종코드/오른쪽 노무비율 컬럼을 누락하는 PDF가 있어
    PyMuPDF 텍스트 라인에서 해당 값을 보강한다.
    """
    limit = max_rows if max_rows and max_rows > 0 else _env_int("PDF_TABLE_MAX_ROWS", 500)
    if limit <= 0:
        limit = 10**9

    rows: List[Dict[str, Any]] = []
    seen: set[str] = set()
    row_re = re.compile(r"^(?P<code>[A-Z]{1,3}\d{3}\.\d{5})\s+(?P<body>.+)$")
    for page in _split_text_pages(text):
        page_no = page.get("page")
        for line in _coalesce_standard_market_lines(str(page.get("text") or "")):
            match = row_re.match(line)
            if not match:
                continue
            code = match.group("code").strip()
            body = match.group("body").strip()
            item_name, spec, unit, price, labor = _split_standard_market_body(body)
            if not price or not labor:
                continue
            key = compact_text(f"{code}|{item_name}|{spec}|{unit}|{price}|{labor}")
            if key in seen:
                continue
            seen.add(key)
            row = enrich_row_units({
                "construction_code": code,
                "item_name": clean_cell_text(item_name),
                "spec": clean_cell_text(spec),
                "unit": clean_cell_text(unit),
                "unit_price": clean_number(price),
                "labor_ratio": clean_cell_text(labor),
                "source_page": f"p.{page_no}" if page_no else "",
            })
            rows.append(row)
            if len(rows) >= limit:
                return rows
    return rows


def merge_standard_market_rows(table_rows: List[Dict[str, Any]], text_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """pdfplumber 표 행과 PyMuPDF 텍스트 행을 같은 순서 기준으로 병합한다."""
    if not text_rows:
        return table_rows
    if not table_rows:
        return text_rows

    merged: List[Dict[str, Any]] = []
    total = max(len(table_rows), len(text_rows))
    for idx in range(total):
        base = dict(table_rows[idx]) if idx < len(table_rows) else {}
        extra = text_rows[idx] if idx < len(text_rows) else {}
        row = dict(base)
        # pdfplumber가 더 정확한 셀을 가진 경우는 유지하고, 누락된 코드/노무비율/페이지를 보강한다.
        for key in ["construction_code", "labor_ratio", "source_page"]:
            if not str(row.get(key, "")).strip() and str(extra.get(key, "")).strip():
                row[key] = extra.get(key)
        for key in ["item_name", "spec", "unit", "unit_price"]:
            if not str(row.get(key, "")).strip() and str(extra.get(key, "")).strip():
                row[key] = extra.get(key)
        row = enrich_row_units(row)
        merged.append(row)
    return merged


def prune_empty_columns(columns: List[Dict[str, Any]], rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """값이 전혀 없는 컬럼은 응답에서 제외한다."""
    if not rows:
        return columns
    visible: List[Dict[str, Any]] = []
    for col in columns:
        key = col.get("key")
        if not key:
            continue
        if any(str(row.get(key, "")).strip() for row in rows):
            visible.append(col)
    return visible or columns


def _looks_like_section_heading(line: str) -> bool:
    line = _clean_line(line, 120)
    if not line or len(line) > 90:
        return False
    if re.match(r"^(?:제\s*\d+\s*[장절]|\d+(?:\.\d+){0,4}(?:[-ㅡ]\d+)?\.?)\s*[^0-9=]{1,70}$", line):
        return True
    if re.match(r"^[가-힣A-Za-z][가-힣A-Za-z0-9\s/·()\-'\[\]]{1,50}$", line) and any(k in line for k in ["기준", "단가", "방법", "요령", "할증", "품셈"]):
        return True
    return False


def _find_section_heading(lines: List[str], index: int) -> str:
    for back in range(index, max(-1, index - 18), -1):
        if back < 0:
            break
        candidate = _clean_line(lines[back], 120)
        if _looks_like_section_heading(candidate):
            return candidate
    return "본문 기준"


def is_reference_row_supported(row: Dict[str, Any], source_text: str) -> bool:
    if not isinstance(row, dict):
        return False
    basis = str(row.get("basis_item") or row.get("application_basis") or row.get("unit_price_basis") or "").strip()
    if not basis:
        return False
    # 기준서 표는 문장 일부가 원문에 있으면 근거 있음으로 판단한다.
    for key in ("basis_item", "application_basis", "calculation_method", "unit_price_basis"):
        value = str(row.get(key) or "").strip()
        if value and source_has_value(source_text, value[:80]):
            return True
    return False


def extract_reference_guideline_rows(text: str, user_request: str = "", max_rows: int | None = None) -> List[Dict[str, Any]]:
    """기준서/지침서 문서를 '품목 단가표'가 아니라 '기준 항목 표'로 변환한다.

    원문에 없는 품목·금액을 만들지 않고, 문서 안에 실제 등장한 기준/단가/산정 문장만 행으로 만든다.
    """
    if not is_reference_or_guideline_document(text):
        return []
    max_rows = max_rows or max(_env_int("REFERENCE_TABLE_MAX_ROWS", 120), 20)
    request = compact_text(user_request)
    focus_price = any(k in request for k in ["단가", "가격", "견적", "비교", "원가"])
    focus_table = any(k in request for k in ["표", "테이블", "정리", "엑셀", "양식"])
    keywords = ["단가", "가격", "계약단가", "거래실례가격", "노임단가", "정부노임단가", "기본요금"] if focus_price else REFERENCE_TABLE_KEYWORDS
    if not focus_table and not focus_price:
        keywords = ["적용기준", "산정기준", "계산", "단가", "가격", "요율"]

    rows: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for page in _split_text_pages(text):
        page_no = page.get("page")
        raw_lines = [line for line in str(page.get("text") or "").splitlines()]
        lines = [_clean_line(line) for line in raw_lines]
        for idx, line in enumerate(lines):
            if not line or len(line) < 5:
                continue
            if not any(keyword in line for keyword in keywords):
                continue

            section = _find_section_heading(lines, idx)
            context_parts = [line]
            # 한 줄짜리 PDF 줄바꿈을 완화하기 위해 바로 다음 1~2줄만 근거 문맥으로 붙인다.
            for next_idx in range(idx + 1, min(len(lines), idx + 3)):
                nxt = lines[next_idx]
                if not nxt or _looks_like_section_heading(nxt):
                    break
                if len(" ".join(context_parts)) < 180:
                    context_parts.append(nxt)
            basis = _clean_line(" ".join(context_parts), 260)
            calculation = ""
            if re.search(r"[=*×xX]|곱|나누|계산|산출|적용|가산|공제|요율|%", basis):
                calculation = basis
            unit_basis = basis if any(k in basis for k in ["단가", "가격", "요금", "노임", "거래실례"]) else ""
            row_key = compact_text(f"{page_no}|{section}|{basis[:120]}")
            if row_key in seen:
                continue
            seen.add(row_key)
            rows.append({
                "section": section,
                "basis_item": section if section != "본문 기준" else basis[:40],
                "application_basis": basis,
                "calculation_method": calculation,
                "unit_price_basis": unit_basis,
                "source_page": f"p.{page_no}" if page_no else "",
                "remark": "기준서 원문 키워드 기반 추출",
            })
            if len(rows) >= max_rows:
                return rows

    # 키워드가 너무 적을 때는 주요 장절 제목만 최소 표로 제공한다.
    if not rows:
        for page in _split_text_pages(text):
            page_no = page.get("page")
            for line in [_clean_line(v) for v in str(page.get("text") or "").splitlines()]:
                if _looks_like_section_heading(line) and any(k in line for k in ["기준", "단가", "계산", "산출", "품셈"]):
                    row_key = compact_text(f"{page_no}|{line}")
                    if row_key in seen:
                        continue
                    seen.add(row_key)
                    rows.append({
                        "section": line,
                        "basis_item": line,
                        "application_basis": line,
                        "calculation_method": "",
                        "unit_price_basis": line if any(k in line for k in ["단가", "가격", "요금"]) else "",
                        "source_page": f"p.{page_no}" if page_no else "",
                        "remark": "기준서 장절 제목 기반 추출",
                    })
                    if len(rows) >= max_rows:
                        return rows
    return rows



def _strip_header_unit_suffix(header: Any) -> str:
    text = clean_cell_text(header)
    text = re.sub(r"[\(\[\{]\s*(원|만원|천원|%|퍼센트|금액|price|krw)\s*[\)\]\}]", "", text, flags=re.I)
    return clean_cell_text(text)


def _field_alias_score(header: Any, aliases: List[str]) -> int:
    compact = compact_text(header)
    best = 0
    for alias in aliases:
        c_alias = compact_text(alias)
        if not c_alias:
            continue
        if compact == c_alias:
            best = max(best, 100 + len(c_alias))
        elif c_alias in compact:
            best = max(best, 50 + len(c_alias))
    return best


def _classify_header_role(header: Any) -> Dict[str, Any]:
    """헤더 역할을 분류한다. A회사/B회사 같은 특정 문자열을 조건으로 쓰지 않는다."""
    raw = clean_cell_text(header)
    compact = compact_text(_strip_header_unit_suffix(raw))
    if not compact:
        return {"role": "field", "source_header": raw, "confidence": 0.0}
    scored: List[Tuple[int, str]] = []
    for key, aliases in FIELD_ALIASES.items():
        score = _field_alias_score(compact, aliases)
        if score:
            scored.append((score, key))
    if any(term in compact for term in ["표준", "기준", "시장단가"]):
        scored.append((220, "standard_unit_price"))
    if any(term in compact for term in ["견적", "제안", "제시", "업체", "공급"]):
        scored.append((210, "vendor_unit_price"))
    if any(term in compact for term in ["차이", "차액", "증감"]):
        scored.append((205, "price_diff"))
    if "%" in raw or any(term in compact for term in ["대비", "비율", "증감률"]):
        scored.append((205, "diff_rate"))
    if scored:
        scored.sort(reverse=True)
        return {"role": scored[0][1], "source_header": raw, "confidence": min(scored[0][0] / 220, 1.0)}
    return {"role": compact or "field", "source_header": raw, "confidence": 0.0}


def _normalize_price_role_fallback(row: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(row, dict):
        return row
    if row.get("vendor_unit_price") and not row.get("unit_price"):
        row["unit_price"] = row.get("vendor_unit_price")
    if row.get("standard_unit_price") and row.get("vendor_unit_price"):
        row["source_price_role"] = "vendor_unit_price"
    elif row.get("vendor_unit_price"):
        row["source_price_role"] = "vendor_unit_price"
    elif row.get("unit_price"):
        row["source_price_role"] = "unit_price"
    elif row.get("amount"):
        row["source_price_role"] = "amount"
    return row


def normalize_header(header: Any) -> str:
    return str(_classify_header_role(header).get("role") or "field")


def clean_number(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    # 숫자, 음수, 소수, 천 단위 구분자만 남긴다. 원/₩ 등은 제거한다.
    cleaned = re.sub(r"[^0-9.,-]", "", text)
    # 1,000.00 형태는 유지하고, 1.000,00 같은 유럽식은 운영 범위 밖이므로 원문 확인 대상으로 둔다.
    return cleaned


def to_number(value: Any) -> float:
    cleaned = str(value or "").replace(",", "").strip()
    if not cleaned:
        return 0.0
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def rows_to_table(raw_rows: List[List[Any]], source_text: str = "", filename: str = "") -> List[Dict[str, Any]]:
    if not raw_rows:
        return []

    cleaned_rows: List[List[str]] = []
    for row in raw_rows:
        values = ["" if v is None else clean_cell_text(v) for v in row]
        while values and not values[-1]:
            values.pop()
        if any(values):
            cleaned_rows.append(values)
    if not cleaned_rows:
        return []

    header_idx = 0
    best_score = -1
    for idx, row in enumerate(cleaned_rows[:25]):
        joined = " ".join(map(str, row))
        score = 0
        for cell in row:
            if _classify_header_role(cell).get("confidence", 0) > 0:
                score += 1
        score += sum(1 for aliases in FIELD_ALIASES.values() for alias in aliases if alias.lower() in joined.lower())
        if score > best_score:
            best_score = score
            header_idx = idx
        if score >= 3:
            break

    if best_score < 2:
        return []

    raw_headers = cleaned_rows[header_idx]
    headers = [normalize_header(cell) for cell in raw_headers]
    for i in range(1, len(headers)):
        raw_cell = re.sub(r"\s+", "", str(raw_headers[i] or "")).lower()
        prev_raw = re.sub(r"\s+", "", str(raw_headers[i - 1] or "")).lower()
        if headers[i] in {"field", "가", "unit_price"} and raw_cell in {"가", "가격"} and (headers[i - 1] == "unit" or "단위" in prev_raw):
            headers[i] = "unit_price"

    normalized_headers: List[str] = []
    seen: Dict[str, int] = {}
    for col in headers:
        base = col or "field"
        seen[base] = seen.get(base, 0) + 1
        normalized_headers.append(base if seen[base] == 1 else f"{base}_{seen[base]}")

    rows: List[Dict[str, Any]] = []
    for raw in cleaned_rows[header_idx + 1 : header_idx + 151]:
        if not any(str(v or "").strip() for v in raw):
            continue
        item: Dict[str, Any] = {}
        for idx, (key, value) in enumerate(zip(normalized_headers, raw)):
            base_key = re.sub(r"_\d+$", "", key)
            if base_key in NUMBER_KEYS:
                item[key] = clean_number(value)
            elif base_key == "diff_rate":
                item[key] = clean_cell_text(value)
            else:
                item[key] = clean_cell_text(value)
            if idx < len(raw_headers):
                item.setdefault("_source_headers", {})[key] = raw_headers[idx]

        for key in list(item.keys()):
            if re.match(r"^vendor_unit_price_\d+$", key) and not item.get("vendor_unit_price"):
                item["vendor_unit_price"] = item.get(key)
            if re.match(r"^standard_unit_price_\d+$", key) and not item.get("standard_unit_price"):
                item["standard_unit_price"] = item.get(key)
            if re.match(r"^price_diff_\d+$", key) and not item.get("price_diff"):
                item["price_diff"] = item.get(key)
            if re.match(r"^diff_rate_\d+$", key) and not item.get("diff_rate"):
                item["diff_rate"] = item.get(key)

        item = _normalize_price_role_fallback(item)
        rows.append(enrich_row_units(item))
    return rows


def table_to_markdown(rows: List[Dict[str, Any]], max_rows: int = 40) -> str:
    if not rows:
        return ""
    preferred = ["vendor_name", "construction_code", "item_name", "spec", "quantity", "unit", "standard_unit_price", "vendor_unit_price", "unit_price", "amount", "price_diff", "diff_rate", "remark"]
    keys = [key for key in preferred if any(str(row.get(key, "")).strip() for row in rows)]
    if not keys:
        seen = []
        for row in rows[:max_rows]:
            for key in row.keys():
                if key.startswith("_"):
                    continue
                if key not in seen:
                    seen.append(key)
        keys = seen[:12] or [col["key"] for col in DEFAULT_COLUMNS]
    header = "| " + " | ".join(keys) + " |"
    sep = "| " + " | ".join(["---"] * len(keys)) + " |"
    body = []
    for row in rows[:max_rows]:
        body.append("| " + " | ".join(str(row.get(key, "")).replace("\n", " ")[:80] for key in keys) + " |")
    return "\n".join([header, sep, *body])


def read_xlsx(content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return "", []

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    all_rows: List[List[Any]] = []
    text_lines: List[str] = []
    for ws in wb.worksheets[:5]:
        text_lines.append(f"[sheet {ws.title}]")
        sheet_rows: List[List[Any]] = []
        for row in ws.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                sheet_rows.append(values)
                all_rows.append(values)
                text_lines.append("\t".join(values))
            if len(sheet_rows) >= 200:
                break
    return "\n".join(text_lines), rows_to_table(all_rows)


def _read_pdf_with_pymupdf(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, int | None, List[Dict[str, Any]], Dict[str, Any]]:
    """PyMuPDF 기반 PDF 텍스트 추출.

    - OCR을 사용하지 않는다.
    - 기본값은 전체 페이지 추출이다(PDF_MAX_PAGES=0).
    - pypdf의 /UniKS-UTF16-H 경고를 피하기 위해 pypdf를 기본 경로에서 제외한다.
    """
    started = time.perf_counter()
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        if log:
            log("error", "PyMuPDF import failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pymupdf", "error": str(exc)}

    try:
        doc = fitz.open(stream=content, filetype="pdf")
        page_count = int(doc.page_count or 0)
        pages_to_read = _limit_pages(page_count, "PDF_MAX_PAGES")
        log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
        parts: List[str] = []
        page_meta: List[Dict[str, Any]] = []
        empty_pages = 0
        total_chars = 0

        if log:
            log(
                "info",
                "PDF parse start",
                filename=filename,
                engine="PyMuPDF",
                text_layer=True,
                page_count=page_count,
                pages_to_read=pages_to_read,
                size_bytes=len(content),
            )

        for idx in range(pages_to_read):
            page = doc.load_page(idx)
            text = page.get_text("text", sort=True) or ""
            char_count = len(text)
            total_chars += char_count
            if not text.strip():
                empty_pages += 1
            parts.append(f"[page {idx + 1} / {page_count}]\n{text}".rstrip())
            page_meta.append({
                "page": idx + 1,
                "pageCount": page_count,
                "engine": "PyMuPDF",
                "status": "TEXT_EXTRACTED" if text.strip() else "NO_TEXT_LAYER",
                "charCount": char_count,
            })
            if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                log(
                    "info",
                    "PDF parse progress",
                    filename=filename,
                    page=f"{idx + 1}/{page_count}",
                    chars=total_chars,
                    empty_pages=empty_pages,
                )

        elapsed = round(time.perf_counter() - started, 3)
        truncated_by_limit = page_count > pages_to_read
        if log:
            log(
                "info",
                "PDF parse finish",
                filename=filename,
                engine="PyMuPDF",
                text_layer=True,
                page_count=page_count,
                pages_read=pages_to_read,
                chars=total_chars,
                empty_pages=empty_pages,
                elapsed_sec=elapsed,
                truncated_by_limit=truncated_by_limit,
            )
        return "\n\n".join(parts), page_count, page_meta, {
            "engine": "PyMuPDF",
            "noOcr": True,
            "pageCount": page_count,
            "pagesRead": pages_to_read,
            "charCount": total_chars,
            "emptyPages": empty_pages,
            "elapsedSec": elapsed,
            "truncatedByLimit": truncated_by_limit,
        }
    except Exception as exc:
        if log:
            log("error", "PyMuPDF parse failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "PyMuPDF", "noOcr": True, "error": str(exc)}


def _read_pdf_text_with_pdfplumber(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, int | None, List[Dict[str, Any]], Dict[str, Any]]:
    try:
        import pdfplumber
    except Exception as exc:
        if log:
            log("error", "pdfplumber import failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pdfplumber", "error": str(exc)}

    started = time.perf_counter()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            pages_to_read = _limit_pages(page_count, "PDF_MAX_PAGES")
            log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
            parts: List[str] = []
            page_meta: List[Dict[str, Any]] = []
            empty_pages = 0
            total_chars = 0
            if log:
                log("info", "PDF fallback parse start", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_to_read=pages_to_read)
            for idx, page in enumerate(pdf.pages[:pages_to_read]):
                text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                char_count = len(text)
                total_chars += char_count
                if not text.strip():
                    empty_pages += 1
                parts.append(f"[page {idx + 1} / {page_count}]\n{text}".rstrip())
                page_meta.append({
                    "page": idx + 1,
                    "pageCount": page_count,
                    "engine": "pdfplumber",
                    "status": "TEXT_EXTRACTED" if text.strip() else "NO_TEXT_LAYER",
                    "charCount": char_count,
                })
                if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                    log("info", "PDF fallback parse progress", filename=filename, page=f"{idx + 1}/{page_count}", chars=total_chars, empty_pages=empty_pages)
            elapsed = round(time.perf_counter() - started, 3)
            if log:
                log("info", "PDF fallback parse finish", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_read=pages_to_read, chars=total_chars, empty_pages=empty_pages, elapsed_sec=elapsed, truncated_by_limit=page_count > pages_to_read)
            return "\n\n".join(parts), page_count, page_meta, {
                "engine": "pdfplumber",
                "noOcr": True,
                "pageCount": page_count,
                "pagesRead": pages_to_read,
                "charCount": total_chars,
                "emptyPages": empty_pages,
                "elapsedSec": elapsed,
                "truncatedByLimit": page_count > pages_to_read,
            }
    except Exception as exc:
        if log:
            log("error", "pdfplumber fallback parse failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pdfplumber", "noOcr": True, "error": str(exc)}


def _read_pdf_tables_with_pdfplumber(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
    source_text: str = "",
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    if not _env_bool("PDF_EXTRACT_TABLES", True):
        if log:
            log("info", "PDF table extraction skipped", filename=filename, reason="PDF_EXTRACT_TABLES=false")
        return rows, {"enabled": False, "rowCount": 0}

    try:
        import pdfplumber
    except Exception as exc:
        if log:
            log("warning", "pdfplumber table import failed", filename=filename, error=str(exc))
        return rows, {"enabled": True, "error": str(exc), "rowCount": 0}

    started = time.perf_counter()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            pages_to_read = _limit_pages(page_count, "PDF_TABLE_MAX_PAGES")
            configured_max_rows = _env_int("PDF_TABLE_MAX_ROWS", 5000)
            max_rows = configured_max_rows if configured_max_rows > 0 else 10**9
            log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
            if log:
                log("info", "PDF table extraction start", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_to_read=pages_to_read)
            for idx, page in enumerate(pdf.pages[:pages_to_read]):
                table_count = 0
                for table in page.extract_tables() or []:
                    table_count += 1
                    parsed = rows_to_table(table, source_text=source_text, filename=filename)
                    rows.extend(parsed)
                    if len(rows) >= max_rows:
                        rows = rows[:max_rows]
                        if log:
                            log("warning", "PDF table row limit reached", filename=filename, row_limit=max_rows, page=idx + 1)
                        elapsed = round(time.perf_counter() - started, 3)
                        return rows, {"enabled": True, "engine": "pdfplumber", "pageCount": page_count, "pagesRead": idx + 1, "rowCount": len(rows), "elapsedSec": elapsed, "rowLimitReached": True}
                if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                    log("info", "PDF table extraction progress", filename=filename, page=f"{idx + 1}/{page_count}", rows=len(rows), tables_on_page=table_count)
            elapsed = round(time.perf_counter() - started, 3)
            if log:
                log("info", "PDF table extraction finish", filename=filename, engine="pdfplumber", page_count=page_count, pages_read=pages_to_read, rows=len(rows), elapsed_sec=elapsed, truncated_by_limit=page_count > pages_to_read)
            return rows, {"enabled": True, "engine": "pdfplumber", "pageCount": page_count, "pagesRead": pages_to_read, "rowCount": len(rows), "elapsedSec": elapsed, "truncatedByLimit": page_count > pages_to_read}
    except Exception as exc:
        if log:
            log("warning", "PDF table extraction failed", filename=filename, error=str(exc))
        return rows, {"enabled": True, "error": str(exc), "rowCount": len(rows)}


def read_pdf(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, List[Dict[str, Any]], int | None, List[Dict[str, Any]], Dict[str, Any]]:
    preferred = os.getenv("PDF_TEXT_ENGINE", "pymupdf").strip().lower()
    if preferred not in {"pymupdf", "pdfplumber"}:
        preferred = "pymupdf"

    if preferred == "pdfplumber":
        text, page_count, page_meta, text_metrics = _read_pdf_text_with_pdfplumber(content, filename=filename, log=log)
        if len(text.strip()) < 80:
            fallback_text, fallback_count, fallback_pages, fallback_metrics = _read_pdf_with_pymupdf(content, filename=filename, log=log)
            if len(fallback_text.strip()) > len(text.strip()):
                text, page_count, page_meta, text_metrics = fallback_text, fallback_count, fallback_pages, fallback_metrics
    else:
        text, page_count, page_meta, text_metrics = _read_pdf_with_pymupdf(content, filename=filename, log=log)
        if len(text.strip()) < 80:
            fallback_text, fallback_count, fallback_pages, fallback_metrics = _read_pdf_text_with_pdfplumber(content, filename=filename, log=log)
            if len(fallback_text.strip()) > len(text.strip()):
                text, page_count, page_meta, text_metrics = fallback_text, fallback_count, fallback_pages, fallback_metrics

    rows, table_metrics = _read_pdf_tables_with_pdfplumber(content, filename=filename, log=log, source_text=text)

    ocr_metrics: Dict[str, Any] = {"enabled": _ocr_enabled(), "ocrUsed": False}
    should_ocr = _ocr_enabled() and (_env_bool("OCR_FORCE", False) or (len(text.strip()) < _ocr_min_text_chars() and not rows))
    if should_ocr:
        ocr_text, ocr_rows, ocr_pages, ocr_metrics = _read_pdf_with_ocr_fallback(content, filename=filename, page_count=page_count, log=log)
        if ocr_text.strip():
            text = (text + "\n\n" + ocr_text).strip() if text.strip() else ocr_text
        if ocr_rows:
            rows = ocr_rows if not rows else [*rows, *ocr_rows]
        if ocr_pages:
            page_meta = page_meta or []
            page_meta.extend(ocr_pages)

    metrics = {"text": text_metrics, "tables": table_metrics, "ocr": ocr_metrics, "ocrUsed": bool(ocr_metrics.get("ocrUsed"))}
    return text, rows, page_count, page_meta, metrics



# -----------------------------------------------------------------------------
# PP-Structure / PaddleOCR fallback
# -----------------------------------------------------------------------------

PADDLE_OCR_ENGINE = None
PADDLE_STRUCTURE_ENGINE = None


def _ocr_enabled() -> bool:
    return _env_bool("OCR_ENABLED", True)


def _ocr_lang() -> str:
    return os.getenv("OCR_LANG", "korean").strip() or "korean"


def _ocr_dpi() -> int:
    return max(_env_int("OCR_DPI", 160), 96)


def _ocr_max_pages(total_pages: int) -> int:
    # 0 또는 음수는 전체 페이지 OCR. 기본값 10은 스캔 PDF에서 속도 폭주를 막기 위한 값이다.
    limit = _env_int("OCR_MAX_PAGES", 10)
    if limit <= 0:
        return total_pages
    return min(total_pages, limit)


def _ocr_min_text_chars() -> int:
    return max(_env_int("OCR_MIN_TEXT_CHARS", 80), 0)


def _get_paddle_ocr_engine():
    global PADDLE_OCR_ENGINE
    if PADDLE_OCR_ENGINE is not None:
        return PADDLE_OCR_ENGINE
    from paddleocr import PaddleOCR
    try:
        PADDLE_OCR_ENGINE = PaddleOCR(use_angle_cls=True, lang=_ocr_lang(), show_log=False)
    except TypeError:
        # PaddleOCR 3.x 계열 일부 옵션 호환 처리
        PADDLE_OCR_ENGINE = PaddleOCR(lang=_ocr_lang())
    return PADDLE_OCR_ENGINE


def _get_pp_structure_engine():
    global PADDLE_STRUCTURE_ENGINE
    if PADDLE_STRUCTURE_ENGINE is not None:
        return PADDLE_STRUCTURE_ENGINE
    from paddleocr import PPStructure
    try:
        PADDLE_STRUCTURE_ENGINE = PPStructure(show_log=False, lang=_ocr_lang())
    except TypeError:
        PADDLE_STRUCTURE_ENGINE = PPStructure(lang=_ocr_lang())
    return PADDLE_STRUCTURE_ENGINE


def _image_bytes_to_numpy(image_bytes: bytes):
    from PIL import Image
    import numpy as np

    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    return np.array(image)


def _render_pdf_page_to_png(content: bytes, page_index: int, dpi: int) -> bytes:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    page = doc.load_page(page_index)
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return pix.tobytes("png")


def _flatten_paddle_ocr_result(result: Any) -> List[Tuple[str, float]]:
    """PaddleOCR 2.x/3.x 응답 차이를 최대한 흡수해 (text, score) 목록으로 변환한다."""
    pairs: List[Tuple[str, float]] = []

    def walk(node: Any) -> None:
        if node is None:
            return
        if isinstance(node, dict):
            text = node.get("text") or node.get("rec_text") or node.get("transcription")
            score = node.get("score") or node.get("confidence") or node.get("rec_score") or 0
            if text:
                try:
                    pairs.append((str(text), float(score or 0)))
                except Exception:
                    pairs.append((str(text), 0.0))
            for value in node.values():
                if isinstance(value, (list, tuple, dict)):
                    walk(value)
            return
        if isinstance(node, (list, tuple)):
            if len(node) >= 2 and isinstance(node[1], (list, tuple)) and len(node[1]) >= 1 and isinstance(node[1][0], str):
                text = node[1][0]
                score = node[1][1] if len(node[1]) > 1 else 0
                try:
                    pairs.append((str(text), float(score or 0)))
                except Exception:
                    pairs.append((str(text), 0.0))
                return
            for item in node:
                walk(item)

    walk(result)
    # 중복 제거: PP-Structure/PaddleOCR 응답을 같이 훑으면 같은 문자가 중복될 수 있다.
    seen = set()
    deduped: List[Tuple[str, float]] = []
    for text, score in pairs:
        clean = re.sub(r"\s+", " ", text).strip()
        if not clean or clean in seen:
            continue
        seen.add(clean)
        deduped.append((clean, score))
    return deduped


def _html_table_to_rows(html: str) -> List[Dict[str, Any]]:
    if not html:
        return []
    table_rows: List[List[str]] = []
    tr_matches = re.findall(r"<tr[^>]*>(.*?)</tr>", html, flags=re.I | re.S)
    for tr in tr_matches:
        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", tr, flags=re.I | re.S)
        cleaned = []
        for cell in cells:
            value = re.sub(r"<[^>]+>", " ", cell)
            value = re.sub(r"\s+", " ", value).strip()
            cleaned.append(value)
        if any(cleaned):
            table_rows.append(cleaned)
    return rows_to_table(table_rows)


def _extract_pp_structure_rows(image_np: Any, log: Callable[..., None] | None = None, filename: str = "") -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    if not _env_bool("OCR_USE_PP_STRUCTURE", True):
        return [], {"enabled": False, "reason": "OCR_USE_PP_STRUCTURE=false"}
    try:
        engine = _get_pp_structure_engine()
        result = engine(image_np)
    except Exception as exc:
        if log:
            log("warning", "PP-Structure failed", filename=filename, error=str(exc))
        return [], {"enabled": True, "engine": "PP-Structure", "error": str(exc), "rowCount": 0}

    rows: List[Dict[str, Any]] = []
    block_count = 0
    for block in result or []:
        block_count += 1
        if not isinstance(block, dict):
            continue
        block_type = str(block.get("type") or "").lower()
        res = block.get("res")
        html = ""
        if isinstance(res, dict):
            html = str(res.get("html") or res.get("html_table") or "")
        elif isinstance(res, str):
            html = res
        if "table" in block_type or "<table" in html.lower():
            rows.extend(_html_table_to_rows(html))
    return rows, {"enabled": True, "engine": "PP-Structure", "blockCount": block_count, "rowCount": len(rows)}


def read_image_with_ocr(
    content: bytes,
    filename: str = "",
    page_label: str | int | None = None,
    log: Callable[..., None] | None = None,
) -> Tuple[str, List[Dict[str, Any]], Dict[str, Any]]:
    """이미지 1장을 PP-Structure/PaddleOCR로 처리한다.

    반환값은 (text, rows, metrics). 이미지 LLM은 절대 호출하지 않는다.
    """
    if not _ocr_enabled():
        return "", [], {"enabled": False, "ocrUsed": False, "reason": "OCR_ENABLED=false"}

    started = time.perf_counter()
    try:
        image_np = _image_bytes_to_numpy(content)
    except Exception as exc:
        if log:
            log("warning", "OCR image decode failed", filename=filename, page=page_label, error=str(exc))
        return "", [], {"enabled": True, "ocrUsed": False, "engine": "PIL", "error": str(exc)}

    rows: List[Dict[str, Any]] = []
    structure_metrics: Dict[str, Any] = {"enabled": False}
    if _env_bool("OCR_USE_PP_STRUCTURE", True):
        rows, structure_metrics = _extract_pp_structure_rows(image_np, log=log, filename=filename)

    try:
        ocr = _get_paddle_ocr_engine()
        if hasattr(ocr, "ocr"):
            result = ocr.ocr(image_np, cls=True)
        elif hasattr(ocr, "predict"):
            result = ocr.predict(image_np)
        else:
            raise RuntimeError("PaddleOCR engine has neither ocr() nor predict().")
        pairs = _flatten_paddle_ocr_result(result)
        text = "\n".join(t for t, _ in pairs)
        avg_conf = round(sum(score for _, score in pairs) / len(pairs), 4) if pairs else 0.0
    except Exception as exc:
        if log:
            log("warning", "PaddleOCR text extraction failed", filename=filename, page=page_label, error=str(exc))
        return "", rows, {
            "enabled": True,
            "ocrUsed": bool(rows),
            "engine": "PaddleOCR",
            "structure": structure_metrics,
            "error": str(exc),
            "rowCount": len(rows),
        }

    if not rows:
        rows = infer_rows_from_text(text, filename or "ocr-image")

    elapsed = round(time.perf_counter() - started, 3)
    if log:
        log("info", "OCR image parse finish", filename=filename, page=page_label, chars=len(text), rows=len(rows), confidence=avg_conf, elapsed_sec=elapsed)
    return text, rows, {
        "enabled": True,
        "ocrUsed": True,
        "engine": "PP-Structure/PaddleOCR" if structure_metrics.get("enabled") else "PaddleOCR",
        "structure": structure_metrics,
        "charCount": len(text),
        "rowCount": len(rows),
        "confidence": avg_conf,
        "elapsedSec": elapsed,
    }


def _read_pdf_with_ocr_fallback(
    content: bytes,
    filename: str = "",
    page_count: int | None = None,
    log: Callable[..., None] | None = None,
) -> Tuple[str, List[Dict[str, Any]], List[Dict[str, Any]], Dict[str, Any]]:
    if not _ocr_enabled():
        return "", [], [], {"enabled": False, "ocrUsed": False, "reason": "OCR_ENABLED=false"}

    started = time.perf_counter()
    try:
        import fitz  # noqa: F401  # PyMuPDF import check
    except Exception as exc:
        if log:
            log("warning", "PDF OCR skipped", filename=filename, reason="PyMuPDF import failed", error=str(exc))
        return "", [], [], {"enabled": True, "ocrUsed": False, "error": str(exc)}

    if page_count is None:
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            page_count = int(doc.page_count or 0)
        except Exception:
            page_count = 0

    pages_to_ocr = _ocr_max_pages(int(page_count or 0))
    dpi = _ocr_dpi()
    text_parts: List[str] = []
    all_rows: List[Dict[str, Any]] = []
    page_meta: List[Dict[str, Any]] = []

    if log:
        log("info", "PDF OCR fallback start", filename=filename, engine="PP-Structure/PaddleOCR", page_count=page_count, pages_to_ocr=pages_to_ocr, dpi=dpi)

    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
    except Exception as exc:
        if log:
            log("warning", "PDF OCR render open failed", filename=filename, error=str(exc))
        return "", [], [], {"enabled": True, "ocrUsed": False, "engine": "PP-Structure/PaddleOCR", "error": str(exc)}

    for idx in range(pages_to_ocr):
        try:
            page = doc.load_page(idx)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            png = pix.tobytes("png")
            page_text, page_rows, metrics = read_image_with_ocr(png, filename=filename, page_label=idx + 1, log=log)
        except Exception as exc:
            page_text, page_rows, metrics = "", [], {"ocrUsed": False, "error": str(exc)}
        text_parts.append(f"[page {idx + 1} / {page_count} OCR]\n{page_text}".rstrip())
        all_rows.extend(page_rows)
        page_meta.append({
            "page": idx + 1,
            "pageCount": page_count,
            "engine": "PP-Structure/PaddleOCR",
            "status": "OCR_EXTRACTED" if page_text.strip() or page_rows else "OCR_EMPTY",
            "charCount": len(page_text),
            "rowCount": len(page_rows),
            "metrics": metrics,
        })

    elapsed = round(time.perf_counter() - started, 3)
    metrics = {
        "enabled": True,
        "ocrUsed": True,
        "engine": "PP-Structure/PaddleOCR",
        "pageCount": page_count,
        "pagesRead": pages_to_ocr,
        "charCount": sum(len(part) for part in text_parts),
        "rowCount": len(all_rows),
        "dpi": dpi,
        "elapsedSec": elapsed,
        "truncatedByLimit": bool(page_count and page_count > pages_to_ocr),
    }
    if log:
        log("info", "PDF OCR fallback finish", filename=filename, pages_ocr=pages_to_ocr, chars=metrics["charCount"], rows=len(all_rows), elapsed_sec=elapsed)
    return "\n\n".join(text_parts), all_rows, page_meta, metrics


def read_docx(content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_rows: List[List[Any]] = []
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                raw_rows.append(values)
        table_lines = ["\t".join(map(str, row)) for row in raw_rows]
        return "\n".join(paragraphs + table_lines), rows_to_table(raw_rows)
    except Exception:
        return "", []


def decode_text(content: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def parse_delimited_text(text: str) -> List[Dict[str, Any]]:
    lines = [line for line in text.splitlines() if line.strip()]
    if len(lines) < 2:
        return []

    sample = "\n".join(lines[:30])
    delimiter_candidates = ["\t", ",", "|", ";"]
    delimiter = max(delimiter_candidates, key=lambda d: sample.count(d))
    if sample.count(delimiter) < 2:
        spaced_rows = [re.split(r"\s{2,}", line.strip()) for line in lines if len(re.split(r"\s{2,}", line.strip())) >= 3]
        return rows_to_table(spaced_rows) if len(spaced_rows) >= 2 else []

    reader = csv.reader(lines, delimiter=delimiter)
    raw_rows = [row for row in reader]
    return rows_to_table(raw_rows)


def infer_rows_from_text(text: str, filename: str) -> List[Dict[str, Any]]:
    """텍스트에서 업무 표 행을 보수적으로 추정한다.

    이전 버전처럼 본문에 숫자나 '견적서'라는 단어가 일부 있다는 이유로 임의 행을 만들지 않는다.
    표 형태가 명확하지 않으면 빈 배열을 반환하고, 분석 요약/확인 필요로 처리한다.
    """
    if is_text_only_vendor_comparison_report(text):
        return []
    if is_reference_or_guideline_document(text):
        return extract_reference_guideline_rows(text, user_request="")
    if is_narrative_document(text):
        return []

    table_rows = parse_delimited_text(text)
    table_rows = filter_grounded_rows(table_rows, text)
    if table_rows:
        return table_rows

    rows: List[Dict[str, Any]] = []
    # 예: 품목명 규격 수량 단위 단가 금액 형태가 한 줄에 존재할 때만 추정한다.
    line_pattern = re.compile(
        r"(?P<item>[가-힣A-Za-z0-9()\[\]#/+\-.\s]{2,40})\s+"
        r"(?P<spec>[A-Za-z0-9가-힣ΦØ㎡㎥㎜㎝㎏./*xX\-\s]{0,30})\s+"
        r"(?P<qty>[0-9]+(?:\.[0-9]+)?)\s*"
        r"(?P<unit>EA|PCS|SET|BOX|LOT|M|m|개|식|본|장|대|조|롤|포|kg|KG|㎡|㎥)?\s+"
        r"(?P<unit_price>[0-9]{1,3}(?:,[0-9]{3})+|[0-9]{3,})\s+"
        r"(?P<amount>[0-9]{1,3}(?:,[0-9]{3})+|[0-9]{3,})"
    )
    for line in text.splitlines()[:1000]:
        if not line.strip():
            continue
        m = line_pattern.search(line)
        if not m:
            continue
        row = enrich_row_units({
            "vendor_name": "",
            "item_name": m.group("item").strip(),
            "spec": m.group("spec").strip(),
            "quantity": clean_number(m.group("qty")),
            "unit": (m.group("unit") or "").strip(),
            "unit_price": clean_number(m.group("unit_price")),
            "amount": clean_number(m.group("amount")),
            "remark": "텍스트 행에서 추정",
        })
        if is_business_row_supported(row, text):
            rows.append(row)
        if len(rows) >= 100:
            break
    return rows


def validate_rows(rows: List[Dict[str, Any]], table_type: str = "NORMAL_TABLE") -> List[Dict[str, Any]]:
    # 기준서/일반 표준단가표는 업체별 비교표가 아니므로 행별 단위 경고를 남발하지 않는다.
    if table_type in {"IMAGE_TABLE", "GENERAL_TABLE", "OCR_TABLE", "REFERENCE_GUIDELINE_TABLE", "GUIDELINE_SUMMARY_TABLE", "STANDARD_MARKET_PRICE_TABLE", TEXT_VENDOR_COMPARISON_TABLE_TYPE}:
        return []

    issues: List[Dict[str, Any]] = []
    for idx, row in enumerate(rows):
        qty = to_number(row.get("quantity"))
        unit_price = to_number(row.get("vendor_unit_price") or row.get("unit_price"))
        amount = to_number(row.get("amount"))
        if qty and unit_price and amount and abs(qty * unit_price - amount) > 1:
            issues.append({
                "rowIndex": idx,
                "issueType": "AMOUNT_MISMATCH",
                "severity": "WARNING",
                "fieldKey": "amount",
                "fieldLabel": "금액",
                "message": f"{idx + 1}행 금액이 수량×단가와 다릅니다. 계산값={qty * unit_price:,.0f}, 입력값={amount:,.0f}",
            })

        # 실제 업체별 비교표에서만 핵심 식별값 누락을 경고한다.
        if table_type == "PRICE_COMPARISON" and not row.get("item_name") and not row.get("vendor_name"):
            issues.append({
                "rowIndex": idx,
                "issueType": "MISSING_KEY_FIELD",
                "severity": "WARNING",
                "fieldKey": "item_name",
                "fieldLabel": "품목명",
                "message": f"{idx + 1}행의 품목명 또는 업체명을 확인하세요.",
            })

        # 단위 인식 자체가 낮은 경우만 경고한다.
        # 업체 비교표는 이미 원문 단위를 그대로 보여주는 산출 표이므로 단순 confidence 누락을 오류로 올리지 않는다.
        known_unit = compact_text(row.get("unit")) in {"개", "본", "식", "시간", "hr", "m", "m2", "m3", "㎡", "㎥", "공m3", "공㎥", "kg", "톤", "대", "장", "매", "조", "세트"}
        if table_type not in {MULTI_VENDOR_COMPARE_TABLE_TYPE, "PRICE_COMPARISON"} and row.get("unit") and not known_unit and float(row.get("unit_confidence") or 0) < 0.5:
            issues.append({
                "rowIndex": idx,
                "issueType": "UNIT_PARSE_LOW_CONFIDENCE",
                "severity": "INFO",
                "fieldKey": "unit",
                "fieldLabel": "단위",
                "message": f"{idx + 1}행 단위 '{row.get('unit_original') or row.get('unit')}' 인식값을 확인하세요.",
            })

    if table_type == "PRICE_COMPARISON":
        grouped: Dict[str, set[str]] = {}
        vendor_count = sum(1 for row in rows if str(row.get("vendor_name") or "").strip())
        # 업체명이 2개 이상 있는 진짜 비교표에서만 동일 품목·규격 단위 불일치를 경고한다.
        if vendor_count >= 2:
            for row in rows:
                item_key = str(row.get("item_name") or "").strip()
                spec_key = str(row.get("spec") or "").strip()
                if not item_key:
                    continue
                key = f"{item_key}|{spec_key}"
                grouped.setdefault(key, set()).add(str(row.get("unit_normalized") or row.get("unit") or "").strip())
            for key, units in grouped.items():
                real_units = {u for u in units if u}
                if len(real_units) >= 2:
                    item_name, spec = key.split("|", 1)
                    label = f"{item_name} / {spec}" if spec else item_name
                    issues.append({
                        "rowIndex": None,
                        "issueType": "UNIT_MISMATCH_BETWEEN_VENDORS",
                        "severity": "WARNING",
                        "fieldKey": "unit",
                        "fieldLabel": "단위",
                        "message": f"'{label}' 항목은 업체별 단위가 달라 직접 단가 비교 전에 환산 기준 확인이 필요합니다. 단위={', '.join(sorted(real_units))}",
                    })
    return issues

def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n...[TRUNCATED]"


def _build_llm_intent_prompt(user_request: str, file_summaries: List[Dict[str, Any]], row_samples: List[Dict[str, Any]]) -> str:
    compact_files = json.dumps(file_summaries[:8], ensure_ascii=False, default=str)
    compact_rows = json.dumps(row_samples[:6], ensure_ascii=False, default=str)
    return f"""
너는 문서 분석 시스템의 요청 의도 분석기다.
반드시 JSON 객체 1개만 반환한다. 마크다운/설명문 금지.

[사용자 입력]
{user_request or ''}

[현재 업로드/작업 파일 요약]
{compact_files}

[파서가 원문에서 추출한 표 행 샘플]
{compact_rows}

[반환 JSON 스키마]
{{
  "intent": "DOCUMENT_SUMMARY|STANDARD_TABLE|COMPANY_COMPARISON|TABLE_FILTER|EXCEL_CREATE|DOCUMENT_QA|UNKNOWN",
  "targetKeywords": ["사용자가 직접 입력한 검색/비교 대상 단어"],
  "includePreviousFiles": true,
  "requiresStandardPrice": false,
  "requiresVendorQuotes": false,
  "outputFormat": "analysis|table|excel",
  "reason": "판단 근거 1문장"
}}

[중요 규칙]
1. targetKeywords에는 사용자 입력에 실제로 들어간 단어/구만 넣는다. 원문 표에만 있고 사용자가 말하지 않은 공종명은 넣지 않는다.
2. "이것도 추가해서", "다시", "전에 올린 것"은 includePreviousFiles=true로 해석한다.
3. 금액, 단가, 회사명, 공종명을 새로 만들지 않는다.
4. 사용자가 "회사별", "업체별", "비교", "견적"을 말하면 COMPANY_COMPARISON으로 판단한다.
5. 사용자가 "표준시장단가만", "기준단가만"을 말하면 STANDARD_TABLE로 판단한다.
""".strip()


def infer_request_intent_by_rule(user_request: str, parsed_files: List[Dict[str, Any]], rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Gemini JSON 응답이 실패해도 화면/표 생성이 멈추지 않도록 하는 결정론적 의도 분석."""
    request = str(user_request or "")
    compact = compact_text(request)
    if not compact:
        return {}
    if _request_wants_company_comparison(request, None):
        intent = "COMPANY_COMPARISON"
    elif _request_wants_standard_price(request, None):
        intent = "STANDARD_TABLE"
    elif any(word in compact for word in ["엑셀", "다운로드", "만들어줘", "생성"]):
        intent = "EXCEL_CREATE"
    elif any(word in compact for word in ["표", "정리"]):
        intent = "TABLE_FILTER"
    elif any(word in compact for word in ["문서", "뭐야", "요약", "분석"]):
        intent = "DOCUMENT_SUMMARY"
    else:
        intent = "UNKNOWN"

    terms = []
    try:
        terms = _extract_focus_terms(request, rows or [], [], llm_terms=[])[:20]
    except Exception:
        terms = []
    return {
        "intent": intent,
        "targetKeywords": terms,
        "includePreviousFiles": any(word in compact for word in ["이것도", "추가", "전에", "기존", "같이"]),
        "requiresStandardPrice": _request_wants_standard_price(request, None),
        "requiresVendorQuotes": _request_wants_company_comparison(request, None),
        "outputFormat": "table" if intent in {"COMPANY_COMPARISON", "TABLE_FILTER", "STANDARD_TABLE"} else "analysis",
        "reason": "LLM JSON 실패 또는 미사용 시 규칙 기반으로 보정",
        "_llmIntentUsed": False,
        "_intentSource": "rule_fallback",
    }


async def interpret_request_with_llm(user_request: str, parsed_files: List[Dict[str, Any]], rows: List[Dict[str, Any]]) -> Tuple[Dict[str, Any], str]:
    """LLM은 사용자 요청 의도 분석에 사용하되, 실패하면 규칙 보정으로 계속 진행한다.

    이전 안정화 패치처럼 업체 비교 요청에서 LLM을 완전히 생략하지 않는다.
    단, LLM이 품목/단가를 새로 만들지는 못하게 하고, 숫자 표 추출은 pdfplumber/규칙 파서가 담당한다.
    """
    rule_intent = infer_request_intent_by_rule(user_request, parsed_files, rows)
    cfg = get_llm_config()
    if not cfg.enabled:
        return rule_intent, ""

    file_summaries = []
    for item in parsed_files[:8]:
        text = str(item.get("extractedText") or item.get("extracted_text") or "")
        file_summaries.append({
            "name": item.get("originalName") or item.get("original_name"),
            "pageCount": item.get("pageCount") or item.get("page_count"),
            "charCount": len(text),
            "rowCount": len(item.get("parsedRows") or item.get("rows") or []),
            # 의도분석에서는 전체 본문을 넣지 않는다. 큰 PDF에서 JSON 응답 불안정 방지.
            "textPreview": text[:80],
        })

    # 의도분석은 가볍게: 표 행 샘플 12개만 제공한다.
    prompt = _build_llm_intent_prompt(user_request, file_summaries, rows[:6])
    try:
        result = await call_llm_json(prompt, cfg)
        if not isinstance(result, dict):
            raise ValueError("LLM 의도분석 결과가 JSON 객체가 아닙니다.")
        result.setdefault("intent", rule_intent.get("intent") or "UNKNOWN")
        result.setdefault("targetKeywords", rule_intent.get("targetKeywords") or [])
        result.setdefault("requiresStandardPrice", rule_intent.get("requiresStandardPrice") or False)
        result.setdefault("requiresVendorQuotes", rule_intent.get("requiresVendorQuotes") or False)
        result["_llmIntentUsed"] = True
        result["_intentSource"] = "llm"
        return result, ""
    except Exception as exc:  # noqa: BLE001
        # 오류를 화면의 분석 실패로 올리지 않는다. 의도는 규칙 파서로 보정한다.
        if rule_intent:
            rule_intent["_llmIntentUsed"] = False
            rule_intent["_intentSource"] = "rule_fallback_after_llm_error"
            rule_intent["_llmError"] = str(exc)[:300]
            return rule_intent, ""
        return {}, str(exc)


def should_call_llm(user_request: str, combined_text: str, rows: List[Dict[str, Any]], file_count: int, table_type: str) -> bool:
    """LLM 호출 여부.

    업체별 단가 비교에서도 LLM은 호출한다. 단, 역할은 표/숫자 생성이 아니라
    의도 보조·요약·검증 의견 생성이다. 표 추출/단가 계산은 규칙 파서가 유지한다.
    """
    cfg = get_llm_config()
    if not cfg.enabled:
        return False
    if not combined_text.strip() and not rows:
        return False
    if cfg.use_mode == "off":
        return False
    if cfg.use_mode == "always":
        return True
    request_text = user_request or ""
    if user_request and str(user_request).strip():
        return True
    if table_type in {MULTI_VENDOR_COMPARE_TABLE_TYPE, "PRICE_COMPARISON", "STANDARD_MARKET_PRICE_TABLE"}:
        return True
    if file_count >= 2:
        return True
    if not rows:
        return True
    if any(word in request_text for word in ["단가", "비교", "업체", "회사", "표", "엑셀", "자사", "양식"]):
        return True
    return False


def build_llm_prompt(user_request: str, output_mode: str, template_id: str | None, combined_text: str, rows: List[Dict[str, Any]]) -> str:
    cfg = get_llm_config()
    is_large_table = len(rows or []) > 30 or any(word in str(user_request or "") for word in ["단가", "비교", "업체", "회사", "견적"])
    max_llm_rows = 25 if is_large_table else 80
    max_markdown_rows = 20 if is_large_table else 50
    max_context_chars = min(cfg.context_chars, 7000) if is_large_table else cfg.context_chars
    markdown_table = table_to_markdown(rows, max_rows=max_markdown_rows)
    compact_rows = json.dumps(rows[:max_llm_rows], ensure_ascii=False)
    text_part = _truncate(combined_text, max_context_chars)
    return f"""
너는 건설/전기/설비 업무 문서를 엑셀화하기 위한 문서 구조화 엔진이다.
반드시 JSON 객체 1개만 반환한다. 설명문, 마크다운, 코드블록은 금지한다.

[사용자 요청]
{user_request or '문서를 분석해서 표로 만들어줘'}

[산출 방식]
output_mode={output_mode or 'FREE_FORM'}
template_id={template_id or ''}

[기존 규칙 파서 표 후보]
{markdown_table or '(표 후보 없음)'}

[표 후보 JSON]
{compact_rows}

[PyMuPDF/pdfplumber/PP-Structure/OCR 추출 텍스트]
{text_part}

[반환 JSON 스키마]
{{
  "analysis": {{
    "documentType": "단가 비교 자료|견적서|거래명세서|업무 문서|기타 중 하나 또는 적절한 한글명",
    "purpose": "문서 데이터 엑셀화 목적",
    "summary": "핵심 분석 결과 1~3문장. 추측 금지. 확인 필요 사항 명시.",
    "confidence": 0.0,
    "keyValues": [{{"label": "문서 유형", "value": "..."}}]
  }},
  "table": {{
    "tableName": "문서 표 후보",
    "tableType": "PRICE_COMPARISON 또는 NORMAL_TABLE",
    "columns": [
      {{"key":"vendor_name","label":"업체명"}},
      {{"key":"item_name","label":"품목명"}},
      {{"key":"spec","label":"규격"}},
      {{"key":"quantity","label":"수량"}},
      {{"key":"unit","label":"단위"}},
      {{"key":"unit_price","label":"단가"}},
      {{"key":"amount","label":"금액"}},
      {{"key":"remark","label":"비고"}}
    ],
    "rows": [
      {{
        "vendor_name":"",
        "item_name":"",
        "spec":"",
        "quantity":"",
        "unit":"",
        "unit_original":"",
        "unit_normalized":"",
        "standard_unit_price":"",
        "vendor_unit_price":"",
        "unit_price":"",
        "amount":"",
        "price_diff":"",
        "diff_rate":"",
        "remark":""
      }}
    ]
  }},
  "issues": [
    {{"rowIndex":0,"issueType":"CHECK_REQUIRED","severity":"WARNING","fieldKey":"unit","fieldLabel":"단위","message":"확인 필요 사유"}}
  ]
}}

[중요 규칙]
1. 원문에 없는 업체명/품목명/규격/수량/단위/단가/금액을 절대 만들어내지 말 것.
2. 문서가 유스케이스 명세서, 보고서, 요구사항서, 설명 문서이면 견적서/단가표로 분류하지 말고 table.rows는 빈 배열로 둔다.
3. 기준서/지침서이면 업체명·품목명·금액을 만들지 말고 section, basis_item, application_basis, calculation_method, unit_price_basis, source_page, remark 구조의 REFERENCE_GUIDELINE_TABLE로 정리한다.
4. 본문에 예시로 "견적서", "단가표"라는 단어가 있어도 실제 제목/표 구조가 아니면 견적서로 판단하지 않는다.
5. 실제 표 행이 없으면 rows를 만들지 말고 issues에 TABLE_NOT_FOUND 또는 NO_BUSINESS_TABLE을 추가한다.
5. 단위는 원문 단위를 unit_original에 보존하고, EA/개/PCS는 개, M/m2/m3는 m/㎡/㎥, 공m3는 공㎥, 본/개소/hr는 본/개소/시간으로 정규화한다.
6. 본/개소/공㎥/㎡/㎥는 건설 표준 단위이므로 행별 환산 경고를 만들지 않는다. 업체별 동일 품목·동일 규격인데 단위가 서로 다를 때만 UNIT_MISMATCH_BETWEEN_VENDORS를 추가한다.
7. 표준단가/기준단가와 업체 견적단가를 구분한다. 업체 비교에는 vendor_unit_price를 사용하고 standard_unit_price를 업체 단가로 사용하지 않는다.
8. 수량×단가와 금액이 다르면 AMOUNT_MISMATCH issue를 추가한다.
9. 업체별 단위가 다른 단가 비교는 최저가를 확정하지 말고 확인 필요로 둔다.
9. 행/열이 애매하면 추측하지 말고 rows에 넣지 않는다.
""".strip()



def build_llm_grounded_analysis_prompt(
    user_request: str,
    analysis: Dict[str, Any],
    table: Dict[str, Any],
    issues: List[Dict[str, Any]],
    combined_text: str,
) -> str:
    """이미 파서가 만든 표를 LLM이 해석/요약만 하도록 하는 작은 프롬프트.

    큰 PDF 본문 전체를 넣지 않고 표 결과 일부만 넣어 JSON 파싱 실패를 줄인다.
    LLM은 rows/columns를 다시 만들 수 없고, summary/keyValues/issues만 반환한다.
    """
    cfg = get_llm_config()
    rows = table.get("rows") if isinstance(table, dict) and isinstance(table.get("rows"), list) else []
    columns = table.get("columns") if isinstance(table, dict) and isinstance(table.get("columns"), list) else []
    table_type = table.get("tableType") or table.get("table_type") or "NORMAL_TABLE"
    compact_table = {
        "tableName": table.get("tableName") or table.get("table_name") or "문서 표 후보",
        "tableType": table_type,
        "rowCount": len(rows),
        "columns": columns[:30],
        "rowsSample": rows[:30],
    }
    compact_issues = issues[:30]
    text_markers = _truncate(combined_text, min(cfg.context_chars, 2500))
    return f"""
너는 건설 문서 분석 결과를 검토하는 LLM 보조분석기다.
반드시 JSON 객체 1개만 반환한다. 마크다운/설명문/코드블록 금지.

[사용자 요청]
{user_request or ''}

[이미 확정된 파서 결과]
{json.dumps(compact_table, ensure_ascii=False, default=str)}

[기존 확인 필요]
{json.dumps(compact_issues, ensure_ascii=False, default=str)}

[원문 근거 일부]
{text_markers}

[반환 JSON 스키마]
{{
  "analysis": {{
    "documentType": "업체별 단가 비교 자료|표준시장단가 자료|견적서|업무 문서|기타",
    "purpose": "문서 데이터 엑셀화 목적",
    "summary": "파서 결과에 근거한 핵심 분석 3~5문장. 비교 대상, 최저 업체 경향, 수량/금액 반영 여부를 포함",
    "confidence": 0.0,
    "keyValues": [
      {{"label":"LLM 검토", "value":"요청 의도와 파서 결과 기준으로 확인한 핵심 내용"}},
      {{"label":"비교 해석", "value":"어떤 업체가 어떤 품목에서 유리한지 표 데이터 근거로 요약"}},
      {{"label":"비용 포인트", "value":"금액 합계/최저가 경향/차이가 큰 구간을 보수적으로 설명"}},
      {{"label":"확인 추천", "value":"단위, 규격, 수량, 제외/추가 항목 중 사용자가 확인해야 할 점"}}
    ]
  }},
  "issues": [
    {{"rowIndex":null,"issueType":"CHECK_REQUIRED","severity":"INFO","fieldKey":"table","fieldLabel":"표 데이터","message":"필요한 경우만 작성"}}
  ]
}}

[엄격 규칙]
1. rows, columns, 단가, 금액, 수량을 새로 만들거나 수정하지 않는다.
2. 표에 없는 업체명/공종명/단가를 말하지 않는다.
3. 숫자 계산은 이미 Python 파서가 수행한 것으로 보고, 검토 의견만 쓴다.
4. 문제가 없으면 issues는 빈 배열 []로 둔다.
5. summary에는 "LLM이 직접 표를 생성했다"고 쓰지 않는다. "파서 결과를 검토했다"고 표현한다.
6. keyValues는 3개 이상 작성하되, rowsSample에 있는 값만 근거로 삼는다.
7. 표에서 제외되거나 추가된 항목은 table rows 기준으로만 말한다. 없으면 없다고 말한다.
""".strip()


def normalize_llm_analysis_only(llm_result: Dict[str, Any], base_analysis: Dict[str, Any], base_issues: List[Dict[str, Any]]) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """LLM 보조분석 결과에서 analysis/issue만 안전하게 병합한다."""
    merged = dict(base_analysis or {})
    raw_analysis = llm_result.get("analysis") if isinstance(llm_result.get("analysis"), dict) else {}
    for key in ["documentType", "purpose", "summary", "confidence"]:
        value = raw_analysis.get(key) or raw_analysis.get(key[0].lower() + key[1:])
        if value not in (None, ""):
            if key == "confidence":
                try:
                    merged[key] = min(float(value), 0.9)
                except Exception:
                    continue
            else:
                merged[key] = str(value).strip()

    raw_kvs = raw_analysis.get("keyValues") if isinstance(raw_analysis.get("keyValues"), list) else []
    grounded_llm_notes = []
    for kv in raw_kvs[:8]:
        if not isinstance(kv, dict):
            continue
        label = str(kv.get("label") or "").strip()
        value = str(kv.get("value") or "").strip()
        if label and value:
            grounded_llm_notes.append({"label": label, "value": value})
    if grounded_llm_notes:
        merged.setdefault("processingMeta", {})["llmReviewNotes"] = grounded_llm_notes

    merged.setdefault("llmMeta", llm_result.get("_llm", {}))

    merged_issues = list(base_issues or [])
    raw_issues = llm_result.get("issues") if isinstance(llm_result.get("issues"), list) else []
    for issue in raw_issues[:30]:
        if not isinstance(issue, dict):
            continue
        message = str(issue.get("message") or "").strip()
        if not message:
            continue
        merged_issues.append({
            "rowIndex": issue.get("rowIndex") if issue.get("rowIndex") is not None else None,
            "issueType": issue.get("issueType") or "CHECK_REQUIRED",
            "severity": issue.get("severity") or "INFO",
            "fieldKey": issue.get("fieldKey") or issue.get("field") or None,
            "fieldLabel": issue.get("fieldLabel") or None,
            "message": message,
            "suggestedValue": issue.get("suggestedValue") or None,
        })
    return merged, merged_issues
def normalize_llm_result(llm_result: Dict[str, Any], fallback_rows: List[Dict[str, Any]], fallback_table_type: str, source_text: str, user_request: str = "") -> Tuple[Dict[str, Any], Dict[str, Any], List[Dict[str, Any]]]:
    analysis = llm_result.get("analysis") if isinstance(llm_result.get("analysis"), dict) else {}
    table = llm_result.get("table") if isinstance(llm_result.get("table"), dict) else {}
    if not table and isinstance(llm_result.get("tables"), list) and llm_result["tables"]:
        table = llm_result["tables"][0]

    rows = table.get("rows") if isinstance(table.get("rows"), list) else fallback_rows
    normalized_rows = []
    for row in rows[:250]:
        if not isinstance(row, dict):
            continue
        clean_row = {key: (clean_number(value) if key in NUMBER_KEYS else str(value or "").strip()) for key, value in row.items()}
        enriched = enrich_row_units(clean_row)
        if is_business_row_supported(enriched, source_text):
            normalized_rows.append(enriched)

    # LLM이 원문에 없는 예시 행을 생성한 경우 제거하고, 근거 있는 규칙 파서 행만 fallback으로 사용한다.
    if not normalized_rows:
        normalized_rows = filter_grounded_rows(fallback_rows, source_text)

    table_type = table.get("tableType") or table.get("table_type") or fallback_table_type
    if table_type not in {"PRICE_COMPARISON", "NORMAL_TABLE", "REFERENCE_GUIDELINE_TABLE", "GUIDELINE_SUMMARY_TABLE", "STANDARD_MARKET_PRICE_TABLE", TEXT_VENDOR_COMPARISON_TABLE_TYPE, MULTI_VENDOR_COMPARE_TABLE_TYPE}:
        table_type = fallback_table_type
    if is_reference_or_guideline_document(source_text) and (normalized_rows or fallback_table_type in REFERENCE_TABLE_TYPES):
        table_type = "REFERENCE_GUIDELINE_TABLE"
    if not normalized_rows:
        table_type = "REFERENCE_GUIDELINE_TABLE" if fallback_table_type in REFERENCE_TABLE_TYPES else "NORMAL_TABLE"

    default_cols_for_type = REFERENCE_GUIDELINE_COLUMNS if table_type in REFERENCE_TABLE_TYPES else (STANDARD_MARKET_PRICE_COLUMNS if table_type in STANDARD_MARKET_TABLE_TYPES else (TEXT_VENDOR_COMPARISON_COLUMNS if table_type == TEXT_VENDOR_COMPARISON_TABLE_TYPE else DEFAULT_COLUMNS))
    normalized_table = {
        "tableName": table.get("tableName") or table.get("table_name") or ("기준서 항목 표" if table_type in REFERENCE_TABLE_TYPES else "문서 표 후보"),
        "tableType": table_type,
        "columns": table.get("columns") if isinstance(table.get("columns"), list) and table.get("columns") else default_cols_for_type,
        "rows": normalized_rows,
    }

    profile = infer_document_profile(source_text, user_request)
    inferred_doc_type = profile["documentType"]
    llm_doc_type = str(analysis.get("documentType") or analysis.get("document_type") or "").strip()
    # narrative 문서인데 LLM이 견적서/단가표로 과잉 분류한 경우 원문 기반 판별을 우선한다.
    if is_narrative_document(source_text) or (llm_doc_type in {"견적서", "단가 비교 자료", "견적서/단가표"} and not normalized_rows):
        doc_type = inferred_doc_type
    else:
        doc_type = llm_doc_type or inferred_doc_type or ("단가 비교 자료" if table_type == "PRICE_COMPARISON" else "업무 문서")

    key_values = analysis.get("keyValues") if isinstance(analysis.get("keyValues"), list) else []
    grounded_key_values = []
    for kv in key_values[:20]:
        if not isinstance(kv, dict):
            continue
        label = str(kv.get("label") or "").strip()
        value = str(kv.get("value") or "").strip()
        if label and value and source_has_value(source_text, value):
            grounded_key_values.append({"label": label, "value": value})
    # 원문에서 안정적으로 뽑을 수 있는 메타값을 보강한다.
    for kv in extract_key_values_from_text(source_text):
        if kv not in grounded_key_values:
            grounded_key_values.append(kv)

    summary = str(analysis.get("summary") or "").strip()
    if not summary or ("견적" in summary and not normalized_rows and is_narrative_document(source_text)):
        if normalized_rows:
            summary = _build_document_only_summary(source_text, table_type, len(normalized_rows), doc_type)
        else:
            summary = _build_document_only_summary(source_text, table_type, 0, doc_type)

    normalized_analysis = {
        "documentType": doc_type,
        "purpose": analysis.get("purpose") or profile.get("purpose") or "문서 데이터 엑셀화",
        "summary": summary,
        "confidence": min(float(analysis.get("confidence") or profile.get("confidence") or 0.7), 0.88),
        "keyValues": grounded_key_values[:20],
        "llmMeta": llm_result.get("_llm", {}),
    }

    issues = llm_result.get("issues") if isinstance(llm_result.get("issues"), list) else []
    normalized_issues = []
    for issue in issues[:100]:
        if not isinstance(issue, dict):
            continue
        normalized_issues.append({
            "rowIndex": issue.get("rowIndex") if issue.get("rowIndex") is not None else None,
            "issueType": issue.get("issueType") or "CHECK_REQUIRED",
            "severity": issue.get("severity") or "WARNING",
            "fieldKey": issue.get("fieldKey") or issue.get("field") or None,
            "fieldLabel": issue.get("fieldLabel") or None,
            "message": issue.get("message") or "LLM 분석 결과 확인이 필요합니다.",
            "suggestedValue": issue.get("suggestedValue") or None,
        })

    return normalized_analysis, normalized_table, normalized_issues


def _first_text(*values: Any) -> str:
    for value in values:
        if value is None:
            continue
        if isinstance(value, list):
            parts = []
            for item in value:
                if isinstance(item, dict):
                    parts.append(str(item.get("value") or item.get("content") or item.get("text") or ""))
                else:
                    parts.append(str(item or ""))
            text = "\n".join(part.strip() for part in parts if part and str(part).strip()).strip()
            if text:
                return text
            continue
        if isinstance(value, dict):
            text = str(value.get("value") or value.get("content") or value.get("text") or value.get("summary") or "").strip()
            if text:
                return text
            continue
        text = str(value or "").strip()
        if text:
            return text
    return ""


def _clean_business_sentence(value: Any) -> str:
    text = clean_cell_text(value)
    # PDF 페이지 마커 제거: [page N / T] 또는 [page N / T OCR]
    text = re.sub(r"\[page\s*\d+\s*/\s*\d+(?:\s*OCR)?\]", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"^[•\-–—*\s]+", "", text).strip()
    return text




def _looks_like_system_or_processing_meta(value: Any) -> bool:
    text = _clean_business_sentence(value)
    if not text:
        return False
    # PDF 페이지 마커 자체가 남아있으면 메타로 처리
    if re.search(r"\[page\s*\d+\s*/\s*\d+", text, re.I):
        return True
    meta_tokens = [
        "첨부 파일", "파일 수", "총 페이지", "확인된 총 페이지", "페이지 수",
        "표 후보", "파싱", "PyMuPDF", "pdfplumber", "PP-Structure", "PaddleOCR",
        "OCR", "LLM", "ai-server", "저장 위치", "산출 방식", "요청 내용은",
        "문서 분석", "엑셀 미리보기", "백그라운드", "분석 결과", "처리했습니다",
        "텍스트 전용 샘플", "텍스트전용", "샘플 문서",
    ]
    return any(token.lower() in text.lower() for token in meta_tokens)


def _document_only_sentences(items: List[str], limit: int = 8) -> List[str]:
    out: List[str] = []
    seen = set()
    for item in items or []:
        clean = _clean_business_sentence(item)
        if not clean:
            continue
        if _looks_like_user_format_request(clean) or _looks_like_system_or_processing_meta(clean):
            continue
        if re.search(r"(보고서|검토보고|검토본)$", clean) or "/ 검토보고" in clean:
            continue
        key = compact_text(clean)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(clean)
        if len(out) >= limit:
            break
    return out

def _looks_like_user_format_request(value: Any) -> bool:
    text = _clean_business_sentence(value)
    if not text:
        return False
    prompt_tokens = [
        "정리해줘", "작성해줘", "만들어줘", "써줘", "출력해줘", "보여줘",
        "보고서 형식", "보고서 형태", "보고서로", "업무보고서 형식",
        "회사 업무보고서 형식", "회사 보고서 형식", "핵심 내용만", "요약과 검토내용",
        "상세하게", "풍부하게", "보고 목적·", "검토 결과·", "조치 계획 중심",
        "원문에 없는 내용", "임의로 만들지", "확인 필요로 표시",
    ]
    return any(token in text for token in prompt_tokens)


def _safe_document_purpose(value: Any) -> str:
    text = _clean_business_sentence(value)
    if not text or _looks_like_user_format_request(text):
        return ""
    generic_values = {
        compact_text("문서 데이터 엑셀화 목적"),
        compact_text("문서 내용 요약 및 표 데이터 추출 가능 여부 확인"),
        compact_text("문서 데이터 엑셀화"),
    }
    if compact_text(text) in generic_values:
        return ""
    return text


def _infer_report_purpose_from_document(analysis: Dict[str, Any], combined_text: str, topics: List[str]) -> str:
    source = " ".join([
        _clean_business_sentence(analysis.get("documentType") or analysis.get("document_type") or ""),
        _clean_business_sentence(analysis.get("summary") or ""),
        _clean_business_sentence(combined_text[:3000]),
        " ".join(topics or []),
    ])
    if any(word in source for word in ["점검", "안전", "위험", "현장", "감리", "지적사항"]):
        return "첨부 문서의 현장 점검 내용과 확인 필요 사항을 검토하기 위한 보고입니다."
    if any(word in source for word in ["견적", "단가", "금액", "업체", "비교", "견적서"]):
        return "첨부 문서의 견적·단가·업체 비교 내용을 검토하기 위한 보고입니다."
    if any(word in source for word in ["회의", "안건", "결정", "조치사항", "협의"]):
        return "첨부 문서의 논의 내용과 후속 조치 사항을 정리하기 위한 보고입니다."
    if any(word in source for word in ["작업일보", "작업일지", "공정", "물량"]):
        return "첨부 문서의 작업 현황과 후속 관리 사항을 정리하기 위한 보고입니다."
    return "첨부 문서의 주요 내용과 확인 필요 사항을 업무 보고 형식으로 정리하기 위한 보고입니다."


def _split_business_sentences(text: str, limit: int = 8) -> List[str]:
    normalized = str(text or "").replace("\r", "\n")
    chunks: List[str] = []
    for line in normalized.splitlines():
        line = _clean_business_sentence(line)
        if not line:
            continue
        # 너무 긴 줄은 문장 단위로 한 번 더 나눈다.
        pieces = re.split(r"(?<=[.。!?])\s+|(?<=다\.)\s+|(?<=요\.)\s+", line)
        for piece in pieces:
            item = _clean_business_sentence(piece)
            if 8 <= len(item) <= 220:
                chunks.append(item)
            elif len(item) > 220:
                chunks.append(item[:220].rstrip() + "…")
        if len(chunks) >= limit:
            break
    # 중복 제거
    out: List[str] = []
    seen = set()
    for chunk in chunks:
        key = compact_text(chunk)
        if key and key not in seen:
            seen.add(key)
            out.append(chunk)
        if len(out) >= limit:
            break
    return out


def _bullet_text(items: List[str], empty: str = "") -> str:
    clean_items = [_clean_business_sentence(item) for item in items if _clean_business_sentence(item)]
    if not clean_items:
        return empty
    return "\n".join(f"• {item}" for item in clean_items)


def _extract_document_title_from_text(text: str, user_request: str = "") -> str:
    request = _clean_business_sentence(user_request)
    # 사용자가 특정 제목을 짧게 지정한 경우만 제목 근거로 사용한다.
    # "회사 업무보고서 형식으로 정리해줘" 같은 작성 지시는 제목/보고목적에 넣지 않는다.
    if request and not _looks_like_user_format_request(request):
        m = re.search(r"([가-힣A-Za-z0-9·ㆍ\s/()\-]{4,60})(보고서|회의록|공문|검토서)", request)
        if m:
            return _clean_business_sentence(m.group(0))[:70]
    for line in str(text or "").splitlines()[:40]:
        line = _clean_business_sentence(line)
        if 4 <= len(line) <= 70 and any(token in line for token in ["보고서", "회의록", "검토", "현황", "공문", "작업일보", "지시사항"]):
            return line[:70]
    if request and not _looks_like_user_format_request(request):
        return request[:60]
    return "업무 문서 검토 보고서"


def _extract_key_topics(text: str, user_request: str = "", limit: int = 6) -> List[str]:
    # 보고서 본문 후보는 반드시 원문 텍스트에서만 뽑는다.
    # 사용자 지시문(예: "보고서 형식으로 작성")은 제목/본문 근거로 사용하지 않는다.
    source = str(text or "")[:6000]
    topics: List[str] = []
    patterns = [
        r"([가-힣A-Za-z0-9·ㆍ/()\-\s]{2,40})(?:\s*[:：]\s*)([가-힣A-Za-z0-9·ㆍ/()\-\s]{2,80})",
        r"(설계변경|협력업체|작업일보|물량|일정|감리\s*지적사항|발주처\s*회의\s*지시사항|임시전력\s*안전점검|안전점검|EPS\s*변경|위험요인)",
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, source, re.I):
            value = _clean_business_sentence(" ".join(g for g in m.groups() if g))
            if value and len(value) >= 2:
                topics.append(value)
    # 일반 명사성 키워드 보강. 단, 단일어만 있을 때는 보고서 본문에서 '확인항목'으로만 사용한다.
    for line in _split_business_sentences(source, limit=30):
        if any(word in line for word in ["설계변경", "협력업체", "작업일보", "물량", "일정", "감리", "발주처", "회의", "지시", "안전", "점검", "위험", "EPS"]):
            topics.append(line)
    out: List[str] = []
    seen = set()
    for topic in topics:
        key = compact_text(topic)
        if key and key not in seen:
            seen.add(key)
            out.append(topic)
        if len(out) >= limit:
            break
    return out


def _extract_action_candidates(text: str, issues: List[Dict[str, Any]], user_request: str = "", limit: int = 6) -> List[Dict[str, str]]:
    # 조치계획 후보도 원문에 있는 조치/확인 문장만 사용한다.
    source_lines = _split_business_sentences(str(text or "")[:9000], limit=80)
    action_words = ["조치", "처리", "미종결", "지시", "점검", "보완", "협의", "변경", "수립", "제출", "요청", "재확인"]
    stop_exact = {"물량", "일정", "품질", "안전", "회의", "작업일보", "협력업체"}
    candidates: List[str] = []
    for line in source_lines:
        compact = compact_text(line)
        if compact in {compact_text(x) for x in stop_exact}:
            continue
        if len(line) < 8:
            continue
        if re.search(r"(보고서|검토보고|검토본)$", line) or "/ 검토보고" in line:
            continue
        if "조치 현황" in line and not any(word in line for word in ["미종결", "재확인", "보완", "처리", "지시", "계획"]):
            continue
        if any(word in line for word in action_words):
            candidates.append(line)
    for issue in issues or []:
        message = _clean_business_sentence(issue.get("message") or issue.get("fieldLabel") or "")
        if message and len(message) >= 8:
            candidates.append(message)
    if not candidates:
        topics = _extract_key_topics(text, user_request=user_request, limit=4)
        for topic in topics:
            if any(word in topic for word in ["미종결", "지시", "점검", "변경", "위험", "감리", "발주처"]):
                candidates.append(f"{topic}에 대한 담당자 지정 및 처리 현황 확인")
    out: List[Dict[str, str]] = []
    seen = set()
    for candidate in candidates:
        candidate = _clean_business_sentence(candidate)
        key = compact_text(candidate)
        if not key or key in seen:
            continue
        seen.add(key)
        # 문장 자체가 너무 제목형이면 업무 지시문으로 바꾼다. 원문 항목명은 보존한다.
        if not any(word in candidate for word in ["확인", "조치", "처리", "점검", "보완", "제출", "협의", "수립", "재확인"]):
            action_text = f"{candidate} 관련 처리 필요 여부 확인"
        else:
            action_text = candidate
        out.append({
            "action_item": action_text[:180],
            "owner": "확인 필요",
            "due_date": "미정",
            "status": "확인 필요",
            "remark": "원문 근거 확인 후 담당자/기한 확정",
        })
        if len(out) >= limit:
            break
    return out


def _make_business_drafts(
    user_request: str,
    analysis: Dict[str, Any],
    table: Dict[str, Any],
    issues: List[Dict[str, Any]],
    combined_text: str,
    file_profiles: List[Dict[str, Any]],
) -> Dict[str, Any]:
    rows = table.get("rows") if isinstance(table, dict) and isinstance(table.get("rows"), list) else []
    title = _extract_document_title_from_text(combined_text, "")
    topics = _document_only_sentences(_extract_key_topics(combined_text, user_request="", limit=8), limit=6)
    sentences = _document_only_sentences(_split_business_sentences(combined_text, limit=10), limit=6)
    actions = _extract_action_candidates(combined_text, issues, user_request="", limit=6)

    purpose = _first_text(
        _safe_document_purpose(analysis.get("purpose")),
        _safe_document_purpose(analysis.get("documentPurpose") or analysis.get("document_purpose")),
        _infer_report_purpose_from_document(analysis, combined_text, topics),
    )

    review_items = _document_only_sentences(sentences[:7], limit=7)
    if not review_items:
        review_items = _document_only_sentences(topics[:7], limit=7)
    if not review_items:
        summary_candidate = _clean_business_sentence(analysis.get("summary") or "")
        review_items = _document_only_sentences([summary_candidate], limit=3)

    issue_items = []
    if actions:
        issue_items.extend([item["action_item"] for item in actions[:4]])
    if issues:
        issue_items.extend([_clean_business_sentence(item.get("message") or item.get("fieldLabel") or "") for item in issues[:4]])
    issue_items = _document_only_sentences(issue_items, limit=5)

    if actions:
        action_plan = _bullet_text([item["action_item"] for item in actions])
    else:
        action_plan = ""

    report = {
        "report_title": title if "보고" in title else f"{title} 보고서",
        "report_purpose": purpose,
        "summary": _bullet_text(review_items[:7]),
        "issue_summary": _bullet_text(issue_items[:5]),
        "review_opinion": _bullet_text(issue_items[:5]),
        "action_plan": action_plan,
        "footer_note": "",
    }

    meeting = {
        "meeting_title": title if "회의" in title else f"{title} 검토 회의록",
        "meeting_date": "",
        "meeting_place": "",
        "attendees": "",
        "agenda": _bullet_text([purpose, *topics[:4]], empty=purpose),
        "discussion": _bullet_text(review_items[:5]),
        "decision": "원문에 명시된 최종 결정사항은 확인되지 않았습니다. 회의 확정 후 결정사항을 입력하세요.",
        "remark": "문서 분석 결과 기준 회의록 초안입니다. 참석자, 장소, 최종 결정사항은 회의 후 확정 입력하세요.",
        "action_items": actions,
    }

    body_lines = [
        "1. 귀 부서의 업무 협조에 감사드립니다.",
        f"2. {purpose}",
        "3. 주요 검토 내용은 아래와 같습니다.",
        *[f"   - {item}" for item in review_items[:4]],
    ]
    if actions:
        body_lines.extend(["4. 아래 사항에 대한 확인 및 조치를 요청드립니다.", *[f"   - {item['action_item']}" for item in actions[:4]]])
    else:
        body_lines.append("4. 원문에 명시된 별도 조치사항은 확인되지 않았으나, 필요 시 담당자 지정 후 후속 관리 바랍니다.")
    official = {
        "letter_title": "공 문",
        "document_no": "",
        "recipient": "수신처 확인 필요",
        "reference": "",
        "document_title": title,
        "body": "\n".join(body_lines),
        "attachment_note": "첨부 문서 참조",
        "sender": "공사팀",
    }

    return {
        "report": report,
        "meeting": meeting,
        "officialLetter": official,
    }


def _build_document_only_summary(combined_text: str, table_type: str, row_count: int, document_type: str) -> str:
    sentences = _document_only_sentences(_split_business_sentences(combined_text, limit=8), limit=3)
    if sentences:
        return " ".join(sentences)[:700]
    if table_type == MULTI_VENDOR_COMPARE_TABLE_TYPE:
        return "업체별 견적 단가 비교 내용이 확인됩니다."
    if table_type == TEXT_VENDOR_COMPARISON_TABLE_TYPE:
        return "서술형 업체별 단가 비교 검토 내용이 확인됩니다."
    if table_type in STANDARD_MARKET_TABLE_TYPES:
        return "표준시장단가 관련 공종·규격·단가 내용이 확인됩니다."
    if table_type in REFERENCE_TABLE_TYPES:
        return "기준서 또는 지침서의 검토 항목과 확인 사항이 확인됩니다."
    return f"{document_type or '업무 문서'}의 주요 내용이 확인됩니다."


def _build_source_key_values(combined_text: str, file_profiles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    values = extract_key_values_from_text(combined_text)
    # 파일 수, 페이지 수, OCR/LLM 등 처리 메타는 보고서/본문 후보에 넣지 않는다.
    return values[:20]


async def analyze_uploads(files: List[UploadFile], user_request: str, output_mode: str, template_id: str | None) -> Dict[str, Any]:
    parsed_files = []
    all_rows: List[Dict[str, Any]] = []
    combined_text_parts = []
    all_parse_logs: List[Dict[str, Any]] = []
    total_page_count = 0
    total_text_chars = 0

    for file in files:
        original_filename = repair_mojibake_filename(file.filename)
        file_logs, file_log = _new_parse_logger(f"ANALYZE:{original_filename}")
        file_log("info", "File received", filename=original_filename, content_type=file.content_type or "")
        saved = await save_upload_file(file, "documents")
        target_path = validate_storage_path(saved["filePath"])
        content = target_path.read_bytes()
        suffix = target_path.suffix.lower()
        text = ""
        rows: List[Dict[str, Any]] = []
        page_count = None
        pages_meta: List[Dict[str, Any]] = []
        parse_metrics: Dict[str, Any] = {"ocrUsed": False}

        if suffix in {".xlsx", ".xlsm"}:
            file_log("info", "Spreadsheet parse start", filename=saved["originalName"], engine="openpyxl")
            text, rows = read_xlsx(content)
            parse_metrics = {"engine": "openpyxl", "rowCount": len(rows), "charCount": len(text), "ocrUsed": False}
            file_log("info", "Spreadsheet parse finish", filename=saved["originalName"], rows=len(rows), chars=len(text))
        elif suffix == ".pdf":
            text, rows, page_count, pages_meta, parse_metrics = read_pdf(content, filename=saved["originalName"], log=file_log)
        elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
            file_log("info", "Image OCR parse start", filename=saved["originalName"], engine="PP-Structure/PaddleOCR")
            text, rows, image_metrics = read_image_with_ocr(content, filename=saved["originalName"], log=file_log)
            page_count = 1
            pages_meta = [{
                "page": 1,
                "pageCount": 1,
                "engine": image_metrics.get("engine") or "PP-Structure/PaddleOCR",
                "status": "OCR_EXTRACTED" if text.strip() or rows else "OCR_EMPTY",
                "charCount": len(text),
                "rowCount": len(rows),
            }]
            parse_metrics = image_metrics
        elif suffix == ".docx":
            file_log("info", "DOCX parse start", filename=saved["originalName"], engine="python-docx")
            text, rows = read_docx(content)
            parse_metrics = {"engine": "python-docx", "rowCount": len(rows), "charCount": len(text), "ocrUsed": False}
            file_log("info", "DOCX parse finish", filename=saved["originalName"], rows=len(rows), chars=len(text))
        elif suffix in {".txt", ".csv", ".tsv", ".md", ".json"}:
            file_log("info", "Plain text parse start", filename=saved["originalName"], suffix=suffix)
            text = decode_text(content)
            parse_metrics = {"engine": "text-decode", "charCount": len(text), "ocrUsed": False}
            file_log("info", "Plain text parse finish", filename=saved["originalName"], chars=len(text))
        else:
            file_log("warning", "Unknown extension parsed as text", filename=saved["originalName"], suffix=suffix)
            text = decode_text(content)
            parse_metrics = {"engine": "text-decode", "charCount": len(text), "ocrUsed": False}

        if not rows:
            rows = infer_rows_from_text(text, saved["originalName"] or "file")
        rows = [enrich_row_units(row) for row in rows]
        file_log("info", "File parse summary", filename=saved["originalName"], pages=page_count or 0, chars=len(text), table_rows=len(rows), ocr_used=bool(parse_metrics.get("ocrUsed")))
        all_rows.extend(rows)
        combined_text_parts.append(text)
        all_parse_logs.extend(file_logs)
        total_page_count += int(page_count or 0)
        total_text_chars += len(text or "")
        parsed_files.append({
            **saved,
            "pageCount": page_count,
            "page_count": page_count,
            "extractedText": text,
            "extracted_text": text,
            "pages": pages_meta,
            "parseLogs": file_logs,
            "parse_logs": file_logs,
            "parseMetrics": parse_metrics,
            "parse_metrics": parse_metrics,
            "parsedRows": rows,
            "rows": rows,
        })

    combined_text = "\n\n".join(combined_text_parts)
    llm_intent, llm_intent_error = await interpret_request_with_llm(user_request, parsed_files, all_rows)
    intent_keywords = llm_intent.get("targetKeywords") or llm_intent.get("target_keywords") or [] if isinstance(llm_intent, dict) else []
    intent_name = str(llm_intent.get("intent") or "").strip() if isinstance(llm_intent, dict) else ""
    profile = infer_document_profile(combined_text, user_request)
    file_profiles = build_file_profiles(parsed_files, user_request=user_request, llm_intent=llm_intent)
    wants_price_compare = any(word in (user_request or "") for word in ["단가", "비교", "가격", "견적", "업체", "회사", "최저"]) or intent_name == "COMPANY_COMPARISON"
    text_only_compare_report = is_text_only_vendor_comparison_report(combined_text)
    multi_compare_table = None
    if not text_only_compare_report:
        multi_compare_table = build_multi_vendor_price_comparison(parsed_files, user_request=user_request, llm_intent=llm_intent)
        if not multi_compare_table:
            multi_compare_table = build_single_file_multi_vendor_price_comparison(parsed_files, user_request=user_request, llm_intent=llm_intent)
    narrative_compare_table = None
    if text_only_compare_report:
        # 표가 없는 비교보고서에서 사용자가 특정 품목/업체/표 생성을 요청하면
        # 총괄 금액 행이 아니라 공종별 문장 기반 비교 행을 먼저 만든다.
        # 원문에 없는 특정 업체 단가는 만들지 않고 확인 필요로 남긴다.
        narrative_compare_table = build_text_vendor_comparison_item_table(combined_text, user_request=user_request)
        if not narrative_compare_table:
            narrative_compare_table = build_text_vendor_comparison_summary_table(combined_text, user_request=user_request)
    is_standard_market_doc = is_standard_market_price_document(combined_text) and not _request_wants_company_comparison(user_request, llm_intent) and not text_only_compare_report

    if multi_compare_table:
        all_rows = multi_compare_table.get("rows", [])
        table_type = MULTI_VENDOR_COMPARE_TABLE_TYPE
        document_type = "업체별 단가 비교 자료"
        profile = {
            "documentType": document_type,
            "purpose": "업체별 견적 단가를 공종별로 비교",
            "confidence": 0.88,
        }
    elif narrative_compare_table:
        all_rows = narrative_compare_table.get("rows", [])
        table_type = TEXT_VENDOR_COMPARISON_TABLE_TYPE
        document_type = "업체별 단가 비교 검토보고서"
        profile = {
            "documentType": document_type,
            "purpose": "서술형 업체별 단가 비교 결과와 확인 필요 사항 검토",
            "confidence": 0.9,
        }
    else:
        if is_standard_market_doc:
            text_market_rows = extract_standard_market_rows_from_text(combined_text)
            all_rows = [
                row for row in all_rows
                if str(row.get("item_name", "")).strip()
                and str(row.get("unit", "")).strip()
                and str(row.get("unit_price", "")).strip()
            ]
            all_rows = merge_standard_market_rows(all_rows, text_market_rows)
            for row in all_rows:
                row.pop("remark", None)
            all_rows = filter_grounded_rows(all_rows, combined_text)
        elif is_reference_or_guideline_document(combined_text):
            reference_rows = extract_reference_guideline_rows(combined_text, user_request=user_request)
            if reference_rows:
                all_rows = reference_rows
        all_rows = filter_grounded_rows(all_rows, combined_text)
        is_reference_doc = is_reference_or_guideline_document(combined_text) and not is_standard_market_doc
        has_price_rows = bool(all_rows) and profile.get("documentType") in {"견적서/단가표"} and not is_reference_doc
        if is_standard_market_doc:
            table_type = "STANDARD_MARKET_PRICE_TABLE"
        elif is_reference_doc:
            table_type = "REFERENCE_GUIDELINE_TABLE"
        else:
            table_type = "PRICE_COMPARISON" if (wants_price_compare and bool(all_rows)) or has_price_rows else "NORMAL_TABLE"
        document_type = profile.get("documentType") or ("단가 비교 자료" if table_type == "PRICE_COMPARISON" else "업무 문서")

    model_name = "pymupdf-pdfplumber-ppstructure-rule-parser"
    prompt_version = "pymupdf-pdfplumber-ppstructure-v1"
    llm_used = False
    llm_error = ""
    analysis = {
        "documentType": document_type,
        "purpose": profile.get("purpose") or "문서 데이터 엑셀화",
        "summary": _build_document_only_summary(combined_text, table_type, len(all_rows), document_type),
        "confidence": profile.get("confidence") if profile else (0.86 if all_rows else 0.58),
        "fileProfiles": file_profiles,
        "keyValues": _build_source_key_values(combined_text, file_profiles),
        "processingMeta": {
            "fileCount": len(files),
            "totalPageCount": total_page_count,
            "parsedTextChars": total_text_chars,
            "rowCount": len(all_rows),
            "outputMode": output_mode,
            "llmIntentUsed": bool(isinstance(llm_intent, dict) and llm_intent.get("_llmIntentUsed")),
            "llmIntent": intent_name or None,
            "llmIntentError": llm_intent_error[:120] if llm_intent_error else "",
        },
    }
    if multi_compare_table:
        table = multi_compare_table
        table_columns = table.get("columns", [])
    elif narrative_compare_table:
        table = narrative_compare_table
        table_columns = table.get("columns", [])
    else:
        table_columns = REFERENCE_GUIDELINE_COLUMNS if table_type in REFERENCE_TABLE_TYPES else (STANDARD_MARKET_PRICE_COLUMNS if table_type in STANDARD_MARKET_TABLE_TYPES else (TEXT_VENDOR_COMPARISON_COLUMNS if table_type == TEXT_VENDOR_COMPARISON_TABLE_TYPE else DEFAULT_COLUMNS))
        table_columns = prune_empty_columns(table_columns, all_rows)
        table = {
            "tableName": "기준서 항목 표" if table_type in REFERENCE_TABLE_TYPES else ("표준시장단가 표" if table_type in STANDARD_MARKET_TABLE_TYPES else ("서술형 업체별 단가 비교 요약" if table_type == TEXT_VENDOR_COMPARISON_TABLE_TYPE else "문서 표 후보")),
            "tableType": table_type,
            "columns": table_columns,
            "rows": all_rows,
        }
    issues = validate_rows(all_rows, table_type=table_type)
    business_drafts = _make_business_drafts(user_request, analysis, table, issues, combined_text, file_profiles)
    analysis["drafts"] = business_drafts
    analysis["reportDraft"] = business_drafts.get("report")
    analysis["meetingDraft"] = business_drafts.get("meeting")
    analysis["officialLetterDraft"] = business_drafts.get("officialLetter")
    if isinstance(table, dict):
        meta = dict(table.get("meta") or {})
        meta["drafts"] = business_drafts
        meta["draftPolicy"] = "업무 양식 미리보기용 초안. 원문에 없는 담당자/기한은 확인 필요 또는 미정으로 표시"
        table["meta"] = meta
    for parsed in parsed_files:
        table_metrics = ((parsed.get("parseMetrics") or {}).get("tables") or {}) if isinstance(parsed, dict) else {}
        if table_metrics.get("rowLimitReached"):
            issues.append({
                "rowIndex": None,
                "issueType": "TABLE_ROW_LIMIT_REACHED",
                "severity": "INFO",
                "fieldKey": "table",
                "fieldLabel": "표 추출",
                "message": f"표 행 제한({table_metrics.get('rowCount')}행)에 도달하여 {table_metrics.get('pagesRead')}페이지까지 표를 추출했습니다. 전체 표가 필요하면 PDF_TABLE_MAX_ROWS 값을 늘리세요.",
            })
    result_tables: List[Dict[str, Any]] = [table]

    llm_intent_used = bool(isinstance(llm_intent, dict) and llm_intent.get("_llmIntentUsed"))
    llm_intent_source = str((llm_intent or {}).get("_intentSource") or ("llm" if llm_intent_used else "rule")) if isinstance(llm_intent, dict) else "none"
    llm_summary_used = False
    llm_structure_used = False
    llm_summary_error = ""
    llm_structure_error = ""

    current_rows_for_llm: List[Dict[str, Any]] = []
    for result_table in result_tables:
        if isinstance(result_table, dict):
            current_rows_for_llm.extend(result_table.get("rows", []) or [])

    # 1) 표가 있는 경우: LLM은 표/금액을 다시 만들지 않고 분석 요약·검증 의견만 작성한다.
    if current_rows_for_llm and should_call_llm(user_request, combined_text, all_rows, len(files), table_type):
        cfg = get_llm_config()
        prompt = build_llm_grounded_analysis_prompt(user_request, analysis, table, issues, combined_text)
        try:
            llm_result = await call_llm_json(prompt, cfg)
            analysis, merged_issues = normalize_llm_analysis_only(llm_result, analysis, issues)
            # 업체 비교표는 파서/계산 결과가 기준이다. LLM이 임의 확인사항을 늘리는 것은 막는다.
            if table_type not in {MULTI_VENDOR_COMPARE_TABLE_TYPE, "PRICE_COMPARISON", "STANDARD_MARKET_PRICE_TABLE"}:
                issues = merged_issues
            llm_summary_used = True
            llm_used = True
            model_name = f"gemini:{cfg.model}+grounded-parser"
            prompt_version = "gemini-grounded-summary-v1"
            analysis.setdefault("processingMeta", {})["llmSummaryUsed"] = True
            analysis.setdefault("processingMeta", {})["llmModel"] = cfg.model
        except Exception as exc:  # noqa: BLE001
            llm_summary_error = str(exc)
            # 실패해도 분석 자체는 성공이다. 화면에는 오류 대신 fallback 상태만 남긴다.
            analysis.setdefault("keyValues", []).extend([
                {"label": "LLM 요약/검증", "value": "실패 → 파서 요약 유지"},
                {"label": "LLM 역할", "value": "표 추출·단가 계산은 Python 파서로 완료"},
            ])

    # 2) 표가 전혀 없는 경우에만 LLM 구조화를 시도한다. 이때도 원문 근거 검증을 통과한 행만 사용한다.
    if not current_rows_for_llm and should_call_llm(user_request, combined_text, all_rows, len(files), table_type):
        cfg = get_llm_config()
        prompt = build_llm_prompt(user_request, output_mode, template_id, combined_text, all_rows)
        try:
            llm_result = await call_llm_json(prompt, cfg)
            llm_analysis, llm_table, llm_issues = normalize_llm_result(llm_result, all_rows, table_type, combined_text, user_request)
            analysis = llm_analysis
            table = llm_table
            result_tables = [table]
            system_issues = validate_rows(table.get("rows", []), table_type=table.get("tableType", table_type))
            dedup: Dict[str, Dict[str, Any]] = {}
            for issue in [*llm_issues, *system_issues]:
                key = f"{issue.get('rowIndex')}|{issue.get('issueType')}|{issue.get('fieldKey')}|{issue.get('message')}"
                dedup[key] = issue
            issues = list(dedup.values())
            llm_structure_used = True
            llm_used = True
            model_name = f"gemini:{cfg.model}"
            prompt_version = "gemini-structure-v1"
            analysis.setdefault("processingMeta", {})["llmStructureUsed"] = True
            analysis.setdefault("processingMeta", {})["llmModel"] = cfg.model
        except Exception as exc:  # noqa: BLE001
            llm_structure_error = str(exc)
            analysis.setdefault("processingMeta", {})["llmStructureError"] = llm_structure_error[:120]

    # LLM이 analysis 객체를 보정해도 업무 양식 초안은 유지한다.
    if "business_drafts" in locals():
        analysis["drafts"] = business_drafts
        analysis["reportDraft"] = business_drafts.get("report")
        analysis["meetingDraft"] = business_drafts.get("meeting")
        analysis["officialLetterDraft"] = business_drafts.get("officialLetter")
        if isinstance(table, dict):
            meta = dict(table.get("meta") or {})
            meta["drafts"] = business_drafts
            table["meta"] = meta

    # 처리 방식 정보는 보고서 본문/keyValues에 섞지 않고 별도 메타에만 보관한다.
    analysis.setdefault("processingMeta", {}).update({
        "tableExtractionSource": "PyMuPDF/pdfplumber",
        "priceCalculationSource": "Python rule parser",
        "llmDirectTableGeneration": "unused" if current_rows_for_llm else ("used" if llm_structure_used else "skipped_or_failed"),
    })

    # 표 행이 없는데 LLM/규칙 파서가 값을 만들지 못한 경우, 빈 표를 유지하고 확인 필요만 표시한다.
    current_rows = []
    for result_table in result_tables:
        if isinstance(result_table, dict):
            current_rows.extend(result_table.get("rows", []) or [])
    if not current_rows and not any(issue.get("issueType") == "NO_BUSINESS_TABLE" for issue in issues):
        is_ref_table = any(isinstance(t, dict) and (t.get("tableType") in REFERENCE_TABLE_TYPES) for t in result_tables)
        issues.append({
            "rowIndex": None,
            "issueType": "NO_BUSINESS_TABLE",
            "severity": "INFO",
            "fieldKey": "table",
            "fieldLabel": "표 데이터",
            "message": "원문에서 기준서 표로 정리할 수 있는 기준 문장을 찾지 못했습니다." if is_ref_table else "원문에서 견적/단가표 형태의 품목·수량·단가 행은 확인되지 않았습니다. 근거 없는 표 행은 생성하지 않았습니다.",
        })
        analysis["summary"] = analysis.get("summary") or "문서 내용은 확인되었지만 업무 표 행은 추출되지 않았습니다."

    # confidence는 이슈가 많으면 보수적으로 낮춘다.
    if issues:
        analysis["confidence"] = min(float(analysis.get("confidence") or 0.7), 0.82)

    llm_used_final = bool(llm_used or llm_intent_used or llm_summary_used or llm_structure_used)
    llm_usage = {
        "used": llm_used_final,
        "intentAnalysis": {
            "used": llm_intent_used,
            "source": llm_intent_source,
            "status": "LLM 사용" if llm_intent_used else ("규칙 보정" if llm_intent else "미사용"),
        },
        "tableExtraction": {
            "used": False,
            "source": "PyMuPDF/pdfplumber",
            "status": "LLM 미사용 - 원문 파서 담당",
        },
        "priceCalculation": {
            "used": False,
            "source": "Python rule parser",
            "status": "LLM 미사용 - 코드 계산",
        },
        "summaryAnalysis": {
            "used": llm_summary_used,
            "source": f"gemini:{get_llm_config().model}" if llm_summary_used else "rule summary",
            "status": "LLM 사용" if llm_summary_used else ("파서 요약 유지" if current_rows else "미사용"),
        },
        "structureGeneration": {
            "used": llm_structure_used,
            "source": f"gemini:{get_llm_config().model}" if llm_structure_used else "parser/fallback",
            "status": "LLM 사용" if llm_structure_used else "미사용 또는 실패 시 파서 유지",
        },
    }

    return {
        "model": model_name,
        "promptVersion": prompt_version,
        "llmUsed": llm_used_final,
        "llmIntentUsed": llm_intent_used,
        "llmIntent": llm_intent,
        "llmUsage": llm_usage,
        "llmError": "" if llm_used_final or llm_intent else (llm_error or llm_intent_error),
        "analysis": analysis,
        "tables": result_tables,
        "issues": issues,
        "files": parsed_files,
        "parseLogs": all_parse_logs,
        "parse_logs": all_parse_logs,
        "parseMetrics": {
            "fileCount": len(files),
            "totalPages": total_page_count,
            "totalChars": total_text_chars,
            "ocrUsed": any(bool((pf.get("parseMetrics") or {}).get("ocrUsed")) for pf in parsed_files),
            "textForLlmTruncatedToChars": get_llm_config().context_chars,
        },
    }
