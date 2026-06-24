from __future__ import annotations

import csv
import io
import os
import re
import time
from typing import Any, Callable, Dict, List, Tuple

from app.services.unit_normalizer import clean_cell_text, enrich_row_units
from app.services.document_analyzer.table_utils import (
    _env_bool,
    _env_int,
    _limit_pages,
    rows_to_table,
    clean_number,
    compact_text,
    source_has_value,
    FIELD_ALIASES,
    BUSINESS_TABLE_REQUIRED_KEYS,
)


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def read_xlsx(content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return "", []

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    all_rows: List[List[Any]] = []
    text_lines: List[str] = []
    for ws in wb.worksheets[:5]:
        text_lines.append(f"[sheet {ws.title}]")
        sheet_rows: List[List[Any]] = []
        for row in ws.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                sheet_rows.append(values)
                all_rows.append(values)
                text_lines.append("\t".join(values))
            if len(sheet_rows) >= 200:
                break
    return "\n".join(text_lines), rows_to_table(all_rows)


# ---------------------------------------------------------------------------
# PDF – PyMuPDF
# ---------------------------------------------------------------------------

def _read_pdf_with_pymupdf(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, int | None, List[Dict[str, Any]], Dict[str, Any]]:
    started = time.perf_counter()
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        if log:
            log("error", "PyMuPDF import failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pymupdf", "error": str(exc)}

    try:
        doc = fitz.open(stream=content, filetype="pdf")
        page_count = int(doc.page_count or 0)
        pages_to_read = _limit_pages(page_count, "PDF_MAX_PAGES")
        log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
        parts: List[str] = []
        page_meta: List[Dict[str, Any]] = []
        empty_pages = 0
        total_chars = 0

        if log:
            log(
                "info",
                "PDF parse start",
                filename=filename,
                engine="PyMuPDF",
                text_layer=True,
                page_count=page_count,
                pages_to_read=pages_to_read,
                size_bytes=len(content),
            )

        for idx in range(pages_to_read):
            page = doc.load_page(idx)
            text = page.get_text("text", sort=True) or ""
            char_count = len(text)
            total_chars += char_count
            if not text.strip():
                empty_pages += 1
            parts.append(f"[page {idx + 1} / {page_count}]\n{text}".rstrip())
            page_meta.append({
                "page": idx + 1,
                "pageCount": page_count,
                "engine": "PyMuPDF",
                "status": "TEXT_EXTRACTED" if text.strip() else "NO_TEXT_LAYER",
                "charCount": char_count,
            })
            if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                log(
                    "info",
                    "PDF parse progress",
                    filename=filename,
                    page=f"{idx + 1}/{page_count}",
                    chars=total_chars,
                    empty_pages=empty_pages,
                )

        elapsed = round(time.perf_counter() - started, 3)
        truncated_by_limit = page_count > pages_to_read
        if log:
            log(
                "info",
                "PDF parse finish",
                filename=filename,
                engine="PyMuPDF",
                text_layer=True,
                page_count=page_count,
                pages_read=pages_to_read,
                chars=total_chars,
                empty_pages=empty_pages,
                elapsed_sec=elapsed,
                truncated_by_limit=truncated_by_limit,
            )
        return "\n\n".join(parts), page_count, page_meta, {
            "engine": "PyMuPDF",
            "noOcr": True,
            "pageCount": page_count,
            "pagesRead": pages_to_read,
            "charCount": total_chars,
            "emptyPages": empty_pages,
            "elapsedSec": elapsed,
            "truncatedByLimit": truncated_by_limit,
        }
    except Exception as exc:
        if log:
            log("error", "PyMuPDF parse failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "PyMuPDF", "noOcr": True, "error": str(exc)}


# ---------------------------------------------------------------------------
# PDF – pdfplumber text
# ---------------------------------------------------------------------------

def _read_pdf_text_with_pdfplumber(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, int | None, List[Dict[str, Any]], Dict[str, Any]]:
    try:
        import pdfplumber
    except Exception as exc:
        if log:
            log("error", "pdfplumber import failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pdfplumber", "error": str(exc)}

    started = time.perf_counter()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            pages_to_read = _limit_pages(page_count, "PDF_MAX_PAGES")
            log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
            parts: List[str] = []
            page_meta: List[Dict[str, Any]] = []
            empty_pages = 0
            total_chars = 0
            if log:
                log("info", "PDF fallback parse start", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_to_read=pages_to_read)
            for idx, page in enumerate(pdf.pages[:pages_to_read]):
                text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                char_count = len(text)
                total_chars += char_count
                if not text.strip():
                    empty_pages += 1
                parts.append(f"[page {idx + 1} / {page_count}]\n{text}".rstrip())
                page_meta.append({
                    "page": idx + 1,
                    "pageCount": page_count,
                    "engine": "pdfplumber",
                    "status": "TEXT_EXTRACTED" if text.strip() else "NO_TEXT_LAYER",
                    "charCount": char_count,
                })
                if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                    log("info", "PDF fallback parse progress", filename=filename, page=f"{idx + 1}/{page_count}", chars=total_chars, empty_pages=empty_pages)
            elapsed = round(time.perf_counter() - started, 3)
            if log:
                log("info", "PDF fallback parse finish", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_read=pages_to_read, chars=total_chars, empty_pages=empty_pages, elapsed_sec=elapsed, truncated_by_limit=page_count > pages_to_read)
            return "\n\n".join(parts), page_count, page_meta, {
                "engine": "pdfplumber",
                "noOcr": True,
                "pageCount": page_count,
                "pagesRead": pages_to_read,
                "charCount": total_chars,
                "emptyPages": empty_pages,
                "elapsedSec": elapsed,
                "truncatedByLimit": page_count > pages_to_read,
            }
    except Exception as exc:
        if log:
            log("error", "pdfplumber fallback parse failed", filename=filename, error=str(exc))
        return "", None, [], {"engine": "pdfplumber", "noOcr": True, "error": str(exc)}


# ---------------------------------------------------------------------------
# PDF – pdfplumber tables
# ---------------------------------------------------------------------------

def _read_pdf_tables_with_pdfplumber(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
    source_text: str = "",
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    if not _env_bool("PDF_EXTRACT_TABLES", True):
        if log:
            log("info", "PDF table extraction skipped", filename=filename, reason="PDF_EXTRACT_TABLES=false")
        return rows, {"enabled": False, "rowCount": 0}

    try:
        import pdfplumber
    except Exception as exc:
        if log:
            log("warning", "pdfplumber table import failed", filename=filename, error=str(exc))
        return rows, {"enabled": True, "error": str(exc), "rowCount": 0}

    started = time.perf_counter()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            pages_to_read = _limit_pages(page_count, "PDF_TABLE_MAX_PAGES")
            configured_max_rows = _env_int("PDF_TABLE_MAX_ROWS", 5000)
            max_rows = configured_max_rows if configured_max_rows > 0 else 10**9
            log_every = _env_int("PDF_LOG_EVERY_PAGES", 0)
            if log:
                log("info", "PDF table extraction start", filename=filename, engine="pdfplumber", text_layer=True, page_count=page_count, pages_to_read=pages_to_read)
            for idx, page in enumerate(pdf.pages[:pages_to_read]):
                table_count = 0
                for table in page.extract_tables() or []:
                    table_count += 1
                    parsed = rows_to_table(table, source_text=source_text, filename=filename)
                    rows.extend(parsed)
                    if len(rows) >= max_rows:
                        rows = rows[:max_rows]
                        if log:
                            log("warning", "PDF table row limit reached", filename=filename, row_limit=max_rows, page=idx + 1)
                        elapsed = round(time.perf_counter() - started, 3)
                        return rows, {"enabled": True, "engine": "pdfplumber", "pageCount": page_count, "pagesRead": idx + 1, "rowCount": len(rows), "elapsedSec": elapsed, "rowLimitReached": True}
                if log and log_every > 0 and ((idx + 1) == 1 or (idx + 1) % log_every == 0 or (idx + 1) == pages_to_read):
                    log("info", "PDF table extraction progress", filename=filename, page=f"{idx + 1}/{page_count}", rows=len(rows), tables_on_page=table_count)
            elapsed = round(time.perf_counter() - started, 3)
            if log:
                log("info", "PDF table extraction finish", filename=filename, engine="pdfplumber", page_count=page_count, pages_read=pages_to_read, rows=len(rows), elapsed_sec=elapsed, truncated_by_limit=page_count > pages_to_read)
            return rows, {"enabled": True, "engine": "pdfplumber", "pageCount": page_count, "pagesRead": pages_to_read, "rowCount": len(rows), "elapsedSec": elapsed, "truncatedByLimit": page_count > pages_to_read}
    except Exception as exc:
        if log:
            log("warning", "PDF table extraction failed", filename=filename, error=str(exc))
        return rows, {"enabled": True, "error": str(exc), "rowCount": len(rows)}


# ---------------------------------------------------------------------------
# read_pdf (main entry – delegates OCR to ocr_engine)
# ---------------------------------------------------------------------------

def read_pdf(
    content: bytes,
    filename: str = "",
    log: Callable[..., None] | None = None,
) -> Tuple[str, List[Dict[str, Any]], int | None, List[Dict[str, Any]], Dict[str, Any]]:
    # Import here to avoid circular at module load time
    from app.services.document_analyzer.ocr_engine import (
        _ocr_enabled,
        _ocr_min_text_chars,
        _read_pdf_with_ocr_fallback,
    )

    preferred = os.getenv("PDF_TEXT_ENGINE", "pymupdf").strip().lower()
    if preferred not in {"pymupdf", "pdfplumber"}:
        preferred = "pymupdf"

    if preferred == "pdfplumber":
        text, page_count, page_meta, text_metrics = _read_pdf_text_with_pdfplumber(content, filename=filename, log=log)
        if len(text.strip()) < 80:
            fallback_text, fallback_count, fallback_pages, fallback_metrics = _read_pdf_with_pymupdf(content, filename=filename, log=log)
            if len(fallback_text.strip()) > len(text.strip()):
                text, page_count, page_meta, text_metrics = fallback_text, fallback_count, fallback_pages, fallback_metrics
    else:
        text, page_count, page_meta, text_metrics = _read_pdf_with_pymupdf(content, filename=filename, log=log)
        if len(text.strip()) < 80:
            fallback_text, fallback_count, fallback_pages, fallback_metrics = _read_pdf_text_with_pdfplumber(content, filename=filename, log=log)
            if len(fallback_text.strip()) > len(text.strip()):
                text, page_count, page_meta, text_metrics = fallback_text, fallback_count, fallback_pages, fallback_metrics

    rows, table_metrics = _read_pdf_tables_with_pdfplumber(content, filename=filename, log=log, source_text=text)

    ocr_metrics: Dict[str, Any] = {"enabled": _ocr_enabled(), "ocrUsed": False}
    should_ocr = _ocr_enabled() and (_env_bool("OCR_FORCE", False) or (len(text.strip()) < _ocr_min_text_chars() and not rows))
    if should_ocr:
        ocr_text, ocr_rows, ocr_pages, ocr_metrics = _read_pdf_with_ocr_fallback(content, filename=filename, page_count=page_count, log=log)
        if ocr_text.strip():
            text = (text + "\n\n" + ocr_text).strip() if text.strip() else ocr_text
        if ocr_rows:
            rows = ocr_rows if not rows else [*rows, *ocr_rows]
        if ocr_pages:
            page_meta = page_meta or []
            page_meta.extend(ocr_pages)

    metrics = {"text": text_metrics, "tables": table_metrics, "ocr": ocr_metrics, "ocrUsed": bool(ocr_metrics.get("ocrUsed"))}
    return text, rows, page_count, page_meta, metrics


# ---------------------------------------------------------------------------
# DOCX / plain text
# ---------------------------------------------------------------------------

def read_docx(content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_rows: List[List[Any]] = []
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                raw_rows.append(values)
        table_lines = ["\t".join(map(str, row)) for row in raw_rows]
        return "\n".join(paragraphs + table_lines), rows_to_table(raw_rows)
    except Exception:
        return "", []


def decode_text(content: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def parse_delimited_text(text: str) -> List[Dict[str, Any]]:
    lines = [line for line in text.splitlines() if line.strip()]
    if len(lines) < 2:
        return []

    sample = "\n".join(lines[:30])
    delimiter_candidates = ["\t", ",", "|", ";"]
    delimiter = max(delimiter_candidates, key=lambda d: sample.count(d))
    if sample.count(delimiter) < 2:
        spaced_rows = [re.split(r"\s{2,}", line.strip()) for line in lines if len(re.split(r"\s{2,}", line.strip())) >= 3]
        return rows_to_table(spaced_rows) if len(spaced_rows) >= 2 else []

    reader = csv.reader(lines, delimiter=delimiter)
    raw_rows = [row for row in reader]
    return rows_to_table(raw_rows)


def infer_rows_from_text(text: str, filename: str) -> List[Dict[str, Any]]:
    """텍스트에서 업무 표 행을 보수적으로 추정한다."""
    # Import lazily to avoid circular deps
    from app.services.document_analyzer.doc_profiler import (
        is_text_only_vendor_comparison_report,
        is_reference_or_guideline_document,
        is_narrative_document,
    )
    from app.services.document_analyzer.text_extractor import extract_reference_guideline_rows
    from app.services.document_analyzer.row_filters import filter_grounded_rows, is_business_row_supported

    if is_text_only_vendor_comparison_report(text):
        return []
    if is_reference_or_guideline_document(text):
        return extract_reference_guideline_rows(text, user_request="")
    if is_narrative_document(text):
        return []

    table_rows = parse_delimited_text(text)
    table_rows = filter_grounded_rows(table_rows, text)
    if table_rows:
        return table_rows

    rows: List[Dict[str, Any]] = []
    line_pattern = re.compile(
        r"(?P<item>[가-힣A-Za-z0-9()\[\]#/+\-.\s]{2,40})\s+"
        r"(?P<spec>[A-Za-z0-9가-힣ΦØ㎡㎥㎜㎝㎏./*xX\-\s]{0,30})\s+"
        r"(?P<qty>[0-9]+(?:\.[0-9]+)?)\s*"
        r"(?P<unit>EA|PCS|SET|BOX|LOT|M|m|개|식|본|장|대|조|롤|포|kg|KG|㎡|㎥)?\s+"
        r"(?P<unit_price>[0-9]{1,3}(?:,[0-9]{3})+|[0-9]{3,})\s+"
        r"(?P<amount>[0-9]{1,3}(?:,[0-9]{3})+|[0-9]{3,})"
    )
    for line in text.splitlines()[:1000]:
        if not line.strip():
            continue
        m = line_pattern.search(line)
        if not m:
            continue
        row = enrich_row_units({
            "vendor_name": "",
            "item_name": m.group("item").strip(),
            "spec": m.group("spec").strip(),
            "quantity": clean_number(m.group("qty")),
            "unit": (m.group("unit") or "").strip(),
            "unit_price": clean_number(m.group("unit_price")),
            "amount": clean_number(m.group("amount")),
            "remark": "텍스트 행에서 추정",
        })
        if is_business_row_supported(row, text):
            rows.append(row)
        if len(rows) >= 100:
            break
    return rows
